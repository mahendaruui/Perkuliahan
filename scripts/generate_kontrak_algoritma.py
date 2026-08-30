#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generator Dokumen Kontrak Perkuliahan Resmi (Sesuai Siakad UUI)
Mata Kuliah: Algoritma dan Pemrograman (IFR206 - 4 SKS)
Program Studi S1 Informatika - Fakultas Sains dan Teknologi
Universitas Ubudiyah Indonesia (UUI)
Semester Ganjil T.A. 2026/2027 - Kelas A
Dosen Pengampu: Mahendar Dwi Payana, S.ST., M.T.
"""

import os
import shutil
import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
OUTPUT_DIR = os.path.join(REPO_ROOT, "kontrak-kuliah")
os.makedirs(OUTPUT_DIR, exist_ok=True)

DOCX_MAIN = os.path.join(OUTPUT_DIR, "Kontrak_Kuliah_Algoritma_Pemrograman_IFR206_Mahendar_Dwi_Payana.docx")
DOCX_LATEST = os.path.join(OUTPUT_DIR, "Kontrak_Kuliah_Algoritma_Pemrograman_IFR206_Latest.docx")

PDF_MAIN = os.path.join(OUTPUT_DIR, "Kontrak_Kuliah_Algoritma_Pemrograman_IFR206_Mahendar_Dwi_Payana.pdf")
PDF_LATEST = os.path.join(OUTPUT_DIR, "Kontrak_Kuliah_Algoritma_Pemrograman_IFR206_Latest.pdf")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('shd'):
            tcPr.remove(child)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = parse_xml(f'<{tag} {nsdecls("w")} w:val="{edge_data.get("val", "single")}" w:sz="{edge_data.get("sz", "4")}" w:space="0" w:color="{edge_data.get("color", "CCCCCC")}"/>')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def build_docx_kontrak():
    print("-> Membangun Dokumen Word (.docx) Kontrak Kuliah Algoritma Pemrograman...")
    doc = docx.Document()
    
    # Margin Halaman Standar A4
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    # 1. KOP SURAT RESMI UUI
    kop_table = doc.add_table(rows=1, cols=2)
    kop_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kop_table.autofit = False
    
    # Col 1: Logo UUI
    logo_cell = kop_table.rows[0].cells[0]
    logo_cell.width = Cm(2.2)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    logo_path = os.path.join(REPO_ROOT, "docs", "public", "uuilogo.png")
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Cm(1.8))
        
    # Col 2: Teks Kop Kampus
    text_cell = kop_table.rows[0].cells[1]
    text_cell.width = Cm(13.8)
    text_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    p_kop = text_cell.paragraphs[0]
    p_kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kop.paragraph_format.space_after = Pt(2)
    p_kop.paragraph_format.line_spacing = 1.1
    
    r1 = p_kop.add_run("YAYASAN UBUDIYAH INDONESIA\nUNIVERSITAS UBUDIYAH INDONESIA\n")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    r2 = p_kop.add_run("FAKULTAS SAINS DAN TEKNOLOGI - PROGRAM STUDI S1 INFORMATIKA\n")
    r2.bold = True
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
    
    r3 = p_kop.add_run("Jl. Alue Naga, Desa Tibang, Kec. Syiah Kuala, Kota Banda Aceh, Aceh 23114\nLaman: https://uui.ac.id | Pos-el: informatika@uui.ac.id")
    r3.font.name = "Calibri"
    r3.font.size = Pt(8.5)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    # Garis Pembatas
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(4)
    p_line.paragraph_format.space_after = Pt(14)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_line = p_line.add_run("═" * 60)
    r_line.font.name = "Calibri"
    r_line.font.size = Pt(11)
    r_line.font.bold = True
    r_line.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    # JUDUL DOKUMEN
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    r_t1 = p_title.add_run("KONTRAK PERKULIAHAN")
    r_t1.bold = True
    r_t1.font.name = "Calibri"
    r_t1.font.size = Pt(15)
    r_t1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_t2 = p_sub.add_run("SEMESTER GANJIL TAHUN AKADEMIK 2026/2027")
    r_t2.bold = True
    r_t2.font.name = "Calibri"
    r_t2.font.size = Pt(11.5)
    r_t2.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    def add_sec_heading(num_str, title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r_n = p.add_run(num_str + " ")
        r_n.bold = True
        r_n.font.name = "Calibri"
        r_n.font.size = Pt(11.5)
        r_n.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        r_t = p.add_run(title_text)
        r_t.bold = True
        r_t.font.name = "Calibri"
        r_t.font.size = Pt(11.5)
        r_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return p

    def add_body_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.name = "Calibri"
            r_pre.font.size = Pt(10.5)
            r_pre.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        r_b = p.add_run(text)
        r_b.font.name = "Calibri"
        r_b.font.size = Pt(10.5)
        r_b.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    def add_bullet_p(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r_pre = p.add_run(bold_prefix + ": ")
            r_pre.bold = True
            r_pre.font.name = "Calibri"
            r_pre.font.size = Pt(10.5)
            r_pre.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        r_b = p.add_run(text)
        r_b.font.name = "Calibri"
        r_b.font.size = Pt(10.5)
        r_b.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    # 1. IDENTITAS MATA KULIAH
    add_sec_heading("I.", "IDENTITAS MATA KULIAH")
    id_data = [
        ["1. Nama Mata Kuliah", "Algoritma Pemrograman"],
        ["2. Kode Mata Kuliah / Bobot SKS", "IFR206 / 4 SKS (Teori: 2 SKS, Praktik: 2 SKS)"],
        ["3. Program Studi / Jenjang", "S1 - Informatika / Sarjana (Strata-1)"],
        ["4. Fakultas / Universitas", "Fakultas Sains dan Teknologi / Universitas Ubudiyah Indonesia"],
        ["5. Kurikulum / Rumpun MK", "Kurikulum 2023 (OBE) / Mata Kuliah Wajib Keahlian"],
        ["6. Semester / Periode Akademik", "Semester I (Satu) / 2026 Ganjil"],
        ["7. Nama Kelas / Sistem Kuliah", "Kelas A / Reguler"],
        ["8. Kapasitas Kelas / Jumlah Peserta", "60 Mahasiswa / 1 Mahasiswa (Update Siakad)"],
        ["9. Dosen Pengampu Utama", "Mahendar Dwi Payana, S.ST., M.T. (NIDN: 1331108701)"],
        ["10. Ruang Perkuliahan & Praktikum", "Lab. Rekayasa Perangkat Lunak & Sistem Informasi UUI"]
    ]
    
    table_id = doc.add_table(rows=len(id_data), cols=2)
    table_id.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_id.autofit = False
    
    for idx, (label, val) in enumerate(id_data):
        row = table_id.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Cm(6.5)
        c2.width = Cm(9.5)
        
        bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(c1, bg_color)
        set_cell_background(c2, bg_color)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c2, top=60, bottom=60, left=100, right=100)
        set_cell_border(c1, bottom={"sz": "2", "color": "E2E8F0"})
        set_cell_border(c2, bottom={"sz": "2", "color": "E2E8F0"})
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(val)
        r2.font.name = "Calibri"
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # 2. MANFAAT MATA KULIAH
    add_sec_heading("II.", "MANFAAT MATA KULIAH")
    add_body_p("Mata kuliah Algoritma dan Pemrograman (IFR206) merupakan fondasi utama (core foundation) bagi seluruh kurikulum di Program Studi Informatika. Penguasaan algoritma menanamkan kemampuan berpikir komputasional (computational thinking), dekomposisi masalah, pengenalan pola logis, dan penalaran sistematis. Mata kuliah ini menjembatani kemampuan analitis mahasiswa dalam merumuskan solusi logis ke dalam instruksi program komputer yang terstruktur, efisien, bebas dari cacat logika, dan siap dilanjutkan ke tingkat pemrograman lanjut (Struktur Data, OOP, dan Rekayasa Web).")

    # 3. DESKRIPSI SINGKAT MATA KULIAH
    add_sec_heading("III.", "DESKRIPSI MATA KULIAH")
    add_body_p("Mata kuliah ini membahas prinsip dasar logika algoritma, notasi representasi algoritma (Flowchart standar ISO dan Pseudocode), sistem tipe data primitif dan komposit, alokasi memori, operator aritmatika/relasional/logika, struktur kontrol alur percabangan (if-else, switch-case), struktur kontrol perulangan (for, while, do-while, nested loop), struktur data larik (Array 1D dan Multidimensi), manipulasi teks (String), rekayasa fungsi dan prosedur (parameter passing by value/reference), teknik rekursi & call stack memory, serta algoritma pencarian (Linear & Binary Search) dan pengurutan data (Bubble, Selection, Insertion Sort) dengan analisis kompleksitas komputasi dasar (Big-O Notation).")

    # 4. CAPAIAN PEMBELAJARAN (CPL & CPMK)
    add_sec_heading("IV.", "CAPAIAN PEMBELAJARAN LULUSAN & MATA KULIAH (CPL & CPMK)")
    add_bullet_p("CPL01 (Pengetahuan Dasar)", "Memiliki pengetahuan komprehensif tentang teori, prinsip, dan konsep dasar algoritma dan informatika.")
    add_bullet_p("CPL03 (Keterampilan Kerja Umum)", "Mampu berpikir logis, analitis, dan sistematis dalam merancang solusi berbasis komputasi.")
    add_bullet_p("CPL04 (Keterampilan Khusus)", "Mampu menyajikan solusi terprogram atas permasalahan nyata di industri dan masyarakat.")
    add_bullet_p("CPL08 (Sikap & Etika)", "Menjunjung tinggi etika akademik, kejujuran kode program (clean code), dan anti-plagiarisme.")

    add_body_p("Capaian Pembelajaran Mata Kuliah (CPMK):", bold_prefix="Capaian Spesifik:")
    add_bullet_p("CPMK-1", "Mampu menjelaskan hakikat logika algoritma, notasi flowchart, pseudocode, serta sistem tipe data dan alokasi memori komputer.")
    add_bullet_p("CPMK-2", "Mampu merancang dan mengimplementasikan struktur kontrol percabangan (decision making) dan perulangan (looping) kompleks.")
    add_bullet_p("CPMK-3", "Mampu merekayasa struktur data larik (Array 1D dan Matriks Multidimensi) serta operasi manipulasi string.")
    add_bullet_p("CPMK-4", "Mampu menerapkan modularisasi kode menggunakan fungsi, prosedur, dan teknik pemecahan masalah rekursif.")
    add_bullet_p("CPMK-5", "Mampu menganalisis, merancang, dan membandingkan efisiensi algoritma Searching dan Sorting.")
    add_bullet_p("CPMK-6", "Mampu membangun dan mendokumentasikan mini project aplikasi berbasis algoritma terstruktur secara mandiri dan tim.")

    # 5. MATRIKS RENCANA PERKULIAHAN (16 PERTEMUAN)
    add_sec_heading("V.", "MATRIKS RENCANA PERKULIAHAN (16 PERTEMUAN)")
    
    rps_headers = ["Prt", "Kemampuan Akhir (Sub-CPMK)", "Bahan Kajian / Pokok Bahasan", "Bentuk Pembelajaran", "Metode Evaluasi"]
    rps_data = [
        ["1", "Memahami konsep algoritma & notasi", "Kontrak Kuliah, Logika Komputasi, Flowchart ISO, Pseudocode", "Kuliah Interaktif & Diskusi", "Aktivitas Partisipatif"],
        ["2", "Mampu mengelola tipe data & memori", "Variabel, Tipe Data Primitif, Konstanta, Type Casting, RAM", "Kuliah & Praktikum Lab", "Tugas Praktikum 1"],
        ["3", "Mampu mengevaluasi ekspresi logika", "Operator Aritmatika, Relasional, Logika Boolean, Bitwise", "Problem-Based Learning", "Tugas Praktikum 2"],
        ["4", "Mampu merancang percabangan bersyarat", "Struktur Kontrol if, if-else, nested-if, operator ternary", "Live Coding & Studi Kasus", "Tugas Praktikum 3"],
        ["5", "Mampu merekayasa menu switch-case", "Struktur switch-case, fall-through, navigasi menu CLI", "Praktikum Terbimbing", "Quiz 1"],
        ["6", "Mampu mengendalikan perulangan (looping)", "Counted loop (for), uncounted (while, do-while), break/continue", "Praktikum Lab", "Tugas Praktikum 4"],
        ["7", "Mampu memanipulasi array 1 dimensi", "Larik 1D, Indeks Elemen, Input/Output, Min/Max/Rata-rata", "Praktikum Intensif", "Tugas Praktikum 5"],
        ["8", "EVALUASI CAPAIAN PERTEMUAN 1 s.d. 7", "UJIAN TENGAH SEMESTER (UTS) - TEORI & LIVE CODING", "Ujian Tertulis & Praktik Lab", "UTS (30%)"],
        ["9", "Mampu mengolah array multidimensi", "Matriks 2D (Baris-Kolom), Operasi Matriks, String Processing", "Praktikum Laboratorium", "Tugas Praktikum 6"],
        ["10", "Mampu menyusun fungsi & prosedur", "Modularitas, Parameter pass-by-value/reference, Scope Variabel", "Live Coding & Modul", "Tugas Praktikum 7"],
        ["11", "Mampu menerapkan teknik rekursi", "Rekursif, Base Case, Recursive Case, Call Stack, Overflow", "Simulasi Trace Table", "Quiz 2"],
        ["12", "Mampu mengimplementasikan searching", "Algoritma Linear Search, Binary Search, Perbandingan O(n)/O(log n)", "Eksperimen Algoritma", "Tugas Praktikum 8"],
        ["13", "Mampu merekayasa algoritma sorting", "Bubble Sort, Selection Sort, Insertion Sort, Trace Array", "Praktikum Sorting", "Tugas Praktikum 9"],
        ["14", "Mampu merancang struktur aplikasi mini", "Arsitektur Modular Aplikasi CLI, Validasi Input, Error Trapping", "Praktikum Proyek", "Aktivitas Partisipatif"],
        ["15", "Mampu mengembangkan proyek mini mandiri", "Pengembangan Capstone Mini Project: Sistem Manajemen Terstruktur", "PjBL & Asistensi Terbimbing", "Asistensi Proyek"],
        ["16", "EVALUASI CAPAIAN KOMPREHENSIF PROYEK", "UJIAN AKHIR SEMESTER (UAS) - DEMO PROGRAM & PRESENTASI", "Evaluasi Capstone & Uji Program", "UAS (40%)"]
    ]
    
    table_rps = doc.add_table(rows=len(rps_data) + 1, cols=5)
    table_rps.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_rps.autofit = False
    
    col_widths = [Cm(1.1), Cm(4.2), Cm(5.6), Cm(3.2), Cm(1.9)]
    
    # Header Row
    hdr_cells = table_rps.rows[0].cells
    for i, title in enumerate(rps_headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=60, right=60)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(9.0)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    for r_idx, row_values in enumerate(rps_data):
        row = table_rps.rows[r_idx + 1]
        is_highlight = row_values[0] in ["8", "16"]
        bg = "EFF6FF" if is_highlight else ("F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
        
        for c_idx, val in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            set_cell_border(cell, bottom={"sz": "2", "color": "CBD5E1"})
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            
            if c_idx in [0, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(8.5)
            if is_highlight:
                r.bold = True
                r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            else:
                r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # 6. KOMPONEN EVALUASI & MAPPING NILAI SIAKAD
    add_sec_heading("VI.", "KOMPONEN EVALUASI & SKALA PENILAIAN (SIAKAD UUI)")
    add_body_p("Komponen dan bobot evaluasi pembelajaran disesuaikan 100% dengan konfigurasi resmi Portal Siakad Universitas Ubudiyah Indonesia sebagai berikut:")
    
    # Table Evaluasi Siakad
    eval_siakad = [
        ["No.", "Metode Evaluasi", "Jenis Evaluasi", "Bobot Evaluasi", "Syarat Lulus"],
        ["1", "TUGAS", "Kognitif/Pengetahuan - Tugas", "20%", "-"],
        ["2", "QUIZ", "Kognitif/Pengetahuan - Quiz", "5%", "-"],
        ["3", "UTS", "Kognitif/Pengetahuan - Ujian Tengah Semester", "30%", "-"],
        ["4", "UAS", "Kognitif/Pengetahuan - Ujian Akhir Semester", "40%", "-"],
        ["5", "AKTIVITAS PARTISIPATIF", "Aktivitas Partisipatif (Kehadiran & Diskusi)", "5%", "-"],
        ["", "Total Persentase Komponen Evaluasi", "", "100%", ""]
    ]
    
    table_eval = doc.add_table(rows=len(eval_siakad), cols=5)
    table_eval.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_eval.autofit = False
    col_w_eval = [Cm(1.0), Cm(3.2), Cm(7.0), Cm(2.8), Cm(2.0)]
    
    for r_idx, row_values in enumerate(eval_siakad):
        row = table_eval.rows[r_idx]
        is_hdr = (r_idx == 0)
        is_tot = (r_idx == len(eval_siakad) - 1)
        
        bg = "0B4F8A" if is_hdr else ("E2E8F0" if is_tot else ("F8FAFC" if r_idx % 2 == 1 else "FFFFFF"))
        
        for c_idx, val in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.width = col_w_eval[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=70, right=70)
            set_cell_border(cell, bottom={"sz": "2", "color": "94A3B8" if (is_hdr or is_tot) else "E2E8F0"})
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if c_idx in [0, 3, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9.0)
            if is_hdr:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif is_tot:
                r.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            else:
                r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Keterangan Syarat Lulus
    p_ket = doc.add_paragraph()
    p_ket.paragraph_format.space_before = Pt(4)
    p_ket.paragraph_format.space_after = Pt(8)
    r_k1 = p_ket.add_run("Keterangan: ")
    r_k1.bold = True
    r_k1.font.name = "Calibri"
    r_k1.font.size = Pt(9.0)
    r_k2 = p_ket.add_run("Syarat Lulus Mata Kuliah adalah pemenuhan seluruh komponen nilai wajib. Mahasiswa yang tidak memiliki salah satu komponen nilai utama akan dinyatakan tidak lulus.")
    r_k2.font.name = "Calibri"
    r_k2.font.size = Pt(9.0)
    r_k2.font.italic = True

    add_body_p("Tabel Mapping Konversi Grade Nilai Akhir Siakad UUI:", bold_prefix="Tabel Konversi Grade: ")
    
    # Table Mapping Grade Siakad
    grade_mapping = [
        ["Grade", "Bobot", "Nilai Bawah", "Nilai Atas", "Status Kelulusan"],
        ["A", "4,00", "80,00", "100,00", "Sangat Memuaskan / Lulus"],
        ["B", "3,00", "60,00", "79,99", "Baik / Lulus"],
        ["C", "2,00", "40,00", "59,99", "Cukup / Batas Minimal Lulus"],
        ["D", "1,00", "20,00", "39,99", "Kurang / Wajib Mengulang"],
        ["E", "0,00", "0,00", "19,99", "Gagal / Tidak Lulus"]
    ]
    
    table_grade = doc.add_table(rows=len(grade_mapping), cols=5)
    table_grade.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_grade.autofit = False
    col_w_grade = [Cm(2.0), Cm(2.2), Cm(3.2), Cm(3.2), Cm(5.4)]
    
    for r_idx, row_values in enumerate(grade_mapping):
        row = table_grade.rows[r_idx]
        is_hdr = (r_idx == 0)
        bg = "0284C7" if is_hdr else ("F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
        
        for c_idx, val in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.width = col_w_grade[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            set_cell_border(cell, bottom={"sz": "2", "color": "E2E8F0"})
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if c_idx in [0, 1, 2, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9.0)
            if is_hdr:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # 7. TATA TERTIB & NORMA AKADEMIK
    add_sec_heading("VII.", "TATA TERTIB & NORMA AKADEMIK PERKULIAHAN")
    add_bullet_p("1. Kehadiran & Ketepatan Waktu", "Mahasiswa wajib hadir tepat waktu sesuai jadwal perkuliahan. Toleransi keterlambatan maksimal adalah 15 menit. Keterlambatan melebihi batas tersebut tetap diperkenankan mengikuti perkuliahan namun presensi dianggap alpa.")
    add_bullet_p("2. Batas Minimal Kehadiran", "Kehadiran minimal mahasiswa adalah 75% dari total 16 pertemuan riil sebagai syarat mutlak mengikuti Ujian Akhir Semester (UAS). Mahasiswa dengan kehadiran di bawah 75% otomatis dinyatakan tidak berhak mengikuti UAS.")
    add_bullet_p("3. Etika Berbusana & Sikap", "Berpakaian sopan, rapi, berkerah (bukan kaos oblong polos), dan mengenakan sepatu di lingkungan kampus dan laboratorium. Menjaga etika berbicara, sopan santun akademik, serta mematikan/mengatur nada dering senyap (silent mode) pada perangkat komunikasi selama perkuliahan.")
    add_bullet_p("4. Fasilitas & Praktikum Lab", "Setiap mahasiswa wajib membawa laptop kerja pribadi yang telah terinstal environment pemrograman (C / C++ / Python / Compiler IDE, VS Code, dan Git). Dilarang merusak, mengubah konfigurasi jaringan lab, atau mengotori ruang laboratorium.")
    add_bullet_p("5. Penyerahan Tugas & Sanksi Keterlambatan", "Tugas mandiri dan kuis lab wajib diunggah tepat waktu melalui LMS Siakad UUI / Repositori GitHub resmi kelas. Keterlambatan pengumpulan tugas dikenakan penalti pemotongan nilai sebesar 10% per hari (maksimal toleransi 3 hari).")
    add_bullet_p("6. Integritas Akademik & Anti-Plagiarisme", "Kecurangan akademik dalam bentuk apa pun (plagiarisme logika program, menyalin kode sumber teman, menggunakan joki tugas, atau kecurangan saat ujian) dikenakan sanksi tegas pembatalan nilai (Nilai E otomatis untuk mata kuliah ini) dan dilaporkan ke Komite Disiplin Universitas Ubudiyah Indonesia.")

    # 8. PUSTAKA & REFERENSI
    add_sec_heading("VIII.", "SUMBER BELAJAR & DAFTAR PUSTAKA")
    add_bullet_p("Buku Ajar Utama", "Payana, M. D., Wibawa, M. B., & Tim Informatika. (2026). Buku Ajar Algoritma dan Pemrograman Terstruktur. Universitas Ubudiyah Indonesia Press.")
    add_bullet_p("Referensi Internasional", "Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2022). Introduction to Algorithms (4th ed.). MIT Press.")
    add_bullet_p("Logika Pemrograman", "Munir, R., & Lidya, L. (2016). Algoritma dan Pemrograman dalam Bahasa Pascal dan C/C++. Informatika Bandung.")
    add_bullet_p("Standar Koding", "Kernighan, B. W., & Ritchie, D. M. (1988). The C Programming Language (2nd ed.). Prentice Hall.")

    # 9. LEMBAR PERSETUJUAN & PENGESAHAN
    add_sec_heading("IX.", "LEMBAR PENGESAHAN KONTRAK PERKULIAHAN")
    add_body_p("Kontrak perkuliahan ini disepakati bersama secara sadar dan penuh tanggung jawab antara Dosen Pengampu Mata Kuliah dan Perwakilan Mahasiswa Kelas A pada pertemuan pertama perkuliahan Semester Ganjil Tahun Akademik 2026/2027 di Universitas Ubudiyah Indonesia.")

    # Tanda Tangan Table
    doc.add_paragraph().paragraph_format.space_before = Pt(8)
    
    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_loc.paragraph_format.space_after = Pt(8)
    r_loc = p_loc.add_run("Banda Aceh, 1 September 2026")
    r_loc.font.name = "Calibri"
    r_loc.font.size = Pt(10.0)
    
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    # Row 0: Labels
    c_mhs_top, c_dos_top = sig_table.rows[0].cells[0], sig_table.rows[0].cells[1]
    c_mhs_top.width = Cm(8.0)
    c_dos_top.width = Cm(8.0)
    
    p_mt = c_mhs_top.paragraphs[0]
    p_mt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_mt.add_run("Perwakilan Mahasiswa / Komti Kelas A,\n\n\n\n\n( _____________________________ )\nNPM. ....................................")
    r.font.name = "Calibri"
    r.font.size = Pt(10.0)
    
    p_dt = c_dos_top.paragraphs[0]
    p_dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_dt.add_run("Dosen Pengampu Mata Kuliah,\n\n\n\n\n( Mahendar Dwi Payana, S.ST., M.T. )\nNIDN. 1331108701")
    r.font.name = "Calibri"
    r.font.size = Pt(10.0)
    r.bold = True
    
    # Row 1: Mengetahui Kaprodi
    c_kpd = sig_table.rows[1].cells[0]
    c_kpd_right = sig_table.rows[1].cells[1]
    c_kpd.width = Cm(16.0)
    c_kpd.merge(c_kpd_right)
    
    p_kp = c_kpd.paragraphs[0]
    p_kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kp.paragraph_format.space_before = Pt(14)
    r = p_kp.add_run("Mengetahui,\nKetua Program Studi S1 Informatika UUI\n\n\n\n\n( M. Bayu Wibawa, S.Kom., MMSI )\nNIDN. ....................................")
    r.font.name = "Calibri"
    r.font.size = Pt(10.0)

    doc.save(DOCX_MAIN)
    shutil.copyfile(DOCX_MAIN, DOCX_LATEST)
    print(f"✅ Dokumen Kontrak Kuliah Word berhasil disimpan: {DOCX_MAIN}")
    print(f"✅ Salinan rilis terbaru: {DOCX_LATEST}")

# PDF Builder
class SimplePDF:
    def __init__(self):
        self.pages = []
        self.page_width = 595.28 # A4
        self.page_height = 841.89 # A4
        self.current_stream = []
        self.current_page_num = 0
        self.fonts = {"F1": "Helvetica", "F2": "Helvetica-Bold", "F3": "Helvetica-Oblique", "F4": "Times-Roman", "F5": "Times-Bold"}
        
    def new_page(self):
        if self.current_stream:
            self.pages.append(b"\n".join(self.current_stream))
            self.current_stream = []
        self.current_page_num += 1
        
        # Header line
        self.draw_line(40, 805, 555, 805, stroke_color=(0.8, 0.85, 0.9), width=0.5)
        self.draw_text("Universitas Ubudiyah Indonesia | Kontrak Perkuliahan Algoritma Pemrograman (IFR206)", 40, 810, font="F3", size=8, color=(0.4, 0.45, 0.5))
        self.draw_text(f"Kelas A - Ganjil 2026", 470, 810, font="F3", size=8, color=(0.4, 0.45, 0.5))
        
        # Footer line
        self.draw_line(40, 40, 555, 40, stroke_color=(0.8, 0.85, 0.9), width=0.5)
        self.draw_text("Program Studi S1 Informatika - FST UUI", 40, 28, font="F3", size=8, color=(0.4, 0.45, 0.5))
        self.draw_text(f"Halaman {self.current_page_num}", 500, 28, font="F3", size=8, color=(0.4, 0.45, 0.5))

    def set_color(self, r, g, b, stroke=False):
        if stroke:
            self.current_stream.append(f"{r:.3f} {g:.3f} {b:.3f} RG".encode("ascii"))
        else:
            self.current_stream.append(f"{r:.3f} {g:.3f} {b:.3f} rg".encode("ascii"))

    def draw_rect(self, x, y, w, h, fill_color=None, stroke_color=None, width=1.0):
        self.current_stream.append(f"q".encode("ascii"))
        if fill_color:
            self.set_color(*fill_color, stroke=False)
        if stroke_color:
            self.set_color(*stroke_color, stroke=True)
            self.current_stream.append(f"{width:.2f} w".encode("ascii"))
        
        if fill_color and stroke_color:
            self.current_stream.append(f"{x:.2f} {y:.2f} {w:.2f} {h:.2f} re B".encode("ascii"))
        elif fill_color:
            self.current_stream.append(f"{x:.2f} {y:.2f} {w:.2f} {h:.2f} re f".encode("ascii"))
        elif stroke_color:
            self.current_stream.append(f"{x:.2f} {y:.2f} {w:.2f} {h:.2f} re S".encode("ascii"))
        self.current_stream.append(f"Q".encode("ascii"))

    def draw_line(self, x1, y1, x2, y2, stroke_color=(0.1, 0.2, 0.5), width=1.0):
        self.current_stream.append(f"q".encode("ascii"))
        self.set_color(*stroke_color, stroke=True)
        self.current_stream.append(f"{width:.2f} w".encode("ascii"))
        self.current_stream.append(f"{x1:.2f} {y1:.2f} m {x2:.2f} {y2:.2f} l S".encode("ascii"))
        self.current_stream.append(f"Q".encode("ascii"))

    def escape_text(self, text):
        return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    def draw_text(self, text, x, y, font="F1", size=10, color=(0.1, 0.1, 0.1)):
        safe_t = self.escape_text(text)
        self.current_stream.append(f"q".encode("ascii"))
        self.set_color(*color, stroke=False)
        self.current_stream.append(f"BT /{font} {size:.2f} Tf {x:.2f} {y:.2f} Td ({safe_t}) Tj ET".encode("latin1", errors="replace"))
        self.current_stream.append(f"Q".encode("ascii"))

    def close(self):
        if self.current_stream:
            self.pages.append(b"\n".join(self.current_stream))
            self.current_stream = []

    def get_pdf_bytes(self):
        self.close()
        objects = []
        objects.append(b"") # 0
        objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
        
        font_obj_ids = {}
        next_id = 3
        for f_alias, f_name in self.fonts.items():
            font_obj_ids[f_alias] = next_id
            objects.append(f"<< /Type /Font /Subtype /Type1 /BaseFont /{f_name} /Encoding /WinAnsiEncoding >>".encode("ascii"))
            next_id += 1
            
        font_res_str = " ".join([f"/{k} {font_obj_ids[k]} 0 R" for k in self.fonts])
        
        page_obj_ids = []
        for i in range(len(self.pages)):
            page_id = next_id
            page_obj_ids.append(page_id)
            next_id += 2
            
        kids_str = " ".join([f"{pid} 0 R" for pid in page_obj_ids])
        pages_tree = f"<< /Type /Pages /Kids [ {kids_str} ] /Count {len(page_obj_ids)} /MediaBox [ 0 0 {self.page_width:.2f} {self.page_height:.2f} ] >>".encode("ascii")
        objects[2] = pages_tree
        
        for i, page_bytes in enumerate(self.pages):
            page_id = page_obj_ids[i]
            content_id = page_id + 1
            
            p_obj = f"<< /Type /Page /Parent 2 0 R /Resources << /Font << {font_res_str} >> >> /Contents {content_id} 0 R >>".encode("ascii")
            objects.append(p_obj)
            
            c_obj = f"<< /Length {len(page_bytes)} >>\nstream\n".encode("ascii") + page_bytes + b"\nendstream"
            objects.append(c_obj)
            
        out = []
        out.append(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets = [0]
        curr_offset = len(out[0])
        
        for obj_idx in range(1, len(objects)):
            obj_data = objects[obj_idx]
            offsets.append(curr_offset)
            header = f"{obj_idx} 0 obj\n".encode("ascii")
            footer = b"\nendobj\n"
            out.append(header)
            out.append(obj_data)
            out.append(footer)
            curr_offset += len(header) + len(obj_data) + len(footer)
            
        xref_offset = curr_offset
        out.append(f"xref\n0 {len(objects)}\n".encode("ascii"))
        out.append(b"0000000000 65535 f \n")
        for obj_idx in range(1, len(objects)):
            out.append(f"{offsets[obj_idx]:010d} 00000 n \n".encode("ascii"))
            
        out.append(f"trailer\n<< /Size {len(objects)} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii"))
        return b"".join(out)

def build_pdf_document():
    print("-> Membangun Dokumen PDF Kontrak Kuliah Algoritma Pemrograman...")
    pdf = SimplePDF()
    
    # -------------------------------------------------------------
    # HALAMAN 1: KOP, IDENTITAS, MANFAAT, DESKRIPSI, CPL/CPMK
    # -------------------------------------------------------------
    pdf.new_page()
    
    # Header KOP Universitas
    pdf.draw_text("YAYASAN UBUDIYAH INDONESIA", 297.64 - 85, 775, font="F2", size=11, color=(0.12, 0.23, 0.54))
    pdf.draw_text("UNIVERSITAS UBUDIYAH INDONESIA", 297.64 - 105, 760, font="F2", size=12, color=(0.12, 0.23, 0.54))
    pdf.draw_text("FAKULTAS SAINS DAN TEKNOLOGI - PROGRAM STUDI S1 INFORMATIKA", 297.64 - 170, 746, font="F2", size=9.5, color=(0.01, 0.52, 0.78))
    pdf.draw_text("Jl. Alue Naga, Desa Tibang, Kec. Syiah Kuala, Kota Banda Aceh | Laman: https://uui.ac.id", 297.64 - 185, 734, font="F3", size=8, color=(0.4, 0.45, 0.5))
    pdf.draw_line(40, 725, 555, 725, stroke_color=(0.12, 0.23, 0.54), width=1.8)
    pdf.draw_line(40, 722, 555, 722, stroke_color=(0.01, 0.52, 0.78), width=0.8)
    
    # Judul Dokumen
    pdf.draw_text("KONTRAK PERKULIAHAN", 297.64 - 95, 698, font="F2", size=14, color=(0.12, 0.23, 0.54))
    pdf.draw_text("SEMESTER GANJIL TAHUN AKADEMIK 2026/2027", 297.64 - 130, 683, font="F2", size=10, color=(0.01, 0.52, 0.78))
    
    # Section I: Identitas MK
    pdf.draw_rect(40, 660, 515, 18, fill_color=(0.12, 0.23, 0.54))
    pdf.draw_text("I. IDENTITAS MATA KULIAH", 48, 665, font="F2", size=9.5, color=(1, 1, 1))
    
    id_rows = [
        ("Nama Mata Kuliah", "Algoritma Pemrograman", "Kode MK / SKS", "IFR206 / 4 SKS (T=2, P=2)"),
        ("Program Studi / Jenjang", "S1 - Informatika / Sarjana (S1)", "Semester / Periode", "Semester I / 2026 Ganjil"),
        ("Fakultas / Universitas", "Fakultas Sains dan Teknologi / UUI", "Nama Kelas / Sistem", "Kelas A / Reguler"),
        ("Kurikulum / Rumpun", "Kurikulum 2023 (OBE) / Wajib Prodi", "Kapasitas / Peserta", "60 Mahasiswa / 1 Mahasiswa"),
        ("Dosen Pengampu Utama", "Mahendar Dwi Payana, S.ST., M.T.", "Ruang Perkuliahan", "Lab. Rekayasa Perangkat Lunak UUI")
    ]
    
    y_id = 642
    for r1, v1, r2, v2 in id_rows:
        bg = (0.97, 0.98, 0.99) if (y_id % 30 < 15) else (1, 1, 1)
        pdf.draw_rect(40, y_id - 2, 515, 14, fill_color=bg, stroke_color=(0.9, 0.92, 0.95), width=0.5)
        pdf.draw_text(r1 + ":", 45, y_id + 1, font="F2", size=8, color=(0.2, 0.25, 0.35))
        pdf.draw_text(v1, 150, y_id + 1, font="F1", size=8, color=(0.05, 0.1, 0.15))
        pdf.draw_text(r2 + ":", 310, y_id + 1, font="F2", size=8, color=(0.2, 0.25, 0.35))
        pdf.draw_text(v2, 400, y_id + 1, font="F1", size=8, color=(0.05, 0.1, 0.15))
        y_id -= 14
        
    # Section II: Manfaat & Deskripsi
    y_sec = y_id - 8
    pdf.draw_rect(40, y_sec, 515, 18, fill_color=(0.12, 0.23, 0.54))
    pdf.draw_text("II. MANFAAT DAN DESKRIPSI SINGKAT MATA KULIAH", 48, y_sec + 5, font="F2", size=9.5, color=(1, 1, 1))
    
    y_sec -= 16
    pdf.draw_text("Mata kuliah Algoritma dan Pemrograman (IFR206) merupakan fondasi utama kurikulum S1 Informatika UUI untuk menanamkan", 40, y_sec, font="F1", size=8.5, color=(0.1, 0.15, 0.2))
    y_sec -= 11
    pdf.draw_text("kemampuan berpikir komputasional (computational thinking), logika pemecahan masalah sistematis, serta implementasi ke dalam", 40, y_sec, font="F1", size=8.5, color=(0.1, 0.15, 0.2))
    y_sec -= 11
    pdf.draw_text("bahasa pemrograman terstruktur. Mahasiswa mempelajari flowchart, pseudocode, tipe data, alokasi memori, percabangan,", 40, y_sec, font="F1", size=8.5, color=(0.1, 0.15, 0.2))
    y_sec -= 11
    pdf.draw_text("perulangan, array 1D/2D, fungsi modular, rekursi, searching, dan sorting menuju kesiapan rekayasa software tingkat lanjut.", 40, y_sec, font="F1", size=8.5, color=(0.1, 0.15, 0.2))
    
    # Section III: CPL & CPMK
    y_sec -= 18
    pdf.draw_rect(40, y_sec, 515, 18, fill_color=(0.12, 0.23, 0.54))
    pdf.draw_text("III. CAPAIAN PEMBELAJARAN LULUSAN & MATA KULIAH (CPL & CPMK)", 48, y_sec + 5, font="F2", size=9.5, color=(1, 1, 1))
    
    y_sec -= 15
    cpls = [
        ("CPL01 (Pengetahuan Dasar)", "Menguasai teori, prinsip, dan konsep fundamental logika algoritma dan informatika."),
        ("CPL03 (Keterampilan Kerja)", "Mampu berpikir analitis, logis, dan sistematis dalam merancang solusi komputasi terstruktur."),
        ("CPL04 (Keahlian Khusus)", "Mampu menyajikan solusi terprogram atas permasalahan nyata berbasis kode bersih."),
        ("CPL08 (Sikap & Nilai)", "Menjunjung tinggi etika akademik, kejujuran kode program, dan kepatuhan anti-plagiarisme.")
    ]
    for code, desc in cpls:
        pdf.draw_text("• " + code + ":", 45, y_sec, font="F2", size=8, color=(0.12, 0.23, 0.54))
        pdf.draw_text(desc, 175, y_sec, font="F1", size=8, color=(0.1, 0.15, 0.2))
        y_sec -= 11
        
    y_sec -= 4
    cpmks = [
        ("CPMK-1: Notasi & Memori", "Mampu menganalisis logika komputasi, merancang flowchart ISO, pseudocode, & alokasi tipe data."),
        ("CPMK-2: Struktur Kontrol", "Mampu merancang alur logika percabangan (if-else, switch-case) dan perulangan (for, while, do-while)."),
        ("CPMK-3: Array & String", "Mampu merekayasa struktur data larik (Array 1D & Matriks 2D) dan manipulasi teks (String)."),
        ("CPMK-4: Modular & Rekursi", "Mampu menerapkan modularisasi kode (fungsi/prosedur) dan penyelesaian masalah rekursif."),
        ("CPMK-5: Searching & Sorting", "Mampu menganalisis dan membandingkan efisiensi algoritma pencarian dan pengurutan data."),
        ("CPMK-6: Mini Project Mandiri", "Mampu membangun dan mendokumentasikan mini project aplikasi berbasis algoritma terstruktur.")
    ]
    for code, desc in cpmks:
        pdf.draw_text("• " + code + ":", 45, y_sec, font="F2", size=8, color=(0.01, 0.52, 0.78))
        pdf.draw_text(desc, 175, y_sec, font="F1", size=8, color=(0.1, 0.15, 0.2))
        y_sec -= 11

    # -------------------------------------------------------------
    # HALAMAN 2: MATRIKS 16 PERTEMUAN (RPS PERKULIAHAN)
    # -------------------------------------------------------------
    pdf.new_page()
    
    pdf.draw_rect(40, 770, 515, 20, fill_color=(0.12, 0.23, 0.54))
    pdf.draw_text("IV. MATRIKS RENCANA PERKULIAHAN (16 PERTEMUAN)", 48, 776, font="F2", size=10, color=(1, 1, 1))
    
    # Table Header
    y_tab = 745
    pdf.draw_rect(40, y_tab, 515, 16, fill_color=(0.01, 0.52, 0.78))
    pdf.draw_text("Prt", 45, y_tab + 4, font="F2", size=8, color=(1, 1, 1))
    pdf.draw_text("Kemampuan Akhir (Sub-CPMK)", 70, y_tab + 4, font="F2", size=8, color=(1, 1, 1))
    pdf.draw_text("Materi Pembelajaran Pokok", 220, y_tab + 4, font="F2", size=8, color=(1, 1, 1))
    pdf.draw_text("Metode & Bentuk", 410, y_tab + 4, font="F2", size=8, color=(1, 1, 1))
    pdf.draw_text("Evaluasi", 510, y_tab + 4, font="F2", size=8, color=(1, 1, 1))
    
    rps_pdf_rows = [
        ("1", "Memahami konsep algoritma & notasi", "Kontrak Kuliah, Logika Komputasi, Flowchart, Pseudocode", "Ceramah & Diskusi", "Partisipatif"),
        ("2", "Mampu mengelola tipe data & memori", "Variabel, Tipe Data Primitif, Konstanta, Type Casting", "Praktikum Lab", "Tugas 1"),
        ("3", "Mampu mengevaluasi ekspresi logika", "Operator Aritmatika, Relasional, Logika, Bitwise", "Problem Solving", "Tugas 2"),
        ("4", "Mampu merancang percabangan bersyarat", "Struktur Kontrol if, if-else, nested-if, ternary", "Live Coding", "Tugas 3"),
        ("5", "Mampu merekayasa menu switch-case", "Struktur switch-case, fall-through, navigasi menu CLI", "Praktikum Lab", "Quiz 1"),
        ("6", "Mampu mengendalikan perulangan", "Counted loop (for), uncounted (while, do-while)", "Praktikum Lab", "Tugas 4"),
        ("7", "Mampu memanipulasi array 1 dimensi", "Larik 1D, Indeks Elemen, I/O, Nilai Min/Max/Average", "Praktikum Intensif", "Tugas 5"),
        ("8", "EVALUASI CAPAIAN PERTEMUAN 1 s.d. 7", "UJIAN TENGAH SEMESTER (UTS) - TEORI & LIVE CODING", "Ujian Lab Terjadwal", "UTS (30%)"),
        ("9", "Mampu mengolah array multidimensi", "Matriks 2D (Baris-Kolom), Operasi Matriks, String", "Praktik Lab", "Tugas 6"),
        ("10", "Mampu menyusun fungsi & prosedur", "Modularitas, Parameter pass-by-val/ref, Scope Variabel", "Live Coding", "Tugas 7"),
        ("11", "Mampu menerapkan teknik rekursi", "Rekursif, Base Case, Recursive Case, Call Stack", "Simulasi Trace Table", "Quiz 2"),
        ("12", "Mampu mengimplementasikan searching", "Linear Search, Binary Search, Perbandingan O(n)/O(log n)", "Eksperimen Algoritma", "Tugas 8"),
        ("13", "Mampu merekayasa algoritma sorting", "Bubble Sort, Selection Sort, Insertion Sort, Trace", "Praktik Sorting", "Tugas 9"),
        ("14", "Mampu merancang struktur aplikasi mini", "Arsitektur Modular Aplikasi CLI, Validasi Input", "Praktik Proyek", "Partisipatif"),
        ("15", "Mampu mengembangkan proyek mini", "Pengembangan Capstone Mini Project: Sistem Terstruktur", "PjBL & Asistensi", "Asistensi"),
        ("16", "EVALUASI CAPAIAN KOMPREHENSIF PROYEK", "UJIAN AKHIR SEMESTER (UAS) - DEMO & UJI PROGRAM", "Demo & Uji Program", "UAS (40%)")
    ]
    
    y_r = y_tab - 18
    for p_num, subc, mat, met, bot in rps_pdf_rows:
        is_hl = p_num in ["8", "16"]
        bg = (0.92, 0.96, 1.0) if is_hl else ((0.97, 0.98, 0.99) if int(p_num) % 2 == 1 else (1, 1, 1))
        pdf.draw_rect(40, y_r - 2, 515, 17, fill_color=bg, stroke_color=(0.88, 0.9, 0.93), width=0.5)
        
        f_type = "F2" if is_hl else "F1"
        c_text = (0.12, 0.23, 0.54) if is_hl else (0.1, 0.15, 0.2)
        
        pdf.draw_text(p_num, 48, y_r + 3, font="F2", size=8, color=c_text)
        pdf.draw_text(subc[:32], 70, y_r + 3, font=f_type, size=7.5, color=c_text)
        pdf.draw_text(mat[:40], 220, y_r + 3, font=f_type, size=7.5, color=c_text)
        pdf.draw_text(met[:20], 410, y_r + 3, font="F1", size=7.5, color=c_text)
        pdf.draw_text(bot, 508, y_r + 3, font="F2", size=7.5, color=c_text)
        y_r -= 18

    # -------------------------------------------------------------
    # HALAMAN 3: PENILAIAN SIAKAD, TATA TERTIB, PUSTAKA & PENGESAHAN
    # -------------------------------------------------------------
    pdf.new_page()
    
    # Section V: Kriteria Penilaian Sesuai Siakad UUI
    pdf.draw_rect(40, 770, 515, 18, fill_color=(0.12, 0.23, 0.54))
    pdf.draw_text("V. KOMPONEN EVALUASI & SKALA PENILAIAN (PORTAL SIAKAD UUI)", 48, 775, font="F2", size=9.5, color=(1, 1, 1))
    
    # Table 1: Komponen Evaluasi Siakad
    y_ev = 750
    pdf.draw_rect(40, y_ev, 515, 14, fill_color=(0.04, 0.31, 0.54))
    pdf.draw_text("No.", 45, y_ev + 3, font="F2", size=7.5, color=(1, 1, 1))
    pdf.draw_text("Metode Evaluasi", 70, y_ev + 3, font="F2", size=7.5, color=(1, 1, 1))
    pdf.draw_text("Jenis Evaluasi", 180, y_ev + 3, font="F2", size=7.5, color=(1, 1, 1))
    pdf.draw_text("Bobot Evaluasi", 430, y_ev + 3, font="F2", size=7.5, color=(1, 1, 1))
    pdf.draw_text("Syarat Lulus", 505, y_ev + 3, font="F2", size=7.5, color=(1, 1, 1))
    
    eval_pdf_siakad = [
        ("1", "TUGAS", "Kognitif/Pengetahuan - Tugas", "20%", "-"),
        ("2", "QUIZ", "Kognitif/Pengetahuan - Quiz", "5%", "-"),
        ("3", "UTS", "Kognitif/Pengetahuan - Ujian Tengah Semester", "30%", "-"),
        ("4", "UAS", "Kognitif/Pengetahuan - Ujian Akhir Semester", "40%", "-"),
        ("5", "AKTIVITAS PARTISIPATIF", "Aktivitas Partisipatif", "5%", "-"),
        ("", "Total Persentase Komponen Evaluasi", "", "100%", "")
    ]
    
    y_ev -= 12
    for n, met, jen, bob, syr in eval_pdf_siakad:
        is_tot = "Total" in met
        bg = (0.9, 0.94, 0.98) if is_tot else ((0.97, 0.98, 0.99) if y_ev % 24 < 12 else (1, 1, 1))
        pdf.draw_rect(40, y_ev - 1, 515, 12, fill_color=bg, stroke_color=(0.88, 0.9, 0.93), width=0.5)
        pdf.draw_text(n, 48, y_ev + 2, font="F2", size=7.5, color=(0.1, 0.15, 0.2))
        pdf.draw_text(met, 70, y_ev + 2, font="F2" if is_tot else "F1", size=7.5, color=(0.1, 0.15, 0.2))
        pdf.draw_text(jen, 180, y_ev + 2, font="F1", size=7.5, color=(0.2, 0.25, 0.3))
        pdf.draw_text(bob, 440, y_ev + 2, font="F2", size=7.5, color=(0.12, 0.23, 0.54))
        pdf.draw_text(syr, 520, y_ev + 2, font="F1", size=7.5, color=(0.2, 0.25, 0.3))
        y_ev -= 12
        
    pdf.draw_text("Keterangan: Syarat Lulus Mata Kuliah adalah pemenuhan komponen nilai wajib. Mahasiswa tanpa komponen utama dinyatakan tidak lulus.", 40, y_ev, font="F3", size=7, color=(0.35, 0.4, 0.45))
    
    # Table 2: Mapping Grade Siakad
    y_gr = y_ev - 16
    pdf.draw_text("Tabel Mapping Konversi Grade Nilai Akhir Siakad UUI:", 40, y_gr + 4, font="F2", size=8, color=(0.12, 0.23, 0.54))
    y_gr -= 12
    pdf.draw_rect(40, y_gr, 515, 13, fill_color=(0.01, 0.52, 0.78))
    pdf.draw_text("Grade", 70, y_gr + 3, font="F2", size=7.5, color=(1, 1, 1))
    pdf.draw_text("Bobot", 180, y_gr + 3, font="F2", size=7.5, color=(1, 1, 1))
    pdf.draw_text("Nilai Bawah", 290, y_gr + 3, font="F2", size=7.5, color=(1, 1, 1))
    pdf.draw_text("Nilai Atas", 400, y_gr + 3, font="F2", size=7.5, color=(1, 1, 1))
    pdf.draw_text("Status Kelulusan", 475, y_gr + 3, font="F2", size=7.5, color=(1, 1, 1))
    
    grades_siakad = [
        ("A", "4,00", "80,00", "100,00", "Sangat Memuaskan / Lulus"),
        ("B", "3,00", "60,00", "79,99", "Baik / Lulus"),
        ("C", "2,00", "40,00", "59,99", "Cukup (Batas Minimal Lulus)"),
        ("D", "1,00", "20,00", "39,99", "Kurang / Wajib Mengulang"),
        ("E", "0,00", "0,00", "19,99", "Gagal / Tidak Lulus")
    ]
    y_gr -= 11
    for grd, bbt, nb, na, stt in grades_siakad:
        bg = (0.97, 0.98, 0.99) if y_gr % 22 < 11 else (1, 1, 1)
        pdf.draw_rect(40, y_gr - 1, 515, 11, fill_color=bg, stroke_color=(0.88, 0.9, 0.93), width=0.5)
        pdf.draw_text(grd, 80, y_gr + 2, font="F2", size=7.5, color=(0.12, 0.23, 0.54))
        pdf.draw_text(bbt, 190, y_gr + 2, font="F1", size=7.5, color=(0.1, 0.15, 0.2))
        pdf.draw_text(nb, 305, y_gr + 2, font="F1", size=7.5, color=(0.1, 0.15, 0.2))
        pdf.draw_text(na, 415, y_gr + 2, font="F1", size=7.5, color=(0.1, 0.15, 0.2))
        pdf.draw_text(stt, 475, y_gr + 2, font="F1", size=7.0, color=(0.2, 0.25, 0.3))
        y_gr -= 11

    # Section VI: Tata Tertib Singkat
    y_sec = y_gr - 6
    pdf.draw_rect(40, y_sec, 515, 15, fill_color=(0.12, 0.23, 0.54))
    pdf.draw_text("VI. TATA TERTIB & NORMA AKADEMIK PERKULIAHAN", 48, y_sec + 4, font="F2", size=8.5, color=(1, 1, 1))
    
    y_sec -= 12
    rules = [
        ("1. Toleransi Terlambat", "Maksimal 15 menit. Terlambat > 15 menit boleh masuk tetapi presensi alpa."),
        ("2. Batas Minimal Hadir", "Kehadiran minimal 75% (12 pertemuan riil) sebagai syarat mutlak mengikuti UAS."),
        ("3. Etika & Perlengkapan", "Berbusana rapi, sopan, berkerah, bersepatu. Membawa laptop dengan IDE pemrograman C/C++/Python."),
        ("4. Sanksi Plagiarisme", "Kecurangan akademik / copy-paste kode dikenakan pembatalan nilai (Nilai E otomatis).")
    ]
    for r_title, r_desc in rules:
        pdf.draw_text("• " + r_title + ":", 45, y_sec, font="F2", size=7.5, color=(0.12, 0.23, 0.54))
        pdf.draw_text(r_desc, 155, y_sec, font="F1", size=7.5, color=(0.1, 0.15, 0.2))
        y_sec -= 9.5

    # Section VII: Lembar Pengesahan
    y_sec -= 4
    pdf.draw_rect(40, y_sec, 515, 15, fill_color=(0.12, 0.23, 0.54))
    pdf.draw_text("VII. LEMBAR PENGESAHAN KONTRAK PERKULIAHAN", 48, y_sec + 4, font="F2", size=8.5, color=(1, 1, 1))
    
    y_sec -= 11
    pdf.draw_text("Banda Aceh, 1 September 2026", 410, y_sec, font="F1", size=8, color=(0.2, 0.25, 0.3))
    
    y_sec -= 13
    pdf.draw_text("Perwakilan Mahasiswa Kelas A,", 70, y_sec, font="F2", size=8.5, color=(0.12, 0.23, 0.54))
    pdf.draw_text("Dosen Pengampu Mata Kuliah,", 360, y_sec, font="F2", size=8.5, color=(0.12, 0.23, 0.54))
    
    y_sec -= 35
    pdf.draw_text("( _____________________________ )", 50, y_sec, font="F1", size=8.5, color=(0.1, 0.1, 0.1))
    pdf.draw_text("( Mahendar Dwi Payana, S.ST., M.T. )", 340, y_sec, font="F2", size=8.5, color=(0.1, 0.1, 0.1))
    
    y_sec -= 8
    pdf.draw_text("NPM. ............................................", 50, y_sec, font="F1", size=7.5, color=(0.4, 0.45, 0.5))
    pdf.draw_text("NIDN. 1331108701", 340, y_sec, font="F1", size=7.5, color=(0.4, 0.45, 0.5))
    
    y_sec -= 15
    pdf.draw_text("Mengetahui,", 275, y_sec, font="F1", size=8, color=(0.2, 0.25, 0.3))
    y_sec -= 9
    pdf.draw_text("Ketua Program Studi S1 Informatika UUI", 205, y_sec, font="F2", size=8.5, color=(0.12, 0.23, 0.54))
    y_sec -= 30
    pdf.draw_text("( M. Bayu Wibawa, S.Kom., MMSI )", 200, y_sec, font="F1", size=8.5, color=(0.1, 0.1, 0.1))
    y_sec -= 8
    pdf.draw_text("NIDN. ....................................................", 200, y_sec, font="F1", size=7.5, color=(0.4, 0.45, 0.5))

    # Save PDF
    pdf_bytes = pdf.get_pdf_bytes()
    with open(PDF_MAIN, "wb") as f:
        f.write(pdf_bytes)
    shutil.copyfile(PDF_MAIN, PDF_LATEST)
    
    print(f"✅ Dokumen Kontrak Kuliah PDF berhasil dibuat: {PDF_MAIN}")
    print(f"✅ Ukuran Berkas: {len(pdf_bytes) / 1024:.2f} KB (Batas Siakad: 2048 KB / 2 MB)")
    print(f"✅ Salinan rilis: {PDF_LATEST}")

if __name__ == "__main__":
    build_docx_kontrak()
    build_pdf_document()
