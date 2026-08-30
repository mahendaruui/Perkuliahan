#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generator Dokumen Kontrak Perkuliahan Resmi (Sesuai Siakad UUI)
Mata Kuliah: Pemrograman Berorientasi Objek (IFR 214 - 3 SKS)
Program Studi S1 Informatika - Fakultas Sains dan Teknologi
Universitas Ubudiyah Indonesia (UUI)
Semester Ganjil T.A. 2026/2027 - Kelas A
Dosen Pengampu: Mahendar Dwi Payana, S.ST., M.T.
"""

import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
OUTPUT_DIR = os.path.join(REPO_ROOT, "kontrak-kuliah")
os.makedirs(OUTPUT_DIR, exist_ok=True)

DOCX_OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Kontrak_Kuliah_OOP_PHP_IFR214_Mahendar_Dwi_Payana.docx")
DOCX_LATEST_FILE = os.path.join(OUTPUT_DIR, "Kontrak_Kuliah_OOP_PHP_IFR214_Latest.docx")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('shd'):
            tcPr.remove(child)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = parse_xml(f'<{tag} {nsdecls("w")} w:val="{edge_data.get("val", "single")}" w:sz="{edge_data.get("sz", "4")}" w:space="0" w:color="{edge_data.get("color", "CCCCCC")}"/>')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def build_docx_kontrak():
    print("-> Membangun Dokumen Kontrak Perkuliahan Word (.docx) sesuai Siakad UUI...")
    doc = docx.Document()
    
    # Margin Halaman Standar A4
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    # 1. KOP SURAT RESMI UUI
    kop_table = doc.add_table(rows=1, cols=2)
    kop_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kop_table.autofit = False
    
    # Col 1: Logo UUI
    logo_cell = kop_table.rows[0].cells[0]
    logo_cell.width = Cm(2.2)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    logo_path = os.path.join(REPO_ROOT, "docs", "public", "uuilogo.png")
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Cm(1.8))
        
    # Col 2: Teks Kop Kampus
    text_cell = kop_table.rows[0].cells[1]
    text_cell.width = Cm(13.8)
    text_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    p_kop = text_cell.paragraphs[0]
    p_kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kop.paragraph_format.space_after = Pt(2)
    p_kop.paragraph_format.line_spacing = 1.1
    
    r1 = p_kop.add_run("YAYASAN UBUDIYAH INDONESIA\nUNIVERSITAS UBUDIYAH INDONESIA\n")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(13)
    r1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    r2 = p_kop.add_run("FAKULTAS SAINS DAN TEKNOLOGI - PROGRAM STUDI S1 INFORMATIKA\n")
    r2.bold = True
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
    
    r3 = p_kop.add_run("Jl. Alue Naga, Desa Tibang, Kec. Syiah Kuala, Kota Banda Aceh, Aceh 23114\nLaman: https://uui.ac.id | Pos-el: informatika@uui.ac.id")
    r3.font.name = "Calibri"
    r3.font.size = Pt(8.5)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    # Garis Pembatas
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(4)
    p_line.paragraph_format.space_after = Pt(14)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_line = p_line.add_run("═" * 60)
    r_line.font.name = "Calibri"
    r_line.font.size = Pt(11)
    r_line.font.bold = True
    r_line.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    # JUDUL DOKUMEN
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    r_t1 = p_title.add_run("KONTRAK PERKULIAHAN")
    r_t1.bold = True
    r_t1.font.name = "Calibri"
    r_t1.font.size = Pt(15)
    r_t1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_t2 = p_sub.add_run("SEMESTER GANJIL TAHUN AKADEMIK 2026/2027")
    r_t2.bold = True
    r_t2.font.name = "Calibri"
    r_t2.font.size = Pt(11.5)
    r_t2.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    def add_sec_heading(num_str, title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r_n = p.add_run(num_str + " ")
        r_n.bold = True
        r_n.font.name = "Calibri"
        r_n.font.size = Pt(11.5)
        r_n.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        r_t = p.add_run(title_text)
        r_t.bold = True
        r_t.font.name = "Calibri"
        r_t.font.size = Pt(11.5)
        r_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return p

    def add_body_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.name = "Calibri"
            r_pre.font.size = Pt(10.5)
            r_pre.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        r_b = p.add_run(text)
        r_b.font.name = "Calibri"
        r_b.font.size = Pt(10.5)
        r_b.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    def add_bullet_p(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r_pre = p.add_run(bold_prefix + ": ")
            r_pre.bold = True
            r_pre.font.name = "Calibri"
            r_pre.font.size = Pt(10.5)
            r_pre.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        r_b = p.add_run(text)
        r_b.font.name = "Calibri"
        r_b.font.size = Pt(10.5)
        r_b.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    # 1. IDENTITAS MATA KULIAH
    add_sec_heading("I.", "IDENTITAS MATA KULIAH")
    id_data = [
        ["1. Nama Mata Kuliah", "Pemrograman Berorientasi Objek"],
        ["2. Kode Mata Kuliah / Bobot SKS", "IFR 214 / 3 SKS (Teori: 2 SKS, Praktik: 1 SKS)"],
        ["3. Program Studi / Jenjang", "S1 - Informatika / Sarjana (Strata-1)"],
        ["4. Fakultas / Universitas", "Fakultas Sains dan Teknologi / Universitas Ubudiyah Indonesia"],
        ["5. Kurikulum / Rumpun MK", "Kurikulum 2023 (OBE) / Mata Kuliah Wajib Keahlian"],
        ["6. Semester / Periode Akademik", "Semester III (Tiga) / 2026 Ganjil"],
        ["7. Nama Kelas / Sistem Kuliah", "Kelas A / Reguler"],
        ["8. Kapasitas Kelas / Jumlah Peserta", "60 Mahasiswa / 37 Mahasiswa"],
        ["9. Dosen Pengampu Utama", "Mahendar Dwi Payana, S.ST., M.T. (NIDN: 1331108701)"],
        ["10. Ruang Perkuliahan & Praktikum", "Lab. Rekayasa Perangkat Lunak & Sistem Informasi UUI"]
    ]
    
    table_id = doc.add_table(rows=len(id_data), cols=2)
    table_id.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_id.autofit = False
    
    for idx, (label, val) in enumerate(id_data):
        row = table_id.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Cm(6.5)
        c2.width = Cm(9.5)
        
        bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(c1, bg_color)
        set_cell_background(c2, bg_color)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c2, top=60, bottom=60, left=100, right=100)
        set_cell_border(c1, bottom={"sz": "2", "color": "E2E8F0"})
        set_cell_border(c2, bottom={"sz": "2", "color": "E2E8F0"})
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(val)
        r2.font.name = "Calibri"
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # 2. MANFAAT MATA KULIAH
    add_sec_heading("II.", "MANFAAT MATA KULIAH")
    add_body_p("Mata kuliah Pemrograman Berorientasi Objek (OOP) merupakan pilar fundamental dalam kurikulum rekayasa perangkat lunak modern. Penguasaan OOP membekali mahasiswa Informatika dengan kemampuan berpikir terstruktur, modular, dan terukur dalam merancang arsitektur aplikasi skala besar. Mahasiswa dipersiapkan untuk menguasai standar industri software engineering terkini berbasis PHP 8+, prinsip SOLID, serta arsitektur Clean Architecture yang menjadi prasyarat mutlak dalam pengembangan perangkat lunak enterprise dan framework industri global (seperti Laravel dan Symfony).")

    # 3. DESKRIPSI SINGKAT MATA KULIAH
    add_sec_heading("III.", "DESKRIPSI MATA KULIAH")
    add_body_p("Mata kuliah ini membahas konsep, filosofi, prinsip, dan implementasi paradigma Object-Oriented Programming (OOP) menggunakan bahasa pemrograman PHP versi 8+ secara komprehensif. Materi pembelajaran mencakup pemahaman mendalam tentang anatomi class, object, typed properties, constructor property promotion, empat pilar utama OOP (Enkapsulasi, Pewarisan, Polimorfisme, dan Abstraksi), trait, interface, backed enums, penanganan galat tangguh (Exception Handling), manipulasi koleksi data bertipe (First-Class Collections), manajemen berkas (File I/O Stream), pengorganisasian namespace berbasis standar PSR-4 dengan Composer, penerapan 5 prinsip desain SOLID, hingga perancangan aplikasi arsitektur berlapis Model-Service-Repository pada studi kasus Point of Sale (POS) terpadu.")

    # 4. CAPAIAN PEMBELAJARAN (CPL & CPMK)
    add_sec_heading("IV.", "CAPAIAN PEMBELAJARAN LULUSAN & MATA KULIAH (CPL & CPMK)")
    add_bullet_p("CPL-P1 (Pengetahuan)", "Menguasai konsep teoretis, paradigma, prinsip, dan teknik OOP dalam rekayasa perangkat lunak modern.")
    add_bullet_p("CPL-KU1 (Keterampilan Umum)", "Mampu berpikir logis, kritis, sistematis, dan solutif dalam merancang abstraksi dan arsitektur perangkat lunak.")
    add_bullet_p("CPL-KK1 (Keterampilan Khusus)", "Mampu merancang, mengimplementasikan, menguji, dan mendokumentasikan aplikasi berorientasi objek menggunakan PHP 8+ sesuai standar industri.")
    add_bullet_p("CPL-S1 (Sikap & Tata Nilai)", "Menunjukkan integritas akademik, etika profesi rekayasa perangkat lunak, komitmen mutu, kerja sama tim, dan kepatuhan anti-plagiarisme.")

    add_body_p("Daftar Capaian Pembelajaran Mata Kuliah (CPMK):", bold_prefix="Capaian Spesifik:")
    add_bullet_p("CPMK-1", "Mampu menganalisis hakikat paradigma OOP dan membandingkannya secara kritis dengan paradigma prosedural pada PHP 8+.")
    add_bullet_p("CPMK-2", "Mampu mengimplementasikan class, object, constructor property promotion, method, dan manajemen referensi memori (Object Handle).")
    add_bullet_p("CPMK-3", "Mampu menerapkan 4 pilar fundamental OOP: Enkapsulasi (Information Hiding), Pewarisan (Inheritance), Polimorfisme (Dynamic Dispatch), dan Abstraksi (Interface & Abstract Class).")
    add_bullet_p("CPMK-4", "Mampu mengelola arsitektur proyek modular melalui Namespace, Standar Autoloading PSR-4 Composer, Exception Handling tangguh, dan First-Class Collections.")
    add_bullet_p("CPMK-5", "Mampu menganalisis dan menerapkan 5 Prinsip Desain SOLID untuk mencegah pembusukan perangkat lunak (Software Rot).")
    add_bullet_p("CPMK-6", "Mampu merancang, membangun, dan mempresentasikan mini project aplikasi berbasis arsitektur Model-Service-Repository secara terpadu.")

    # 5. MATRIKS RENCANA PEMBELAJARAN MINGGUAN (16 PERTEMUAN)
    add_sec_heading("V.", "MATRIKS RENCANA PERKULIAHAN (16 PERTEMUAN)")
    
    rps_headers = ["Prt", "Kemampuan Akhir (Sub-CPMK)", "Bahan Kajian / Pokok Bahasan", "Bentuk Pembelajaran", "Metode Evaluasi"]
    rps_data = [
        ["1", "Memahami konsep dasar & paradigma OOP", "Kontrak Kuliah, Evolusi Paradigma, Ekosistem PHP 8+", "Kuliah Teori & Diskusi Interaktif", "Aktivitas Partisipatif"],
        ["2", "Mampu memodelkan class & object di memori", "Class, Object, Typed Properties, Object Handle vs Clone", "Kuliah & Praktikum Mandiri", "Tugas Praktikum"],
        ["3", "Mampu mengelola siklus hidup & constructor", "Constructor Promotion, Destructor, Named Args, Static", "Praktikum Terbimbing di Lab", "Tugas Praktikum"],
        ["4", "Mampu menerapkan prinsip enkapsulasi", "Visibility Modifiers, Tell Don't Ask, Readonly, Hooks", "Praktikum & Studi Kasus Domain", "Quiz / Tugas"],
        ["5", "Mampu merekayasa pewarisan & trait", "Inheritance (extends), Parent Chaining, Trait, Final", "Praktikum & Analisis Hierarki", "Tugas Praktikum"],
        ["6", "Mampu mengimplementasikan polimorfisme", "Polymorphism, Dynamic Method Dispatch, Type Hinting", "Praktikum & Kuis Konsep", "Quiz"],
        ["7", "Mampu merancang abstraksi kontrak sistem", "Abstract Class, Template Method, Interface, Backed Enum", "Praktikum & Desain Kontrak", "Tugas Praktikum"],
        ["8", "EVALUASI CAPAIAN PERTEMUAN 1 s.d. 7", "UJIAN TENGAH SEMESTER (UTS) - TEORI & LIVE CODING", "Ujian Tertulis & Praktik Lab", "UTS (30%)"],
        ["9", "Mampu mengorganisasi namespace & autoload", "Namespace Architecture, Composer Setup, Standar PSR-4", "Praktikum Konfigurasi Proyek", "Tugas Praktikum"],
        ["10", "Mampu merekayasa penanganan galat tangguh", "Exception Hierarchy, Multi-Catch, Custom Exception", "Praktikum Robust Error Handling", "Tugas Praktikum"],
        ["11", "Mampu memanipulasi first-class collections", "First-Class Collections, SPL Interfaces, Map/Filter", "Praktikum Manipulasi Koleksi", "Quiz / Tugas"],
        ["12", "Mampu mengelola persistensi berkas I/O", "File I/O Stream, SplFileObject, Race Condition & LOCK_EX", "Praktikum File-Based Persistence", "Tugas Praktikum"],
        ["13", "Mampu menerapkan 5 prinsip desain SOLID", "Prinsip SOLID (SRP, OCP, LSP, ISP, DIP) pada PHP Modern", "Studi Kasus & Diskusi Desain", "Aktivitas Partisipatif"],
        ["14", "Mampu merancang arsitektur aplikasi bersih", "Arsitektur Model-Service-Repository & Domain Modeling", "Praktikum Arsitektur Perangkat Lunak", "Tugas Besar"],
        ["15", "Mampu merekayasa mini project terpadu", "Pengembangan Capstone Project: Sistem POS Enterprise", "PjBL & Asistensi Terbimbing", "Asistensi Proyek"],
        ["16", "EVALUASI CAPAIAN KOMPREHENSIF PROYEK", "UJIAN AKHIR SEMESTER (UAS) - DEMO & PRESENTASI PROYEK", "Evaluasi Capstone & Uji Program", "UAS (40%)"]
    ]
    
    table_rps = doc.add_table(rows=len(rps_data) + 1, cols=5)
    table_rps.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_rps.autofit = False
    
    col_widths = [Cm(1.1), Cm(4.2), Cm(5.6), Cm(3.2), Cm(1.9)]
    
    # Header Row
    hdr_cells = table_rps.rows[0].cells
    for i, title in enumerate(rps_headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=60, right=60)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(9.0)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    for r_idx, row_values in enumerate(rps_data):
        row = table_rps.rows[r_idx + 1]
        is_highlight = row_values[0] in ["8", "16"]
        bg = "EFF6FF" if is_highlight else ("F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
        
        for c_idx, val in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            set_cell_border(cell, bottom={"sz": "2", "color": "CBD5E1"})
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            
            if c_idx in [0, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(8.5)
            if is_highlight:
                r.bold = True
                r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            else:
                r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # 6. KOMPONEN EVALUASI & MAPPING NILAI SIAKAD
    add_sec_heading("VI.", "KOMPONEN EVALUASI & SKALA PENILAIAN (SIAKAD UUI)")
    add_body_p("Komponen dan bobot evaluasi pembelajaran disesuaikan 100% dengan konfigurasi resmi Portal Siakad Universitas Ubudiyah Indonesia sebagai berikut:")
    
    # Table Evaluasi Siakad
    eval_siakad = [
        ["No.", "Metode Evaluasi", "Jenis Evaluasi", "Bobot Evaluasi", "Syarat Lulus"],
        ["1", "TUGAS", "Kognitif/Pengetahuan - Tugas", "20%", "-"],
        ["2", "QUIZ", "Kognitif/Pengetahuan - Quiz", "5%", "-"],
        ["3", "UTS", "Kognitif/Pengetahuan - Ujian Tengah Semester", "30%", "-"],
        ["4", "UAS", "Kognitif/Pengetahuan - Ujian Akhir Semester", "40%", "-"],
        ["5", "AKTIVITAS PARTISIPATIF", "Aktivitas Partisipatif (Kehadiran & Diskusi)", "5%", "-"],
        ["", "Total Persentase Komponen Evaluasi", "", "100%", ""]
    ]
    
    table_eval = doc.add_table(rows=len(eval_siakad), cols=5)
    table_eval.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_eval.autofit = False
    col_w_eval = [Cm(1.0), Cm(3.2), Cm(7.0), Cm(2.8), Cm(2.0)]
    
    for r_idx, row_values in enumerate(eval_siakad):
        row = table_eval.rows[r_idx]
        is_hdr = (r_idx == 0)
        is_tot = (r_idx == len(eval_siakad) - 1)
        
        bg = "0B4F8A" if is_hdr else ("E2E8F0" if is_tot else ("F8FAFC" if r_idx % 2 == 1 else "FFFFFF"))
        
        for c_idx, val in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.width = col_w_eval[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=70, right=70)
            set_cell_border(cell, bottom={"sz": "2", "color": "94A3B8" if (is_hdr or is_tot) else "E2E8F0"})
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if c_idx in [0, 3, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9.0)
            if is_hdr:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif is_tot:
                r.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            else:
                r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Keterangan Syarat Lulus
    p_ket = doc.add_paragraph()
    p_ket.paragraph_format.space_before = Pt(4)
    p_ket.paragraph_format.space_after = Pt(8)
    r_k1 = p_ket.add_run("Keterangan: ")
    r_k1.bold = True
    r_k1.font.name = "Calibri"
    r_k1.font.size = Pt(9.0)
    r_k2 = p_ket.add_run("Syarat Lulus Mata Kuliah adalah pemenuhan seluruh komponen nilai wajib. Mahasiswa yang tidak memiliki salah satu komponen nilai utama akan dinyatakan tidak lulus.")
    r_k2.font.name = "Calibri"
    r_k2.font.size = Pt(9.0)
    r_k2.font.italic = True

    add_body_p("Tabel Mapping Konversi Grade Nilai Akhir Siakad UUI:", bold_prefix="Tabel Konversi Grade: ")
    
    # Table Mapping Grade Siakad
    grade_mapping = [
        ["Grade", "Bobot", "Nilai Bawah", "Nilai Atas", "Status Kelulusan"],
        ["A", "4,00", "80,00", "100,00", "Sangat Memuaskan / Lulus"],
        ["B", "3,00", "60,00", "79,99", "Baik / Lulus"],
        ["C", "2,00", "40,00", "59,99", "Cukup / Batas Minimal Lulus"],
        ["D", "1,00", "20,00", "39,99", "Kurang / Wajib Mengulang"],
        ["E", "0,00", "0,00", "19,99", "Gagal / Tidak Lulus"]
    ]
    
    table_grade = doc.add_table(rows=len(grade_mapping), cols=5)
    table_grade.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_grade.autofit = False
    col_w_grade = [Cm(2.0), Cm(2.2), Cm(3.2), Cm(3.2), Cm(5.4)]
    
    for r_idx, row_values in enumerate(grade_mapping):
        row = table_grade.rows[r_idx]
        is_hdr = (r_idx == 0)
        bg = "0284C7" if is_hdr else ("F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
        
        for c_idx, val in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.width = col_w_grade[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            set_cell_border(cell, bottom={"sz": "2", "color": "E2E8F0"})
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if c_idx in [0, 1, 2, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9.0)
            if is_hdr:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # 7. TATA TERTIB & NORMA AKADEMIK
    add_sec_heading("VII.", "TATA TERTIB & NORMA AKADEMIK PERKULIAHAN")
    add_bullet_p("1. Kehadiran & Ketepatan Waktu", "Mahasiswa wajib hadir tepat waktu sesuai jadwal perkuliahan. Toleransi keterlambatan maksimal adalah 15 menit. Keterlambatan melebihi batas tersebut tetap diperkenankan mengikuti perkuliahan namun presensi dianggap alpa.")
    add_bullet_p("2. Batas Minimal Kehadiran", "Kehadiran minimal mahasiswa adalah 75% dari total 16 pertemuan riil sebagai syarat mutlak mengikuti Ujian Akhir Semester (UAS). Mahasiswa dengan kehadiran di bawah 75% otomatis dinyatakan tidak berhak mengikuti UAS.")
    add_bullet_p("3. Etika Berbusana & Sikap", "Berpakaian sopan, rapi, berkerah (bukan kaos oblong polos), dan mengenakan sepatu di lingkungan kampus dan laboratorium. Menjaga etika berbicara, sopan santun akademik, serta mematikan/mengatur nada dering senyap (silent mode) pada perangkat komunikasi selama perkuliahan.")
    add_bullet_p("4. Fasilitas & Praktikum Lab", "Setiap mahasiswa wajib membawa laptop kerja pribadi yang telah terinstal PHP 8.1+, Composer, Visual Studio Code / PHPStorm, dan Git. Dilarang merusak, mengubah konfigurasi jaringan lab, atau mengotori ruang laboratorium.")
    add_bullet_p("5. Penyerahan Tugas & Sanksi Keterlambatan", "Tugas mandiri dan kelompok wajib diunggah tepat waktu melalui LMS Siakad UUI / Repositori GitHub resmi kelas. Keterlambatan pengumpulan tugas dikenakan penalti pemotongan nilai sebesar 10% per hari (maksimal toleransi 3 hari).")
    add_bullet_p("6. Integritas Akademik & Anti-Plagiarisme", "Kecurangan akademik dalam bentuk apa pun (plagiarisme kode program, menyalin pekerjaan teman tanpa izin, menggunakan joki tugas, atau kecurangan saat ujian) dikenakan sanksi tegas pembatalan nilai (Nilai E otomatis untuk mata kuliah ini) dan dilaporkan ke Komite Disiplin Universitas Ubudiyah Indonesia.")

    # 8. PUSTAKA & REFERENSI
    add_sec_heading("VIII.", "SUMBER BELAJAR & DAFTAR PUSTAKA")
    add_bullet_p("Buku Ajar Utama", "Payana, M. D., Yusian TB, D. R., Wibawa, M. B., Hamdi, N., & Musliyana, Z. (2026). Buku Ajar Pemrograman Berorientasi Objek menggunakan PHP 8+. Universitas Ubudiyah Indonesia Press.")
    add_bullet_p("Referensi Industri", "Zandstra, M. (2021). PHP 8 Objects, Patterns, and Practice: Mastering OO Enhancements, Design Patterns, and Test-Driven Development (6th ed.). Apress.")
    add_bullet_p("Prinsip Arsitektur", "Martin, R. C. (2017). Clean Architecture: A Craftsman's Guide to Software Structure and Design. Prentice Hall.")
    add_bullet_p("Standar Koding", "PHP-FIG. (2024). PHP Standard Recommendations: PSR-4 (Autoloading) & PSR-12 (Extended Coding Style Guide). https://www.php-fig.org/")

    # 9. LEMBAR PERSETUJUAN & PENGESAHAN
    add_sec_heading("IX.", "LEMBAR PENGESAHAN KONTRAK PERKULIAHAN")
    add_body_p("Kontrak perkuliahan ini disepakati bersama secara sadar dan penuh tanggung jawab antara Dosen Pengampu Mata Kuliah dan Perwakilan Mahasiswa Kelas A pada pertemuan pertama perkuliahan Semester Ganjil Tahun Akademik 2026/2027 di Universitas Ubudiyah Indonesia.")

    # Tanda Tangan Table
    doc.add_paragraph().paragraph_format.space_before = Pt(8)
    
    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_loc.paragraph_format.space_after = Pt(8)
    r_loc = p_loc.add_run("Banda Aceh, 1 September 2026")
    r_loc.font.name = "Calibri"
    r_loc.font.size = Pt(10.0)
    
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    # Row 0: Labels
    c_mhs_top, c_dos_top = sig_table.rows[0].cells[0], sig_table.rows[0].cells[1]
    c_mhs_top.width = Cm(8.0)
    c_dos_top.width = Cm(8.0)
    
    p_mt = c_mhs_top.paragraphs[0]
    p_mt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_mt.add_run("Perwakilan Mahasiswa / Komti Kelas A,\n\n\n\n\n( _____________________________ )\nNPM. ....................................")
    r.font.name = "Calibri"
    r.font.size = Pt(10.0)
    
    p_dt = c_dos_top.paragraphs[0]
    p_dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_dt.add_run("Dosen Pengampu Mata Kuliah,\n\n\n\n\n( Mahendar Dwi Payana, S.ST., M.T. )\nNIDN. 1331108701")
    r.font.name = "Calibri"
    r.font.size = Pt(10.0)
    r.bold = True
    
    # Row 1: Mengetahui Kaprodi
    c_kpd = sig_table.rows[1].cells[0]
    c_kpd_right = sig_table.rows[1].cells[1]
    c_kpd.width = Cm(16.0)
    c_kpd.merge(c_kpd_right)
    
    p_kp = c_kpd.paragraphs[0]
    p_kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kp.paragraph_format.space_before = Pt(14)
    r = p_kp.add_run("Mengetahui,\nKetua Program Studi S1 Informatika UUI\n\n\n\n\n( ____________________________________ )\nNIDN. ....................................")
    r.font.name = "Calibri"
    r.font.size = Pt(10.0)

    doc.save(DOCX_OUTPUT_FILE)
    import shutil
    shutil.copyfile(DOCX_OUTPUT_FILE, DOCX_LATEST_FILE)
    print(f"✅ Dokumen Kontrak Kuliah Word berhasil disimpan: {DOCX_OUTPUT_FILE}")
    print(f"✅ Salinan rilis terbaru: {DOCX_LATEST_FILE}")

if __name__ == "__main__":
    build_docx_kontrak()
