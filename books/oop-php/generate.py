#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generator Buku Ajar: Pemrograman Berorientasi Objek Menggunakan PHP 8+
Mata Kuliah: Pemrograman Berorientasi Objek (IFR 214 - 3 SKS)
Penulis: Mahendar Dwi Payana, S.ST., M.T.
Program Studi Informatika, Fakultas Sains dan Teknologi, Universitas Ubudiyah Indonesia

Mendukung versioning dokumen Word (.docx)
Penggunaan:
    python generate.py
    python generate.py --version 1.0.0
"""

import os
import sys
import argparse
import shutil

from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Include repository root for module imports
REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
sys.path.insert(0, REPO_ROOT)

from books.core.book_builder import AcademicBookBuilder

def parse_args():
    parser = argparse.ArgumentParser(description="Generator Buku Ajar OOP PHP 8+")
    parser.add_argument(
        "--version", "-v",
        type=str,
        default="v1.0.0",
        help="Versi rilis dokumen (default: v1.0.0)"
    )
    parser.add_argument(
        "--output-dir", "-o",
        type=str,
        default=os.path.join(os.path.dirname(os.path.abspath(__file__)), "output"),
        help="Folder penyimpanan berkas output"
    )
    return parser.parse_args()

def build_oop_php_book(version="v1.0.0", output_dir=None):
    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
        
    os.makedirs(output_dir, exist_ok=True)
    clean_version = version if version.startswith("v") else f"v{version}"
    filename = f"Buku_Ajar_OOP_PHP8_{clean_version}_Mahendar_Dwi_Payana.docx"
    output_path = os.path.join(output_dir, filename)
    latest_path = os.path.join(output_dir, "Buku_Ajar_OOP_PHP8_Latest.docx")

    print("=" * 65)
    print(f" GENERATOR BUKU AJAR: OOP PHP 8+ (UNIVERSITAS UBUDIYAH INDONESIA)")
    print(f" Versi Dokumen : {clean_version}")
    print(f" Target Berkas : {output_path}")
    print("=" * 65)

    builder = AcademicBookBuilder(
        output_path=output_path,
        book_title="Buku Ajar Pemrograman Berorientasi Objek (PHP 8+)",
        version=clean_version,
        course_code="IFR 214"
    )

    # 1. Halaman Sampul
    print("-> 1/6 Membangun Halaman Judul...")
    builder.add_title_page(
        main_title="PEMROGRAMAN BERORIENTASI OBJEK\nMENGGUNAKAN PHP 8+",
        subtitle="Pendekatan Teoretis, Praktik Rekayasa Perangkat Lunak Modern,\ndan Implementasi Arsitektur Bersih",
        author="Mahendar Dwi Payana, S.ST., M.T.",
        year="2025"
    )

    # 2. Halaman Hak Cipta & KDT
    print("-> 2/6 Membangun Halaman Penerbitan & Hak Cipta (KDT)...")
    meta_info = [
        ["Judul Buku", "Buku Ajar Pemrograman Berorientasi Objek menggunakan PHP 8+"],
        ["Penulis", "Mahendar Dwi Payana, S.ST., M.T."],
        ["Editor Ahli", "Tim Pengembang Kurikulum Program Studi Informatika UUI"],
        ["Desain & Tata Letak", "Laboratorium Rekayasa Perangkat Lunak & Sistem Informasi UUI"],
        ["Penerbit", "UUI Press / Program Studi Informatika\nFakultas Sains dan Teknologi, Universitas Ubudiyah Indonesia"],
        ["Alamat Redaksi", "Jl. Alue Naga, Desa Tibang, Kec. Syiah Kuala, Kota Banda Aceh, Aceh 23114\nLaman: https://uui.ac.id | Pos-el: info@uui.ac.id"],
        ["Edisi & Cetakan", f"Cetakan Pertama, 2025 (Versi Rilis: {clean_version})"],
        ["Nomor ISBN", "978-623-XXXX-XX-X (e-Book Digital Reference)"]
    ]
    kdt_text = (
        "PAYANA, Mahendar Dwi\n"
        "    Buku Ajar Pemrograman Berorientasi Objek menggunakan PHP 8+ /\n"
        "    Mahendar Dwi Payana ; editor, Tim Reviewer Informatika UUI.\n"
        "    -- Cet. 1 -- Banda Aceh : UUI Press, 2025.\n"
        "    xvi, 175 hlm. : ilus. ; 29,7 cm. (Format A4)\n\n"
        "    Bibliografi : hlm. 162-164\n"
        "    Glosarium   : hlm. 158-161\n"
        "    ISBN 978-623-XXXX-XX-X\n\n"
        "    1. Pemrograman Berorientasi Objek (Ilmu Komputer)   2. PHP 8+\n"
        "    I. Judul   II. Tim Reviewer Informatika UUI\n"
        "                                                                005.133 -- DDC23"
    )
    builder.add_copyright_page(meta_info, kdt_text)

    # 3. Kata Pengantar
    print("-> 3/6 Membangun Kata Pengantar Penulis...")
    preface_paras = [
        "Puji dan syukur penulis panjatkan ke hadirat Allah SWT, karena atas limpahan rahmat, taufik, dan hidayah-Nya, Buku Ajar berjudul \"Pemrograman Berorientasi Objek menggunakan PHP 8+\" ini dapat diselesaikan dengan baik. Buku ini disusun secara khusus sebagai buku referensi utama dan pegangan perkuliahan bagi mahasiswa Program Studi Informatika, Fakultas Sains dan Teknologi, Universitas Ubudiyah Indonesia, maupun para pembelajar mandiri yang ingin menguasai rekayasa perangkat lunak modern berbasis PHP.",
        "Perkembangan teknologi perangkat lunak saat ini menuntut setiap pengembang untuk menulis kode program yang tidak sekadar bekerja, melainkan juga harus modular, terstruktur, mudah diuji (testable), serta siap dikembangkan dalam skala besar (scalable). Paradigma Pemrograman Berorientasi Objek (Object-Oriented Programming/OOP) telah menjadi landasan fundamental dalam industri perangkat lunak dunia. Bersamaan dengan itu, bahasa pemrograman PHP telah mengalami evolusi revolusioner, khususnya sejak rilis PHP versi 8.0, 8.1, 8.2, dan seterusnya. Fitur modern seperti Constructor Property Promotion, Type Safety yang ketat, Readonly Properties, Backed Enums, serta ekosistem manajemen pustaka Composer dengan standar PSR-4 menjadikan PHP sebagai bahasa kelas industri yang sangat tangguh.",
        "Buku ajar ini disusun secara berjenjang (pedagogical ladder) yang memandu pembaca mulai dari filosofi dasar pergeseran paradigma prosedural menuju berorientasi objek, penguasaan empat pilar utama OOP (Enkapsulasi, Pewarisan, Polimorfisme, dan Abstraksi), penanganan kesalahan tangguh (Exception Handling), manipulasi koleksi data dan file I/O, hingga penerapan prinsip desain SOLID dan arsitektur Model-Service-Repository dalam bentuk proyek terintegrasi.",
        "Penulis menyadari bahwa buku ini masih memiliki ruang untuk perbaikan dan penyempurnaan di masa mendatang. Oleh karena itu, saran, masukan konstruktif, dan kritik yang membangun dari para sejawat akademisi, praktisi industri, serta mahasiswa senantiasa penulis nantikan.",
        "Akhir kata, penulis menyampaikan apresiasi yang mendalam kepada pimpinan Universitas Ubudiyah Indonesia, rekan-rekan dosen di lingkungan Fakultas Sains dan Teknologi, serta para mahasiswa yang senantiasa memberikan inspirasi selama proses belajar mengajar. Semoga buku ajar ini dapat memberikan manfaat nyata, menginspirasi, dan menjadi bekal kompetensi yang kokoh bagi kemajuan keilmuan teknologi informasi di tanah air."
    ]
    builder.add_preface(preface_paras, author="Mahendar Dwi Payana, S.ST., M.T.", location="Banda Aceh, 2025")

    # 4. Capaian Pembelajaran
    print("-> 4/6 Membangun Matriks Capaian Pembelajaran (CPL & CPMK)...")
    p_title = builder.doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(16)
    r = p_title.add_run("CAPAIAN PEMBELAJARAN MATA KULIAH")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    builder.add_paragraph("Mata kuliah Pemrograman Berorientasi Objek (Kode MK: IFR 214, Bobot: 3 SKS) dirancang untuk membekali mahasiswa dengan kompetensi berikut:")
    
    cpl_data = [
        ["CPL-P1", "Menguasai konsep, paradigma, prinsip, dan teknik pemrograman berorientasi objek dalam rekayasa perangkat lunak."],
        ["CPL-KU1", "Mampu berpikir logis, sistematis, kritis, dan solutif dalam merancang abstraksi perangkat lunak."],
        ["CPL-KK1", "Mampu merancang, mengimplementasikan, menguji, dan mendokumentasikan aplikasi berorientasi objek dengan PHP 8+."],
        ["CPL-S1", "Menunjukkan sikap profesional, berintegritas, mandiri, dan beretika akademik dalam penyelesaian tugas perangkat lunak."]
    ]
    builder.add_table(["Kode CPL", "Deskripsi Capaian Pembelajaran Lulusan (CPL)"], cpl_data)
    
    builder.add_heading_2("Daftar Capaian Pembelajaran Mata Kuliah (CPMK)")
    cpmk_data = [
        ["CPMK-1", "Mampu membandingkan paradigma prosedural vs berorientasi objek dan memahami ekosistem PHP 8+."],
        ["CPMK-2", "Mampu mengimplementasikan class, objek, constructor promotion, method, dan struktur memori."],
        ["CPMK-3", "Mampu menerapkan 4 pilar OOP: Encapsulation, Inheritance, Polymorphism, dan Abstraction."],
        ["CPMK-4", "Mampu mengelola arsitektur proyek melalui Namespace, Autoloading PSR-4, Exception Handling, dan Collections."],
        ["CPMK-5", "Mampu menganalisis dan menerapkan prinsip desain perangkat lunak SOLID pada aplikasi."],
        ["CPMK-6", "Mampu merancang dan membangun proyek aplikasi OOP PHP terintegrasi berbasis studi kasus riil."]
    ]
    builder.add_table(["Kode CPMK", "Deskripsi Capaian Pembelajaran Mata Kuliah"], cpmk_data)
    builder.doc.add_page_break()

    # 5. Daftar Isi
    print("-> 5/6 Membangun Daftar Isi Lengkap...")
    toc_entries = [
        ('SECTION', 'BAGIAN AWAL', ''),
        ('PRE', 'Halaman Judul & Informasi Penulis', 'i'),
        ('PRE', 'Informasi Penerbitan & Hak Cipta (KDT)', 'ii'),
        ('PRE', 'Kata Pengantar Penulis', 'iii'),
        ('PRE', 'Matriks Capaian Pembelajaran (CPL & CPMK)', 'iv'),
        ('PRE', 'Daftar Isi Lengkap', 'v'),
        
        ('SECTION', 'BATANG TUBUH PEMBELAJARAN', ''),
        ('BAB', 'BAB 1: PENGANTAR PARADIGMA OOP & EKOSISTEM PHP 8+', '1'),
        ('SUB', '1.1 Hakikat dan Evolusi Paradigma Pemrograman', '1'),
        ('SUB', '1.2 Perbandingan Komprehensif: Prosedural vs Berorientasi Objek', '3'),
        ('SUB', '1.3 Empat Pilar Fundamental Pemrograman Berorientasi Objek', '6'),
        ('SUB', '1.4 Transformasi Modern: Mengapa PHP 8+ Menjadi Standar Baru?', '8'),
        
        ('BAB', 'BAB 2: ANATOMI CLASS, OBJEK, DAN MANAJEMEN MEMORI', '12'),
        ('SUB', '2.1 Konsep Entitas Dunia Nyata dan Pemodelan Class', '12'),
        ('SUB', '2.2 Deklarasi Class dan Typed Properties di PHP 8+', '14'),
        ('SUB', '2.3 Instansiasi Objek dan Operator Arrow (->)', '16'),
        ('SUB', '2.4 Variabel Pseudo $this dan Lingkup Eksekusi', '18'),
        ('SUB', '2.5 Pengelolaan Memori: Object Handle vs Kloning (clone)', '20'),
        
        ('BAB', 'BAB 3: METHOD, CONSTRUCTOR PROMOTION, DAN SIKLUS HIDUP OBJEK', '24'),
        ('SUB', '3.1 Siklus Hidup Objek dan Inisialisasi Otomatis', '24'),
        ('SUB', '3.2 Pergeseran Paradigma: Constructor Property Promotion', '26'),
        ('SUB', '3.3 Named Arguments pada PHP 8.0+', '28'),
        ('SUB', '3.4 Member Statis dan Pola Static Factory Method', '31'),
        
        ('BAB', 'BAB 4: ENKAPSULASI, VISIBILITY MODIFIERS, DAN READONLY', '36'),
        ('SUB', '4.1 Filosofi Enkapsulasi: Information Hiding & State Invariant', '36'),
        ('SUB', '4.2 Tiga Tingkat Hak Akses & Prinsip Tell, Don\'t Ask', '38'),
        ('SUB', '4.3 Praktik Terbaik: Mutasi Terkendali vs Public Properties', '40'),
        ('SUB', '4.4 Modern Immutability: readonly Properties & Class', '42'),
        ('SUB', '4.5 PHP 8.4+: Asymmetric Visibility & Property Hooks', '44'),
        
        ('BAB', 'BAB 5: PEWARISAN (INHERITANCE) DAN KOMPOSISI TRAIT', '48'),
        ('SUB', '5.1 Filosofi Pewarisan: Taksonomi Is-A & Bahaya Fragile Base Class', '48'),
        ('SUB', '5.2 Sintaks Pewarisan, Keyword extends, dan Penimpaan parent::', '50'),
        ('SUB', '5.3 Mengunci Perilaku: final Class, Method, dan Constants (PHP 8.1+)', '53'),
        ('SUB', '5.4 Trait: Horizontal Code Reuse & Resolusi Konflik (insteadof / as)', '55'),
        
        ('BAB', 'BAB 6: POLIMORFISME (POLYMORPHISM) DAN DYNAMIC DISPATCH', '60'),
        ('SUB', '6.1 Hakikat Polimorfisme: Klasifikasi Cardelli & Wegner (1985)', '60'),
        ('SUB', '6.2 Dynamic Method Dispatch pada Runtime Engine', '62'),
        ('SUB', '6.3 Polymorphic Type Hinting & Polymorphic Collections', '64'),
        ('SUB', '6.4 Operator instanceof & Type Narrowing', '67'),
        ('SUB', '6.5 Studi Kasus Enterprise: Multi-Channel Notification Engine', '69'),
        
        ('BAB', 'BAB 7: ABSTRAKSI: INTERFACE, ABSTRACT CLASS, DAN BACKED ENUM', '72'),
        ('SUB', '7.1 Filosofi Abstraksi: Pemisahan Kontrak (What) dan Implementasi (How)', '72'),
        ('SUB', '7.2 Abstract Class & Template Method Pattern', '74'),
        ('SUB', '7.3 Interface: Kontrak Murni & Multiple Interface Implementation', '76'),
        ('SUB', '7.4 Matriks Komparasi Akademis: Abstract Class vs Interface', '79'),
        ('SUB', '7.5 Integrasi Backed Enum di PHP 8.1+ & Pattern Matching', '81'),
        
        ('BAB', 'BAB 8: MANAJEMEN NAMESPACE, STANDAR PSR-4, DAN COMPOSER', '84'),
        ('SUB', '8.1 Filosofi Namespace & Masalah Global Namespace Pollution', '84'),
        ('SUB', '8.2 Resolusi Simbol Namespace: FQN, Qualified, dan Unqualified', '86'),
        ('SUB', '8.3 Standar PSR-4 Autoloading & Struktur Direktori Composer', '88'),
        ('SUB', '8.4 Konfigurasi composer.json & Optimasi Autoloading Produksi', '91'),
        ('SUB', '8.5 Studi Kasus Arsitektur Clean Namespace PSR-4', '93'),
        
        ('BAB', 'BAB 9: PENANGANAN KESALAHAN (EXCEPTION HANDLING) & ERROR FLOW', '96'),
        ('SUB', '9.1 Filosofi Exception Handling & Hierarki Throwable', '96'),
        ('SUB', '9.2 Blok Kontrol try-catch-finally & Multi-Catch Syntax', '98'),
        ('SUB', '9.3 Fitur Modern PHP 8+: throw Expression & Non-Capturing Catches', '101'),
        ('SUB', '9.4 Merancang Custom Domain Exceptions & Exception Chaining ($previous)', '103'),
        ('SUB', '9.5 Studi Kasus Transaksi Perbankan Terpadu', '105'),
        
        ('BAB', 'BAB 10: KOLEKSI OBJEK DAN MANIPULASI ARRAY MODERN', '108'),
        ('SUB', '10.1 Filosofi First-Class Collections & Masalah Type Safety', '108'),
        ('SUB', '10.2 Mengintegrasikan Antarmuka Standar PHP (SPL: Countable, IteratorAggregate)', '110'),
        ('SUB', '10.3 Manipulasi Fungsional: map, filter, reduce, dan Arrow Functions', '113'),
        ('SUB', '10.4 Studi Kasus Analisis Rekapitulasi Yudisium', '116'),
        
        ('BAB', 'BAB 11: MANAJEMEN BERKAS DAN ALIRAN DATA (FILE I/O STREAM)', '120'),
        ('SUB', '11.1 Filosofi Persistensi Berkas & Pencegahan Race Condition (flock LOCK_EX)', '120'),
        ('SUB', '11.2 Pemrosesan Berkas Berorientasi Objek dengan SplFileObject', '122'),
        ('SUB', '11.3 Manipulasi JSON Modern & Validasi Native di PHP 8.3+ (json_validate)', '125'),
        ('SUB', '11.4 Pola Desain: File-Based Repository Pattern', '127'),
        
        ('BAB', 'BAB 12: PRINSIP DESAIN PERANGKAT LUNAK SOLID PADA PHP MODERN', '132'),
        ('SUB', '12.1 Fondasi Teoretis SOLID: Menghindari Pembusukan Perangkat Lunak', '132'),
        ('SUB', '12.2 Bedah Komprehensif 5 Prinsip SOLID di PHP 8+ (SRP, OCP, LSP, ISP, DIP)', '134'),
        ('SUB', '12.3 Pipeline Checkout E-Commerce Berstandar SOLID', '139'),
        
        ('BAB', 'BAB 13: ARSITEKTUR APLIKASI (MODEL-SERVICE-REPOSITORY)', '144'),
        ('SUB', '13.1 Filosofi Clean Architecture & Arsitektur Berlapis (Dependency Rule)', '144'),
        ('SUB', '13.2 Lapisan 1: Entity / Domain Model dengan Backed Enum', '146'),
        ('SUB', '13.3 Repository Interface (Kontrak Abstraksi di Lapisan Domain)', '148'),
        ('SUB', '13.4 Service Layer (Use Case / Business Orchestrator)', '150'),
        ('SUB', '13.5 Implementasi Repository Konkret (Lapisan Infrastructure)', '152'),
        
        ('BAB', 'BAB 14: STUDI KASUS MINI PROJECT: POINT OF SALE (POS) TERPADU', '156'),
        ('SUB', '14.1 Deskripsi Kasus & Kebutuhan Sistem POS Enterprise', '156'),
        ('SUB', '14.2 Desain Domain Model & First-Class Collection', '158'),
        ('SUB', '14.3 Strategi Diskon & Pembayaran Polimorfik (SOLID OCP & DIP)', '161'),
        ('SUB', '14.4 Application Service & Persistensi File-Based Terproteksi', '163'),
        ('SUB', '14.5 Panduan Evaluasi & Rubrik Penilaian Capstone Project', '165'),
        
        ('SECTION', 'BAGIAN AKHIR', ''),
        ('POST', 'Glosarium Istilah Rekayasa Perangkat Lunak & OOP PHP 8+', '168'),
        ('POST', 'Daftar Pustaka & Referensi Akademis Terindeks', '172'),
        ('POST', 'Profil Penulis & Tim Reviewer', '175')
    ]
    builder.add_table_of_contents(toc_entries)

    # 6. Konten 14 BAB
    print("-> 6/6 Membangun Konten 14 BAB Pembelajaran...")
    
    # BAB 1
    builder.add_bab_title(1, "Pengantar Paradigma Pemrograman Berorientasi Objek & Ekosistem PHP 8+")
    builder.add_learning_objectives("Sub-CPMK 1", [
        "Menjelaskan konsep fundamental paradigma pemrograman dan pergeserannya dari prosedural ke berorientasi objek.",
        "Membedakan secara analitis antara struktur program prosedural dan struktur modular berbasis objek.",
        "Mengidentifikasi fitur-fitur revolusioner pada PHP modern (PHP 8.0, 8.1, 8.2, 8.3).",
        "Menyiapkan lingkungan pengembangan perangkat lunak berbasis PHP 8+ dan Composer."
    ])
    builder.add_heading_2("1.1 Hakikat dan Evolusi Paradigma Pemrograman")
    builder.add_paragraph("Dalam ilmu komputasi dan rekayasa perangkat lunak, paradigma pemrograman (programming paradigm) merupakan kerangka konseptual, model berpikir, dan metodologi fundamental yang mendikte bagaimana seorang pengembang menstrukturkan data, mengeksekusi logika, dan menyelesaikan persoalan komputasi. Sebagaimana halnya paradigma dalam sains, cara pandang ini menentukan alat bantu, sintaksis, serta batas-batas arsitektur yang digunakan dalam membangun program komputer.")
    builder.add_paragraph("Sepanjang sejarah perkembangan rekayasa perangkat lunak, terdapat tiga paradigma utama yang paling dominan mempengaruhi cara perangkat lunak dibangun di industri:")
    builder.add_bullet("Pemrograman Prosedural (Procedural Programming)", "Pendekatan berbasis instruksi sekuensial dan pemanggilan fungsi/prosedur yang memanipulasi data global atau lokal secara terpisah.")
    builder.add_bullet("Pemrograman Berorientasi Objek (Object-Oriented Programming - OOP)", "Pendekatan yang memandang sistem sebagai interaksi antar entitas mandiri (objek) yang membungkus status data (properti) dan perilakunya (method) ke dalam satu kesatuan utuh.")
    builder.add_bullet("Pemrograman Fungsional (Functional Programming)", "Pendekatan deklaratif yang memandang komputasi sebagai evaluasi fungsi matematis murni tanpa efek samping (side effects) dan tanpa status yang berubah-ubah (immutable state).")
    
    builder.add_heading_2("1.2 Perbandingan Komprehensif: Prosedural vs Berorientasi Objek")
    builder.add_paragraph("Pada masa-masa awal bahasa PHP (era PHP 3 dan PHP 4), gaya penulisan kode didominasi oleh pendekatan prosedural. Pengembang menuliskan serangkaian berkas PHP yang berisi gabungan kode logika, instruksi SQL, manipulasi string, dan tag HTML dalam satu berkas tunggal yang linier (sering diistilahkan sebagai 'spaghetti code').")
    builder.add_paragraph("Pendekatan prosedural memiliki kelebihan berupa kesederhanaan untuk skrip kecil (quick and dirty scripts). Namun, ketika skala aplikasi berkembang menjadi ribuan baris kode dengan puluhan modul bisnis, pendekatan prosedural mengalami masalah serius: data global mudah terkontaminasi, fungsi saling bergantung secara ketat (tightly coupled), dan pengujian unit (unit testing) menjadi hampir mustahil dilakukan.")
    
    builder.add_heading_3("Pendekatan Prosedural di PHP:")
    builder.add_code("<?php\n// Kode Prosedural: Data dan Fungsi Terpisah\nfunction hitungLuasPersegiPanjang(float $panjang, float $lebar): float {\n    return $panjang * $lebar;\n}\n\n$panjang = 12.5;\n$lebar   = 4.0;\n$luas    = hitungLuasPersegiPanjang($panjang, $lebar);\necho \"Luas Bangun: \" . $luas . \" cm²\";\n")
    
    builder.add_heading_3("Pendekatan Berorientasi Objek di PHP 8+:")
    builder.add_code("<?php\ndeclare(strict_types=1);\n\n// Kode OOP: Data ($panjang, $lebar) dan Perilaku (hitungLuas) Menyatu\nclass PersegiPanjang {\n    public function __construct(\n        private float $panjang,\n        private float $lebar\n    ) {\n        if ($panjang <= 0 || $lebar <= 0) {\n            throw new InvalidArgumentException(\"Dimensi harus bernilai positif.\");\n        }\n    }\n\n    public function hitungLuas(): float {\n        return $this->panjang * $this->lebar;\n    }\n}\n\n$bangun = new PersegiPanjang(12.5, 4.0);\necho \"Luas Bangun: \" . $bangun->hitungLuas() . \" cm²\";\n")
    
    table_comp = [
        ["Dimensi Analisis", "Pemrograman Prosedural", "Pemrograman Berorientasi Objek (OOP)"],
        ["Satuan Dasar Organisasi", "Fungsi, Subroutine, dan Prosedur", "Class dan Instance Objek"],
        ["Relasi Data & Perilaku", "Terpisah; data mengalir bebas ke berbagai fungsi", "Tersatukan secara kohesif di dalam entitas Objek"],
        ["Keamanan Status Data", "Rendah; variabel global rentan diubah tanpa validasi", "Tinggi; dilindungi melalui Enkapsulasi & Access Modifiers"],
        ["Reusabilitas Kode", "Rendah hingga sedang (mengandalkan copy-paste fungsi)", "Sangat tinggi (menggunakan Inheritance, Trait, & Polimorfisme)"],
        ["Kemudahan Pemeliharaan", "Menurun drastis seiring bertambahnya kompleksitas", "Tinggi; modul terisolasi memudahkan debugging & refactoring"],
        ["Penerapan Industri", "Skrip otomatisasi singkat, pemrosesan data mikro", "Framework Enterprise (Laravel, Symfony, Yii2, Laminas)"]
    ]
    builder.add_table(table_comp[0], table_comp[1:])
    
    builder.add_heading_2("1.3 Empat Pilar Fundamental Pemrograman Berorientasi Objek")
    builder.add_paragraph("Arsitektur OOP bertumpu kokoh di atas empat pilar utama. Penguasaan mendalam atas keempat pilar ini merupakan syarat mutlak bagi setiap rekayasawan perangkat lunak:")
    builder.add_bullet("1. Enkapsulasi (Encapsulation)", "Mekanisme penggabungan data (state) dan metode pemroses data (behavior) ke dalam satu wadah tunggal (class), seraya menyembunyikan detail internal yang sensitif dari akses luar langsung (Information Hiding).")
    builder.add_bullet("2. Pewarisan (Inheritance)", "Kemampuan sebuah class baru (subclass/child) untuk mewarisi atribut, metode, dan sifat dari class yang sudah ada (superclass/parent), mendorong efisiensi penulisan kode sesuai prinsip DRY (Don't Repeat Yourself).")
    builder.add_bullet("3. Polimorfisme (Polymorphism)", "Prinsip 'satu antarmuka, banyak rupa', yang memungkinkan objek dari berbagai class turunan berbeda merespons pemanggilan pesan atau method yang sama dengan perilaku spesifik masing-masing melalui dynamic dispatch.")
    builder.add_bullet("4. Abstraksi (Abstraction)", "Proses penyederhanaan kompleksitas sistem dengan hanya menampilkan antarmuka penting dan relevan kepada pengguna luar, seraya menyembunyikan mekanisme teknis internal yang rumit.")

    builder.add_heading_2("1.4 Transformasi Modern: Mengapa PHP 8+ Menjadi Standar Baru?")
    builder.add_paragraph("PHP (PHP: Hypertext Preprocessor) telah bertransformasi secara radikal. PHP 8.x bukan lagi sekadar bahasa skrip pelengkap HTML, melainkan bahasa pemrograman berorientasi objek yang lengkap, tangguh, memiliki sistem tipe data statis (Strict Typing), serta dilengkapi JIT (Just-In-Time) compiler.")
    
    table_features = [
        ["Fitur Modern", "Rilis", "Signifikansi & Dampak Arsitektural"],
        ["Typed Properties", "PHP 7.4", "Menjamin integritas tipe data properti secara statis sejak kompilasi."],
        ["Constructor Property Promotion", "PHP 8.0", "Mengeliminasi boilerplate kode inisialisasi class hingga 70%."],
        ["Named Arguments", "PHP 8.0", "Memungkinkan pemanggilan parameter fungsi secara eksplisit tanpa terikat urutan."],
        ["Union Types (A|B)", "PHP 8.0", "Mendukung deklarasi multi-tipe parameter secara legal dan type-safe."],
        ["Readonly Properties", "PHP 8.1", "Menjamin kekekalan nilai (immutability) properti setelah instansiasi."],
        ["Backed Enums", "PHP 8.1", "Tipe data enumerasi native dengan nilai skalar untuk validasi status yang aman."],
        ["Readonly Classes", "PHP 8.2", "Menjadikan seluruh properti class secara otomatis bersifat readonly."]
    ]
    builder.add_table(table_features[0], table_features[1:])

    builder.add_tip(
        "Tips Praktik Industri: Disiplin Deklarasi Strict Types",
        "Selalu biasakan menyematkan deklarasi 'declare(strict_types=1);' pada baris pertama setiap berkas program PHP Anda. Deklarasi ini memaksa Zend Engine menolak konversi tipe data otomatis (type coercion) yang acapkali menjadi celah bug kritis dan inkonsistensi data di lingkungan produksi enterprise."
    )

    builder.add_summary_and_questions([
        "Paradigma OOP mengorganisasikan sistem komputasi ke dalam entitas-entitas objek yang menggabungkan status data dan perilaku.",
        "Empat pilar utama OOP meliputi Enkapsulasi, Pewarisan, Polimorfisme, dan Abstraksi.",
        "PHP 8+ membawa perubahan masif dengan menghadirkan sistem tipe yang ketat, constructor promotion, readonly properties, dan performa tinggi.",
        "Pengembangan aplikasi skala besar wajib memanfaatkan standar pengorganisasian modern untuk menjaga skalabilitas dan kemudahan uji."
    ], [
        "Jelaskan kelemahan mendasar pemrograman prosedural jika diterapkan pada sistem informasi rumah sakit berskala besar!",
        "Uraikan bagaimana prinsip Information Hiding dalam enkapsulasi melindungi integritas data finansial perbankan!",
        "Sebutkan dan jelaskan 3 fitur baru PHP 8.x yang secara signifikan meningkatkan produktivitas penulisan kode berorientasi objek!",
        "Instal PHP 8.1+ di komputer kerja Anda, periksa versinya melalui terminal (`php -v`), dan buatlah skrip pengujian sederhana bertipe data ketat!"
    ])

    # BAB 2
    builder.add_bab_title(2, "Anatomi Class, Objek, dan Manajemen Memori")
    builder.add_learning_objectives("Sub-CPMK 2", [
        "Mendefinisikan perbedaan ontologis antara Class sebagai blueprint dan Objek sebagai instansi fisik di memori.",
        "Mendeklarasikan Class dengan Typed Properties dan Method fungsional sesuai standar PSR-12.",
        "Melakukan instansiasi objek dan mengakses properti serta method menggunakan operator arrow (`->`).",
        "Menjelaskan peran variabel pseudo `$this` dalam konteks runtime objek.",
        "Menganalisis pengelolaan referensi memori (Object Handle) dan mekanisme kloning objek (`clone`)."
    ])
    builder.add_heading_2("2.1 Konsep Entitas di Dunia Nyata dan Pemodelan Class")
    builder.add_paragraph("Dalam rekayasa perangkat lunak berorientasi objek, dunia nyata dimodelkan sebagai kumpulan objek. Objek nyata memiliki dua karakteristik dasar: karakteristik deskriptif (atribut/ciri-ciri) dan karakteristik operasional (tindakan/perilaku). Sebagai contoh, seorang Mahasiswa memiliki atribut berupa Nomor Induk Mahasiswa (NIM), Nama Lengkap, Program Studi, dan Indeks Prestasi Kumulatif (IPK). Di samping itu, mahasiswa memiliki perilaku seperti belajar, mengambil mata kuliah, dan membayar biaya kuliah.")
    builder.add_paragraph("Class adalah cetak biru (blueprint), deskriptor, atau skema rancangan abstrak yang mendefinisikan atribut apa saja yang akan dimiliki serta perilaku apa saja yang dapat dilakukan oleh suatu kelompok entitas. Sedangkan Objek adalah perwujudan konkret (instance) dari cetak biru tersebut yang dialokasikan di dalam memori komputer dan memiliki nilai data uniknya sendiri.")
    
    builder.add_heading_2("2.2 Deklarasi Class dan Typed Properties di PHP 8+")
    builder.add_paragraph("Sejak era PHP 7.4 dan disempurnakan di PHP 8+, PHP mewajibkan penulisan tipe data pada properti class (Typed Properties). Praktik ini sangat krusial guna mencegah bug akibat ketidaksesuaian tipe data pada saat runtime.")
    
    builder.add_code("<?php\ndeclare(strict_types=1);\n\nnamespace App\\Model;\n\nclass Mahasiswa\n{\n    public string $nim;\n    public string $nama;\n    public string $programStudi = \"Informatika\";\n    public float $ipk = 0.0;\n    public bool $isAktif = true;\n\n    public function perbaruiIpk(float $ipkBaru): void\n    {\n        if ($ipkBaru < 0.0 || $ipkBaru > 4.0) {\n            throw new \\InvalidArgumentException(\"Rentang IPK harus antara 0.00 hingga 4.00!\");\n        }\n        $this->ipk = $ipkBaru;\n    }\n\n    public function cetakKarakteristik(): void\n    {\n        $status = $this->isAktif ? \"Aktif\" : \"Cuti/Non-Aktif\";\n        echo \"=== KARTU TANDA MAHASISWA ===\\n\";\n        echo \"NIM           : {$this->nim}\\n\";\n        echo \"Nama Lengkap  : {$this->nama}\\n\";\n        echo \"Program Studi : {$this->programStudi}\\n\";\n        echo \"IPK Terkini   : \" . number_format($this->ipk, 2) . \"\\n\";\n        echo \"Status        : {$status}\\n\";\n        echo \"=============================\\n\";\n    }\n}\n")
    
    builder.add_heading_2("2.3 Instansiasi Objek dan Operator Arrow (`->`)")
    builder.add_paragraph("Untuk menciptakan objek fisik dari suatu class, bahasa PHP menggunakan kata kunci `new`. Ketika `new Mahasiswa()` dieksekusi, Zend Engine (mesin eksekusi inti PHP) akan mengalokasikan satu blok memori baru, menginisialisasi tabel simbol properti, dan mengembalikan penunjuk (handle) objek tersebut kepada variabel penampung.")
    
    builder.add_code("<?php\nrequire_once 'Mahasiswa.php';\nuse App\\Model\\Mahasiswa;\n\n$mhs1 = new Mahasiswa();\n$mhs1->nim = \"240101001\";\n$mhs1->nama = \"Cut Meurah Intan\";\n$mhs1->perbaruiIpk(3.88);\n\n$mhs2 = new Mahasiswa();\n$mhs2->nim = \"240101002\";\n$mhs2->nama = \"Teuku Rayhan\";\n$mhs2->perbaruiIpk(3.72);\n\n$mhs1->cetakKarakteristik();\n$mhs2->cetakKarakteristik();\n")
    
    builder.add_heading_2("2.4 Variabel Pseudo `$this` dan Lingkup Eksekusi")
    builder.add_paragraph("Variabel pseudo `$this` secara otomatis menunjuk ke instansi objek yang saat itu sedang mengeksekusi method bersangkutan (current object instance). Melalui `$this->namaProperti`, kode di dalam method dapat membaca dan memodifikasi status internal objek tersebut secara akurat.")

    builder.add_heading_2("2.5 Pengelolaan Memori di PHP: Object Handle vs Kloning (`clone`)")
    builder.add_paragraph("Poin penting yang wajib dipahami oleh pengembang PHP adalah bahwa variabel penampung objek di PHP sesungguhnya menyimpan Object Identifier (Handle), bukan salinan data fisik objek. Akibatnya, ketika suatu variabel objek di-assign ke variabel lain (`$b = $a`), PHP hanya menyalin referensi handle-nya, bukan menggandakan objek di memori.")
    
    builder.add_code("<?php\n$a = new Mahasiswa();\n$a->nama = \"Budi\";\n\n$b = $a;\n$b->nama = \"Siti\";\necho $a->nama; // Output: \"Siti\"\n\n$c = clone $a;\n$c->nama = \"Andi\";\necho $a->nama; // Output: Tetap \"Siti\"\necho $c->nama; // Output: \"Andi\"\n")

    builder.add_tip(
        "Tips Efisiensi Memori: Pahami Karakteristik Object Handle",
        "Ingatlah bahwa penugasan '$b = $a' di PHP hanya menyalin Object Handle (penunjuk memori), bukan menggandakan fisik objek. Manfaatkan sifat ini untuk menghemat alokasi RAM ketika melewatkan entitas besar ke service layer. Gunakan kata kunci 'clone' hanya ketika Anda benar-benar memerlukan duplikasi fisik independen."
    )

    builder.add_summary_and_questions([
        "Class bertindak sebagai template atau cetak biru abstrak, sementara Object adalah wujud fisik yang memakan memori komputer.",
        "Typed properties menjamin ketepatan tipe data dan mencegah bug runtime yang tidak terduga.",
        "Variabel pseudo `$this` merujuk ke instance objek yang sedang aktif mengeksekusi method.",
        "Variabel objek di PHP bekerja berdasarkan Object Handle. Untuk menduplikasi objek fisik secara terpisah, gunakan kata kunci `clone`."
    ], [
        "Buatlah class `BukuPerpustakaan` dengan properti `$isbn` (string), `$judul` (string), `$penulis` (string), `$stok` (int), dan method `pinjamBuku(int $jumlah): bool`!",
        "Jelaskan apa yang terjadi di memori komputer ketika instruksi `$obj2 = $obj1` dijalankan dibandingkan dengan `$obj2 = clone $obj1`!",
        "Mengapa PHP melempar pesan error `Typed property must not be accessed before initialization` dan bagaimana cara mencegahnya?",
        "Tuliskan sebuah skrip untuk mendemonstrasikan perilaku dua objek dari class yang sama namun memiliki state data yang berbeda!"
    ])

    # BAB 3
    builder.add_bab_title(3, "Method, Constructor Property Promotion, dan Siklus Hidup Objek")
    builder.add_learning_objectives("Sub-CPMK 2", [
        "Memahami siklus hidup objek mulai dari alokasi awal, inisialisasi constructor, hingga pelepasan oleh destructor.",
        "Menerapkan magic method `__construct()` dan `__destruct()` secara tepat.",
        "Menguasai fitur modern PHP 8.0: Constructor Property Promotion untuk efisiensi kode.",
        "Memanfaatkan Named Arguments dan Union Types pada pemanggilan method.",
        "Membedakan fungsi member statis (`static`) dengan member instansiasi, serta menerapkan pola Static Factory Method."
    ])
    builder.add_heading_2("3.1 Siklus Hidup Objek dan Inisialisasi Otomatis")
    builder.add_paragraph("Setiap objek di dalam memori komputer melalui siklus hidup yang terdefinisi dengan jelas: kelahiran (instansiasi & inisialisasi nilai awal), masa aktif (pemanggilan method dan manipulasi status), serta kematian (pelepasan sumber daya dan destruksi dari memori oleh Garbage Collector).")
    builder.add_heading_2("3.2 Pergeseran Paradigma: Dari Cara Tradisional ke Constructor Promotion")
    builder.add_code("<?php\ndeclare(strict_types=1);\n\nclass RekeningBank {\n    public function __construct(\n        public readonly string $nomorRekening,\n        public string $pemilik,\n        public float $saldo = 0.0\n    ) {\n        if ($saldo < 0.0) {\n            throw new \\InvalidArgumentException(\"Saldo awal tidak boleh negatif!\");\n        }\n    }\n}\n\n$rek = new RekeningBank(\"123-456-789\", \"Mahendar Dwi Payana\", 1_500_000.0);\necho \"Pemilik: \" . $rek->pemilik . \" | Saldo: Rp \" . number_format($rek->saldo);\n")
    
    builder.add_heading_2("3.3 Named Arguments pada PHP 8.0+")
    builder.add_code("<?php\nclass Pengguna {\n    public function __construct(\n        public string $username,\n        public string $email,\n        public string $role = \"Mahasiswa\",\n        public bool $isVerifikasi = false,\n        public string $zonaWaktu = \"Asia/Jakarta\"\n    ) {}\n}\n\n$user = new Pengguna(\n    email: \"mahendar@uui.ac.id\",\n    username: \"mahendar\",\n    zonaWaktu: \"Asia/Banda_Aceh\"\n);\n")

    builder.add_heading_2("3.4 Member Statis dan Pola Static Factory Method")
    builder.add_code("<?php\ndeclare(strict_types=1);\n\nclass AkunPengguna {\n    private function __construct(\n        public readonly string $username,\n        public readonly string $email,\n        public readonly string $levelAkses,\n        public readonly int $kuotaPenyimpananMB\n    ) {}\n\n    public static function buatAkunMahasiswa(string $nim, string $email): self {\n        return new self($nim, $email, 'Mahasiswa', 1024);\n    }\n\n    public static function buatAkunDosen(string $nidn, string $email): self {\n        return new self($nidn, $email, 'Dosen', 10240);\n    }\n}\n")

    builder.add_tip(
        "Tips Rekayasa: Kombinasi Constructor Promotion & Readonly",
        "Gabungkan modifier 'public readonly' langsung di dalam parameter constructor PHP 8+. Pola ini menghasilkan Data Transfer Object (DTO) dan Value Object yang sangat ringkas, aman dari mutasi liar (immutable), serta secara otomatis bersifat self-documenting."
    )
    builder.add_summary_and_questions([
        "Constructor (`__construct`) dieksekusi secara otomatis saat instansiasi untuk menginisialisasi status objek.",
        "Constructor Property Promotion di PHP 8 menyederhanakan deklarasi properti dan parameter menjadi satu kesatuan ringkas.",
        "Named Arguments memungkinkan pengisian parameter secara eksplisit, acak urutan, dan melewati nilai default.",
        "Member statis dapat diakses langsung melalui Class (`self::` atau `ClassName::`), dan Static Factory Method adalah standar industri untuk instansiasi objek multi-varian."
    ], [
        "Jelaskan keuntungan arsitektural penggunaan Constructor Property Promotion dibandingkan cara deklarasi klasik!",
        "Bagaimana peran destructor (`__destruct()`) dalam pengelolaan koneksi jaringan atau berkas?",
        "Rancanglah sebuah class `MataKuliah` dengan properti `$kodeMK`, `$namaMK`, `$sks`, dan `$dosenPengampu` menggunakan PHP 8 Constructor Promotion dan Named Arguments!",
        "Implementasikan Static Factory Method pada class `TransaksiKasir` untuk membuat transaksi tipe 'Tunai', 'Debit', dan 'QRIS'!"
    ])

    # =========================================================================
    # BAB 4: ENKAPSULASI, VISIBILITY MODIFIERS, DAN READONLY PROPERTIES
    # =========================================================================
    builder.add_bab_title(4, "Enkapsulasi, Visibility Modifiers, Readonly, dan Asymmetric Visibility")
    builder.add_learning_objectives("Sub-CPMK 3", [
        "Memahami filosofi fundamental Enkapsulasi, Information Hiding, dan konsep State Invariant.",
        "Menguasai secara mendalam 3 tingkatan Visibility Modifiers: public, protected, dan private pada properti, method, dan konstanta.",
        "Mengimplementasikan Getter (Accessor) dan Setter (Mutator) dengan validasi aturan domain bisnis serta pola Method Chaining.",
        "Menerapkan fitur modern readonly Properties (PHP 8.1+) dan readonly class (PHP 8.2+) untuk menjamin kekekalan status data (Data Immutability).",
        "Menganalisis dan mengimplementasikan fitur mutakhir PHP 8.4: Asymmetric Visibility (public private(set)) dan Property Hooks (get/set).",
        "Memahami mekanisme kendali properti dinamis melalui magic methods (__get, __set, __isset, __unset) serta aturan penolakan Dynamic Properties di PHP 8.2+."
    ])

    builder.add_heading_2("4.1 Filosofi dan Fondasi Teoretis Enkapsulasi")
    builder.add_paragraph(
        "Dalam rekayasa perangkat lunak modern, Enkapsulasi (Encapsulation) merupakan pilar fundamental pertama yang menyatukan data (state/properti) dan perilaku pemroses data (behavior/method) ke dalam satu unit struktural mandiri yang kohesif (Class), seraya menyembunyikan detail representasi internal yang sensitif dari akses luar langsung (Information Hiding)."
    )
    builder.add_paragraph(
        "Prinsip Information Hiding pertama kali dicetuskan secara ilmiah oleh David Parnas pada tahun 1972 dalam karya klasiknya mengenai dekomposisi modular sistem perangkat lunak. Parnas menegaskan bahwa modul perangkat lunak yang bermutu tinggi harus menyembunyikan keputusan perancangan internalnya dari modul lain. Kode pemanggil (client code) cukup berinteraksi dengan antarmuka publik (public interface) yang stabil tanpa perlu mengetahui bagaimana struktur data fisik disimpan atau dimanipulasi di dalam memori."
    )
    builder.add_paragraph(
        "Tujuan utama enkapsulasi adalah memelihara State Invariant, yaitu kondisi atau aturan kebenaran bisnis yang harus selalu terjamin validitasnya di setiap saat sepanjang masa hidup objek. Sebagai contoh, saldo rekening perbankan tidak boleh bernilai negatif secara sembarangan, usia pasien rumah sakit harus berada pada rentang biologis logis (0 hingga 130 tahun), dan koordinat GPS harus berada pada rentang batas bola bumi (-90 s.d. +90 untuk latitude)."
    )
    builder.add_paragraph(
        "Jika seluruh properti dibiarkan public, pihak luar dapat merusak invariant kapan saja ($rekening->saldo = -999999;). Melalui enkapsulasi, akses penulisan dikunci rapat dan hanya dapat dilakukan melalui method yang memeriksa invariant terlebih dahulu."
    )

    builder.add_heading_3("Prinsip 'Tell, Don't Ask' vs Antipattern Anemic Domain Model:")
    builder.add_paragraph(
        "Salah satu jebakan yang sering dialami oleh pemula adalah menerapkan enkapsulasi setengah hati dengan membuat semua properti private, namun langsung membuatkan method getter dan setter untuk setiap properti tanpa logika bisnis apa pun. Pendekatan ini menghasilkan Anemic Domain Model (objek yang hanya menjadi kantong data pasif tanpa kecerdasan bisnis)."
    )
    builder.add_paragraph(
        "Prinsip desain berorientasi objek yang elegan adalah Tell, Don't Ask (Perintahkan, Jangan Bertanya): Alih-alih kode luar mengambil saldo lalu menghitung sendiri pemotongan uang dan meng-set nilai baru, kode luar cukup memerintahkan objek rekening untuk mengeksekusi penarikan uang ($rekening->tarikTunai($nominal, $pin);). Objek rekening secara mandiri akan memvalidasi otorisasi PIN, memastikan kecukupan saldo, mengurangi saldo internal, dan mencatat riwayat mutasi audit."
    )

    builder.add_heading_2("4.2 Tiga Tingkat Hak Akses (Visibility Modifiers) di PHP")
    builder.add_paragraph(
        "PHP menyediakan tiga kata kunci pengatur hak akses (Visibility Modifiers) yang dapat disematkan pada properti, method, dan konstanta class:"
    )

    table_vis = [
        ["Modifier", "Akses dari Dalam Class", "Akses dari Child Class (extends)", "Akses dari Luar (Client Code)", "Tingkat Keamanan"],
        ["public", "✅ Diizinkan penuh", "✅ Diizinkan penuh", "✅ Bebas diakses langsung", "Terbuka (Public API)"],
        ["protected", "✅ Diizinkan penuh", "✅ Diizinkan penuh (Inheritance)", "❌ Dilarang (Fatal Error)", "Internal Keluarga Warisan"],
        ["private", "✅ Diizinkan penuh", "❌ Dilarang (Hanya class asal)", "❌ Dilarang (Terkunci rapat)", "Terkunci Paling Rapat"]
    ]
    builder.add_table(table_vis[0], table_vis[1:])

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "class RekeningInduk\n"
        "{\n"
        "    public string $nomorRekening;      // Bebas dibaca dan ditulis publik\n"
        "    protected float $saldo = 0.0;      // Hanya bisa diakses class ini dan subclass\n"
        "    private string $pinRahasia;        // Terkunci rapat hanya untuk class RekeningInduk\n\n"
        "    public function __construct(string $nomor, float $saldoAwal, string $pin) {\n"
        "        $this->nomorRekening = $nomor;\n"
        "        $this->saldo         = $saldoAwal;\n"
        "        $this->pinRahasia    = $pin;\n"
        "    }\n\n"
        "    protected function verifikasiPin(string $pinInput): bool {\n"
        "        return $this->pinRahasia === $pinInput;\n"
        "    }\n"
        "}\n\n"
        "class RekeningTabungan extends RekeningInduk\n"
        "{\n"
        "    public function ambilUang(float $nominal, string $pinInput): void {\n"
        "        // ✅ Boleh mengakses method protected milik parent:\n"
        "        if (!$this->verifikasiPin($pinInput)) {\n"
        "            throw new \\DomainException(\"Autentikasi PIN gagal!\");\n"
        "        }\n"
        "        if ($nominal > $this->saldo) {\n"
        "            throw new \\UnderflowException(\"Saldo tabungan tidak mencukupi!\");\n"
        "        }\n"
        "        $this->saldo -= $nominal; // ✅ Boleh memodifikasi properti protected parent\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_3("Hak Akses pada Konstanta Class (Class Constants):")
    builder.add_paragraph(
        "Sejak era PHP 7.1+, konstanta di dalam class dapat dikunci menggunakan visibility modifier untuk mencegah kebocoran konfigurasi internal ke pihak luar:"
    )
    builder.add_code(
        "<?php\n"
        "class KonfigurasiKeamanan {\n"
        "    public const VERSI_SISTEM = \"3.2.1\";        // Konfigurasi publik\n"
        "    protected const BATAS_LOGIN_GAGAL = 5;       // Hanya untuk subclass autentikasi\n"
        "    private const KUNCI_ENKRIPSI = \"UUI-KEY-2025\"; // Rahasia internal mutlak\n"
        "}\n"
    )

    builder.add_heading_2("4.3 Getter, Setter, dan Validasi Aturan Domain Bisnis")
    builder.add_paragraph(
        "Getter (Accessor) bertugas membaca nilai properti private secara terkontrol, sedangkan Setter (Mutator) bertugas memodifikasi nilai properti seraya menegakkan aturan validasi bisnis (Business Rules Enforcement)."
    )
    builder.add_paragraph(
        "Penerapan Method Chaining (Fluent Interface) pada setter dilakukan dengan mengembalikan referensi objek saat ini (return $this;), sehingga pengembang dapat memperbarui beberapa nilai properti dalam satu rangkaian instruksi yang ringkas dan ekspresif."
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "class PasienRumahSakit\n"
        "{\n"
        "    private string $rekamMedisId;\n"
        "    private string $namaPasien;\n"
        "    private int $umur;\n"
        "    private float $suhuTubuh;\n\n"
        "    public function __construct(string $id, string $nama, int $umur, float $suhu) {\n"
        "        $this->setRekamMedisId($id);\n"
        "        $this->setNamaPasien($nama);\n"
        "        $this->setUmur($umur);\n"
        "        $this->setSuhuTubuh($suhu);\n"
        "    }\n\n"
        "    public function getRekamMedisId(): string { return $this->rekamMedisId; }\n"
        "    private function setRekamMedisId(string $id): void {\n"
        "        if (!preg_match('/^RM-\\d{5}$/', $id)) {\n"
        "            throw new \\InvalidArgumentException(\"Format ID Rekam Medis harus RM-XXXXX!\");\n"
        "        }\n"
        "        $this->rekamMedisId = $id;\n"
        "    }\n\n"
        "    public function getNamaPasien(): string { return $this->namaPasien; }\n"
        "    public function setNamaPasien(string $nama): self {\n"
        "        $bersih = trim($nama);\n"
        "        if (strlen($bersih) < 3) {\n"
        "            throw new \\InvalidArgumentException(\"Nama pasien minimal 3 karakter!\");\n"
        "        }\n"
        "        $this->namaPasien = $bersih;\n"
        "        return $this; // Fluent Method Chaining\n"
        "    }\n\n"
        "    public function getUmur(): int { return $this->umur; }\n"
        "    public function setUmur(int $umur): self {\n"
        "        if ($umur < 0 || $umur > 130) {\n"
        "            throw new \\InvalidArgumentException(\"Rentang umur tidak realistis: {$umur} tahun!\");\n"
        "        }\n"
        "        $this->umur = $umur;\n"
        "        return $this;\n"
        "    }\n\n"
        "    public function getSuhuTubuh(): float { return $this->suhuTubuh; }\n"
        "    public function setSuhuTubuh(float $suhu): self {\n"
        "        if ($suhu < 30.0 || $suhu > 45.0) {\n"
        "            throw new \\InvalidArgumentException(\"Pengukuran suhu di luar ambang medis manusia!\");\n"
        "        }\n"
        "        $this->suhuTubuh = $suhu;\n"
        "        return $this;\n"
        "    }\n"
        "}\n\n"
        "// Penggunaan Fluent Setter Chaining:\n"
        "$pasien = new PasienRumahSakit(\"RM-10245\", \"Teuku Iskandar\", 28, 36.6);\n"
        "$pasien->setNamaPasien(\"Teuku Iskandar Muda\")\n"
        "       ->setUmur(29)\n"
        "       ->setSuhuTubuh(37.1);\n"
    )

    builder.add_heading_2("4.4 Konsep Kekekalan Data (Immutability): readonly di PHP 8.1 & 8.2")
    builder.add_paragraph(
        "Dalam arsitektur perangkat lunak modern (seperti Domain-Driven Design / DDD), konsep Immutability (ketetapan nilai objek) memegang peranan krusial. Objek yang immutable (disebut Value Object) tidak dapat diubah status datanya setelah pertama kali diinisialisasi di constructor. Sifat kekal ini mencegah side-effects tak terduga, aman saat diproses dalam konkurensi, dan memudahkan debugging."
    )
    builder.add_paragraph(
        "PHP 8.1 memperkenalkan readonly property, dan PHP 8.2 melengkapinya dengan readonly class yang secara otomatis menjadikan seluruh properti class bersifat readonly serta melarang pembuatan properti dinamis."
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "// PHP 8.2: Readonly Class (Value Object Uang)\n"
        "readonly class NilaiMataUang\n"
        "{\n"
        "    public function __construct(\n"
        "        public float $nominal,\n"
        "        public string $kodeMataUang = \"IDR\"\n"
        "    ) {\n"
        "        if ($nominal < 0) {\n"
        "            throw new \\InvalidArgumentException(\"Nominal uang tidak boleh negatif!\");\n"
        "        }\n"
        "    }\n\n"
        "    // Operasi matematika menghasilkan Objek BARU (tidak memodifikasi status objek lama)\n"
        "    public function tambah(NilaiMataUang $lain): self {\n"
        "        if ($this->kodeMataUang !== $lain->kodeMataUang) {\n"
        "            throw new \\InvalidArgumentException(\"Mata uang harus sejenis!\");\n"
        "        }\n"
        "        return new self($this->nominal + $lain->nominal, $this->kodeMataUang);\n"
        "    }\n"
        "}\n\n"
        "$uang1 = new NilaiMataUang(50_000, \"IDR\");\n"
        "$uang2 = new NilaiMataUang(25_000, \"IDR\");\n"
        "$total = $uang1->tambah($uang2); // Objek baru dengan nominal Rp 75.000\n"
    )

    builder.add_heading_2("4.5 Paradigma Mutakhir PHP 8.4: Asymmetric Visibility & Property Hooks")
    builder.add_paragraph(
        "Rilis PHP 8.4 menghadirkan terobosan terbesar dalam evolusi enkapsulasi dengan dua fitur revolusioner:"
    )

    builder.add_heading_3("1. Asymmetric Visibility (public private(set)):")
    builder.add_paragraph(
        "Seringkali kita menginginkan suatu properti dapat dibaca secara bebas oleh pihak luar (public read), namun perubahan nilainya hanya boleh dilakukan oleh internal class itu sendiri (private write). Di masa lalu, kita terpaksa mendeklarasikan properti private dan menuliskan getter method secara berulang-ulang."
    )
    builder.add_paragraph(
        "Di PHP 8.4, hal ini diselesaikan secara native dengan sintaks public private(set) atau public protected(set):"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "// PHP 8.4: Asymmetric Visibility\n"
        "class AnggotaPerpustakaan\n"
        "{\n"
        "    // Publik bebas membaca ($a->id, $a->totalPoin), tapi penulisan dikunci di internal class:\n"
        "    public private(set) string $idAnggota;\n"
        "    public private(set) string $nama;\n"
        "    public private(set) int $totalPoinAktivitas = 0;\n\n"
        "    public function __construct(string $id, string $nama) {\n"
        "        $this->idAnggota = $id;\n"
        "        $this->nama      = $nama;\n"
        "    }\n\n"
        "    public function tambahPoinKunjungan(): void {\n"
        "        $this->totalPoinAktivitas += 10; // ✅ Diizinkan: mutasi dari dalam class\n"
        "    }\n"
        "}\n\n"
        "$mhs = new AnggotaPerpustakaan(\"MHS-001\", \"Cut Nyak Dhien\");\n"
        "echo $mhs->nama;                // ✅ Boleh dibaca langsung! Output: Cut Nyak Dhien\n"
        "echo $mhs->totalPoinAktivitas;  // ✅ Boleh dibaca langsung! Output: 0\n"
        "// $mhs->totalPoinAktivitas = 999; // ❌ FATAL ERROR: Cannot modify private(set) property from outside!\n"
    )

    builder.add_heading_3("2. Property Hooks (get dan set hooks):")
    builder.add_paragraph(
        "PHP 8.4 memperkenalkan Property Hooks yang memungkinkan penyematan logika validasi, transformasi string, maupun komputasi virtual langsung pada definisi properti tanpa memerlukan method pembantu:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "// PHP 8.4: Property Hooks\n"
        "class ProfilAkun\n"
        "{\n"
        "    // Backed Property dengan Hook Validasi Format Email\n"
        "    public string $email {\n"
        "        get => $this->email;\n"
        "        set(string $nilaiBaru) {\n"
        "            if (!filter_var($nilaiBaru, FILTER_VALIDATE_EMAIL)) {\n"
        "                throw new \\InvalidArgumentException(\"Format email tidak valid!\");\n"
        "            }\n"
        "            $this->email = strtolower(trim($nilaiBaru));\n"
        "        }\n"
        "    }\n\n"
        "    public string $namaDepan;\n"
        "    public string $namaBelakang;\n\n"
        "    // Virtual Property: Tidak memakan alokasi RAM, dikomputasi saat diakses\n"
        "    public string $namaLengkap {\n"
        "        get => \"{$this->namaDepan} {$this->namaBelakang}\";\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("4.6 Magic Methods untuk Kendali Akses Properti Dinamis")
    builder.add_paragraph(
        "PHP menyediakan magic methods khusus untuk mencegat (intercept) upaya pembacaan, penulisan, pemeriksaan, maupun penghapusan properti yang tidak dapat diakses langsung:"
    )
    builder.add_bullet("__get(string $name): mixed", "Mencegat pembacaan properti yang private atau belum didefinisikan.")
    builder.add_bullet("__set(string $name, mixed $value): void", "Mencegat penulisan nilai ke properti yang private atau belum terdefinisi.")
    builder.add_bullet("__isset(string $name): bool", "Mencegat pemanggilan fungsi isset() atau empty() pada properti terenkapsulasi.")
    builder.add_bullet("__unset(string $name): void", "Mencegat pemanggilan instruksi unset() pada properti terenkapsulasi.")

    builder.add_callout(
        "Peringatan PHP 8.2+: Deprekasi Dynamic Properties",
        "Mulai PHP 8.2, penulisan properti dinamis tanpa deklarasi ($obj->propertiBaru = 100;) telah resmi di-deprecate dan memicu peringatan runtime. Seluruh properti wajib dideklarasikan secara eksplisit di dalam class, atau class tersebut harus ditandai dengan atribut #[\\AllowDynamicProperties].",
        "warning"
    )

    builder.add_tip(
        "Tips Keamanan Enkapsulasi: Terapkan Prinsip Least Privilege",
        "Terapkan prinsip Least Privilege secara konsisten pada setiap anggota class. Awali seluruh properti data sebagai 'private'. Hanya naikkan hak akses ke 'protected' jika memang disiapkan untuk diwariskan ke subclass, dan manfaatkan fitur 'public private(set)' pada PHP 8.4+ untuk mengekspos data publik yang terlindungi dari manipulasi luar."
    )

    builder.add_summary_and_questions([
        "Enkapsulasi menyatukan data dan metode seraya melindungi status internal objek (Information Hiding).",
        "State Invariant menjamin bahwa status data bisnis objek selalu valid dari lahir hingga selesai dieksekusi.",
        "Prinsip 'Tell, Don't Ask' menolak Anemic Domain Model dengan menempatkan logika bisnis langsung pada class pemilik data.",
        "Tiga level hak akses: public (terbuka), protected (keluarga warisan), dan private (terkunci rapat).",
        "Fitur readonly (PHP 8.1/8.2) mewujudkan Value Object yang immutable dan aman dari efek samping konkurensi.",
        "PHP 8.4 menghadirkan Asymmetric Visibility (public private(set)) dan Property Hooks (get/set) yang menyederhanakan penulisan kode enkapsulasi kelas industri."
    ], [
        "Jelaskan mengapa pendekatan 'Tell, Don't Ask' lebih unggul dibandingkan membuat Getter dan Setter untuk setiap variabel tanpa aturan validasi!",
        "Uraikan perbedaan fungsional antara `readonly property` di PHP 8.1 dengan `public private(set)` di PHP 8.4!",
        "Rancanglah class `NilaiMahasiswa` dengan enkapsulasi ketat untuk properti `$tugas`, `$uts`, `$uas` (rentang 0-100), method `hitungNilaiAkhir()`, dan method `getHurufMutu()`!",
        "Bagaimana cara mengamankan array mutasi transaksi perbankan agar tidak dapat dimodifikasi oleh kode pemanggil saat dikembalikan melalui getter (Defensive Copying)?"
    ])

    # =========================================================================
    # BAB 5: PEWARISAN (INHERITANCE) DAN KOMPOSISI KODE MENGGUNAKAN TRAIT
    # =========================================================================
    builder.add_bab_title(5, "Pewarisan (Inheritance), Final Keyword, dan Komposisi Trait")
    builder.add_learning_objectives("Sub-CPMK 3", [
        "Memahami fondasi teoretis Inheritance (Pewarisan Sifat), prinsip Taksonomi 'Is-A', serta perbedaan antara Subtyping dan Code Reuse.",
        "Mengimplementasikan pewarisan menggunakan kata kunci extends dan mengelola siklus hidup inisialisasi menggunakan parent::__construct().",
        "Menerapkan teknik Method Overriding untuk menyesuaikan perilaku subclass seraya mempertahankan integritas antarmuka induk.",
        "Mengendalikan dan mengamankan rancangan hierarki menggunakan kata kunci final pada class, method, dan class constants (PHP 8.1+).",
        "Mengatasi keterbatasan Single Inheritance di PHP menggunakan Trait (Horizontal Code Reuse), mengelola resolusi konflik (insteadof dan as), serta memanfaatkan konstanta di dalam Trait (PHP 8.2+).",
        "Menganalisis kelemahan desain hierarki (Fragile Base Class Problem) dan menerapkan prinsip 'Favor Composition over Inheritance'."
    ])

    builder.add_heading_2("5.1 Filosofi dan Fondasi Teoretis Pewarisan")
    builder.add_paragraph(
        "Dalam rekayasa perangkat lunak berorientasi objek, Pewarisan (Inheritance) merupakan pilar kedua yang memungkinkan suatu class baru (Subclass / Child Class) mengadopsi seluruh atribut (state) dan method (behavior) yang telah didefinisikan pada class yang sudah ada sebelumnya (Superclass / Parent Class)."
    )
    builder.add_paragraph(
        "Relasi pewarisan mencerminkan hubungan taksonomi 'Is-A' (Adalah Seorang / Adalah Sebuah): Dosen adalah seorang Sivitas Akademika, Mahasiswa adalah seorang Sivitas Akademika, dan Mobil Listrik adalah sebuah Kendaraan Bermotor. Melalui pewarisan, atribut dan perilaku umum yang berlaku untuk seluruh anggota taksonomi (seperti nomor identitas, nama lengkap, dan email) cukup dituliskan satu kali pada superclass. Hal ini secara langsung mewujudkan prinsip DRY (Don't Repeat Yourself) dan mempermudah pemeliharaan sistem."
    )
    builder.add_paragraph(
        "Namun, seorang arsitek perangkat lunak wajib mewaspadai fenomena Fragile Base Class Problem. Ketika pohon pewarisan dibuat terlalu dalam (misalnya lebih dari 3 atau 4 tingkat hierarki), sistem menjadi sangat rapuh: modifikasi kecil pada kode internal superclass dapat memicu kerusakan berantai (ripple effects) pada puluhan subclass di bawahnya. Oleh karena itu, standar industri merekomendasikan prinsip 'Favor Composition over Inheritance' (Utamakan Komposisi dan Trait dibandingkan Pewarisan Kelas yang Terlalu Dalam)."
    )

    builder.add_heading_2("5.2 Anatomi Pewarisan di PHP: `extends`, `parent::`, dan Method Overriding")
    builder.add_paragraph(
        "Di dalam PHP, proses penurunan sifat dinyatakan dengan kata kunci extends. Ketika sebuah subclass mendefinisikan constructor-nya sendiri, Zend Engine tidak secara otomatis memanggil constructor milik superclass. Oleh sebab itu, pengembang wajib memanggil constructor induk secara eksplisit menggunakan sintaks parent::__construct(...)."
    )
    builder.add_paragraph(
        "Method Overriding adalah kemampuan subclass untuk menulis ulang implementasi method yang diwarisinya dari parent class agar sesuai dengan kebutuhan spesifik subclass bersangkutan, seraya tetap dapat memanggil fungsionalitas dasar parent melalui parent::namaMethod()."
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Domain;\n\n"
        "// Superclass (Parent Class)\n"
        "class SivitasAkademika\n"
        "{\n"
        "    public function __construct(\n"
        "        protected readonly string $nomorIdentitas,\n"
        "        protected string $namaLengkap,\n"
        "        protected string $emailKampus\n"
        "    ) {\n"
        "        if (empty($nomorIdentitas) || empty($namaLengkap)) {\n"
        "            throw new \\InvalidArgumentException(\"Nomor Identitas dan Nama wajib diisi!\");\n"
        "        }\n"
        "    }\n\n"
        "    public function cetakKartuIdentitas(): void {\n"
        "        echo \"========================================\\n\";\n"
        "        echo \"KARTU SIVITAS AKADEMIKA UUI\\n\";\n"
        "        echo \"Nomor ID : {$this->nomorIdentitas}\\n\";\n"
        "        echo \"Nama     : {$this->namaLengkap}\\n\";\n"
        "        echo \"Email    : {$this->emailKampus}\\n\";\n"
        "        echo \"Peran    : \" . static::class . \"\\n\";\n"
        "    }\n\n"
        "    public function hitungBantuanFasilitas(): float {\n"
        "        return 100_000.0; // Bantuan kuota internet dasar universitas\n"
        "    }\n"
        "}\n\n"
        "// Subclass Dosen (Mewarisi SivitasAkademika)\n"
        "class Dosen extends SivitasAkademika\n"
        "{\n"
        "    public function __construct(\n"
        "        string $nidn,\n"
        "        string $namaLengkap,\n"
        "        string $emailKampus,\n"
        "        private string $jabatanFungsional = \"Asisten Ahli\",\n"
        "        private int $sksMengajar = 12\n"
        "    ) {\n"
        "        // 1. Constructor Chaining: Inisialisasi properti parent\n"
        "        parent::__construct($nidn, $namaLengkap, $emailKampus);\n"
        "    }\n\n"
        "    // 2. Method Overriding: Menyesuaikan kalkulasi tunjangan khusus dosen\n"
        "    public function hitungBantuanFasilitas(): float {\n"
        "        $dasar = parent::hitungBantuanFasilitas(); // Mengambil nilai Rp 100.000 dari parent\n"
        "        $tunjanganSks = $this->sksMengajar * 50_000.0;\n"
        "        return $dasar + $tunjanganSks;\n"
        "    }\n\n"
        "    public function cetakKartuIdentitas(): void {\n"
        "        parent::cetakKartuIdentitas();\n"
        "        echo \"Jabatan Fungsional : {$this->jabatanFungsional}\\n\";\n"
        "        echo \"Beban Mengajar     : {$this->sksMengajar} SKS\\n\";\n"
        "        echo \"Total Bantuan Fas. : Rp \" . number_format($this->hitungBantuanFasilitas(), 0, ',', '.') . \"\\n\";\n"
        "        echo \"========================================\\n\";\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("5.3 Pengendalian Hierarki dengan Kata Kunci `final`")
    builder.add_paragraph(
        "Kata kunci final digunakan untuk mengunci elemen arsitektur agar tidak dapat diperluas atau diubah oleh pengembang lain. PHP mendukung tiga tingkat penguncian final:"
    )
    builder.add_bullet("1. final class", "Melarang class dijadikan superclass bagi class lain mana pun. Sangat berguna untuk class konfigurasi keamanan atau utilitas tertutup.")
    builder.add_bullet("2. final method", "Melarang method tertentu di-override oleh subclass turunan. Pola ini merupakan inti dari Template Method Pattern.")
    builder.add_bullet("3. final class constant (PHP 8.1+)", "Mencegah nilai konstanta class ditimpa (overridden) oleh subclass turunan.")

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "// Contoh final class: Terkunci rapat dari pewarisan\n"
        "final class EnkripsiKeamananToken\n"
        "{\n"
        "    // PHP 8.1: final constant\n"
        "    final public const ALGORITMA = \"AES-256-GCM\";\n"
        "}\n\n"
        "class AturanKelulusanProdi\n"
        "{\n"
        "    // final method: Rumus penentuan yudisium dilarang diutak-atik subclass\n"
        "    final public function validasiYudisium(float $ipk, int $totalSks): bool {\n"
        "        return $ipk >= 2.00 && $totalSks >= 144;\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("5.4 Trait: Solusi Horizontal Code Reuse di PHP")
    builder.add_paragraph(
        "PHP menganut model Single Inheritance (setiap class hanya boleh memiliki satu parent langsung). Keterbatasan ini sering menimbulkan kesulitan ketika kita ingin berbagi fungsionalitas umum (misalnya logging, serialisasi JSON, soft deletion) ke banyak class yang tidak berada dalam satu pohon taksonomi warisan."
    )
    builder.add_paragraph(
        "Trait (diteliti oleh Scharli et al., 2003) menyediakan mekanisme Horizontal Code Reuse di mana potongan-potongan perilaku dapat disisipkan ke berbagai class independen menggunakan instruksi use."
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Traits;\n\n"
        "trait AuditLoggableTrait\n"
        "{\n"
        "    public function catatAudit(string $aktivitas): void {\n"
        "        $waktu = date('Y-m-d H:i:s');\n"
        "        echo \"[AUDIT LOG] [{$waktu}] [\" . static::class . \"] {$aktivitas}\\n\";\n"
        "    }\n"
        "}\n\n"
        "trait ExportableJsonTrait\n"
        "{\n"
        "    public function keJson(): string {\n"
        "        return json_encode(get_object_vars($this), JSON_PRETTY_PRINT | JSON_UNESCAPED_UNICODE);\n"
        "    }\n"
        "}\n\n"
        "// PHP 8.2+: Konstanta di dalam Trait\n"
        "trait CacheConfigTrait\n"
        "{\n"
        "    public const DEFAULT_TTL = 3600; // 1 Jam\n"
        "}\n"
    )

    builder.add_heading_2("5.5 Resolusi Konflik Trait: `insteadof` dan `as`")
    builder.add_paragraph(
        "Ketika sebuah class mengimpor dua Trait berbeda yang kebetulan memiliki nama method yang sama, Zend Engine akan memicu Fatal Error akibat bentrok nama (Name Collision). PHP menyediakan operator insteadof (untuk memilih implementasi mana yang menang) dan operator as (untuk memberikan nama alias pada implementasi yang kalah):"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "trait EmailNotifierTrait {\n"
        "    public function kirimPesan(string $teks): void {\n"
        "        echo \"📧 Mengirim Email: {$teks}\\n\";\n"
        "    }\n"
        "}\n\n"
        "trait TelegramNotifierTrait {\n"
        "    public function kirimPesan(string $teks): void {\n"
        "        echo \"📱 Mengirim Bot Telegram: {$teks}\\n\";\n"
        "    }\n"
        "}\n\n"
        "class PusatNotifikasi\n"
        "{\n"
        "    use EmailNotifierTrait, TelegramNotifierTrait {\n"
        "        // 1. Pilih TelegramNotifierTrait untuk method kirimPesan utama:\n"
        "        TelegramNotifierTrait::kirimPesan insteadof EmailNotifierTrait;\n\n"
        "        // 2. Beri nama alias untuk EmailNotifierTrait agar tetap dapat dipanggil:\n"
        "        EmailNotifierTrait::kirimPesan as kirimEmail;\n"
        "    }\n"
        "}\n\n"
        "$notif = new PusatNotifikasi();\n"
        "$notif->kirimPesan(\"Server pulih kembali.\"); // Memanggil Telegram Notifier\n"
        "$notif->kirimEmail(\"Laporan rekapitulasi.\"); // Memanggil Email Notifier\n"
    )

    builder.add_heading_2("5.6 Abstract Method di dalam Trait")
    builder.add_paragraph(
        "Trait juga dapat mendeklarasikan Abstract Method untuk menuntut class pengguna menyediakan method atau data tertentu agar logika internal Trait dapat berjalan dengan sempurna:"
    )

    builder.add_code(
        "<?php\n"
        "trait NomorSuratOtomatisTrait\n"
        "{\n"
        "    // Menuntut class pengguna mengembalikan kode unit\n"
        "    abstract public function getKodeUnit(): string;\n\n"
        "    public function buatNomorSurat(int $urutan): string {\n"
        "        $tahun = date('Y');\n"
        "        return sprintf(\"%04d/UUI-%s/%s\", $urutan, $this->getKodeUnit(), $tahun);\n"
        "    }\n"
        "}\n\n"
        "class SuratKeputusanDekan\n"
        "{\n"
        "    use NomorSuratOtomatisTrait;\n\n"
        "    public function getKodeUnit(): string {\n"
        "        return \"FST-DEKAN\";\n"
        "    }\n"
        "}\n"
    )

    builder.add_tip(
        "Tips Desain Pewarisan: Pedoman Komposisi vs Pewarisan",
        "Gunakan Pewarisan Kelas (extends) hanya jika relasi 'Is-A' terpenuhi secara murni dan tidak ada pelanggaran Liskov Substitution Principle. Untuk fungsionalitas utilitas lintas divisi yang bersifat modular (seperti logging, caching, hashing), selalu utamakan penggunaan Trait atau Dependency Injection."
    )

    builder.add_summary_and_questions([
        "Pewarisan ('extends') membangun relasi taksonomi 'Is-A' dan memaksimalkan penggunaan ulang kode (DRY).",
        "Sintaks parent::__construct() wajib dipanggil saat subclass mendefinisikan constructor-nya sendiri.",
        "Method Overriding memungkinkan penyesuaian perilaku induk seraya memperluas fungsionalitas.",
        "Kata kunci final mengunci class, method, dan konstanta (PHP 8.1+) dari risiko modifikasi liar.",
        "Trait menyediakan solusi Horizontal Code Reuse untuk mengatasi batasan Single Inheritance di PHP.",
        "Operator 'insteadof' dan 'as' menyelesaikan konflik bentrok nama method antar-trait secara presisi."
    ], [
        "Jelaskan mengapa hierarki pewarisan yang terlalu dalam memicu Fragile Base Class Problem!",
        "Bagaimana cara menyelesaikan bentrok nama method ketika sebuah class menggunakan dua Trait yang memiliki nama fungsi identik?",
        "Rancanglah hierarki class: Parent `Kendaraan`, Child `MobilListrik` (menambahkan baterai dan kalkulasi jarak tempuh) yang dilengkapi Trait `AuditLoggableTrait`!",
        "Kapan sebuah konstanta class sebaiknya ditandai dengan kata kunci `final const` di PHP 8.1+?"
    ])

    # =========================================================================
    # BAB 6: POLIMORFISME (POLYMORPHISM) DAN DYNAMIC DISPATCH
    # =========================================================================
    builder.add_bab_title(6, "Polimorfisme (Polymorphism) dan Dynamic Dispatch")
    builder.add_learning_objectives("Sub-CPMK 3", [
        "Memahami filosofi fundamental Polymorphism (Polimorfisme), taksonomi teori tipe Cardelli & Wegner, serta prinsip 'Satu Antarmuka, Banyak Perilaku'.",
        "Memahami mekanisme eksekusi Dynamic Method Dispatch pada Zend Engine runtime PHP.",
        "Mengimplementasikan Polymorphic Type Hinting dan memproses Koleksi Objek Polimorfik (Polymorphic Collections).",
        "Membedakan secara analitis Polimorfisme berbasis Pewarisan Class (Class Inheritance) vs Polimorfisme berbasis Kontrak (Interface-based Polymorphism).",
        "Menggunakan operator instanceof secara tepat untuk Type Narrowing seraya menghindari Anti-pattern Type Checking.",
        "Merancang arsitektur perangkat lunak yang mematuhi Open/Closed Principle (OCP) dan Liskov Substitution Principle (LSP)."
    ])

    builder.add_heading_2("6.1 Filosofi dan Fondasi Teoretis Polimorfisme")
    builder.add_paragraph(
        "Secara etimologi bahasa Yunani, Polimorfisme berasal dari kata poly (banyak) dan morph (bentuk atau rupa). Dalam ilmu rekayasa perangkat lunak berorientasi objek, polimorfisme adalah prinsip kemampuan objek-objek dari berbagai class turunan yang berbeda untuk merespons pemanggilan pesan atau method yang sama dengan cara dan implementasi unik mereka masing-masing."
    )
    builder.add_paragraph(
        "Kekuatan utama polimorfisme terletak pada pemisahan yang tegas antara 'Apa yang harus dilakukan' (didefinisikan pada antarmuka umum / superclass) dan 'Bagaimana cara melakukannya' (didefinisikan secara spesifik oleh masing-masing subclass). Kode pemanggil (Client Code) cukup berkomunikasi dengan tipe acuan abstrak tanpa perlu mengetahui secara detail class konkrit mana yang sedang aktif saat runtime."
    )
    builder.add_paragraph(
        "Dalam literatur klasik teori tipe komputasi (Cardelli & Wegner, 1985), polimorfisme pada OOP bertumpu pada Subtyping / Inclusion Polymorphism, di mana sebuah variabel ber-tipe superclass dapat menampung objek dari subclass mana pun yang sah dan secara otomatis mengeksekusi perilaku yang tepat tanpa memerlukan percabangan logika if-else."
    )

    builder.add_heading_2("6.2 Mekanisme Dynamic Method Dispatch pada PHP Runtime")
    builder.add_paragraph(
        "Di dalam PHP, resolusi pemanggilan method ($objek->bayar()) diselesaikan secara dinamis pada saat program berjalan (runtime). Ketika method dipanggil, Zend Engine memeriksa Virtual Method Table (V-Table) atau tabel simbol kelas dari instance objek yang tersimpan di memori dan melompat langsung ke alamat fungsi yang sesuai."
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Pembayaran;\n\n"
        "// Superclass Abstraksi\n"
        "abstract class SaluranPembayaran\n"
        "{\n"
        "    public function __construct(protected float $totalTagihan) {}\n\n"
        "    abstract public function bayar(): string;\n"
        "    abstract public function hitungBiayaAdmin(): float;\n"
        "}\n\n"
        "// Subclass 1: Transfer Bank\n"
        "class TransferBank extends SaluranPembayaran\n"
        "{\n"
        "    public function __construct(float $total, private string $namaBank, private string $rekening) {\n"
        "        parent::__construct($total);\n"
        "    }\n\n"
        "    public function hitungBiayaAdmin(): float { return 4_000.0; }\n\n"
        "    public function bayar(): string {\n"
        "        $totalBayar = $this->totalTagihan + $this->hitungBiayaAdmin();\n"
        "        return \"🏦 [TRANSFER BANK] {$this->namaBank} ({$this->rekening}) | Total: Rp \" . \n"
        "               number_format($totalBayar, 0, ',', '.') . \" (Termasuk Biaya Admin: Rp 4.000)\";\n"
        "    }\n"
        "}\n\n"
        "// Subclass 2: QRIS Instant\n"
        "class QrisInstant extends SaluranPembayaran\n"
        "{\n"
        "    public function __construct(float $total, private string $merchantNMID) {\n"
        "        parent::__construct($total);\n"
        "    }\n\n"
        "    public function hitungBiayaAdmin(): float { return $this->totalTagihan * 0.007; } // 0.7%\n\n"
        "    public function bayar(): string {\n"
        "        $totalBayar = $this->totalTagihan + $this->hitungBiayaAdmin();\n"
        "        return \"📱 [QRIS INSTANT] NMID: {$this->merchantNMID} | Total: Rp \" . \n"
        "               number_format($totalBayar, 0, ',', '.') . \" [LUNAS REALTIME]\";\n"
        "    }\n"
        "}\n\n"
        "// Subclass 3: Dompet Digital\n"
        "class EWalletGoPay extends SaluranPembayaran\n"
        "{\n"
        "    public function __construct(float $total, private string $nomorHp) {\n"
        "        parent::__construct($total);\n"
        "    }\n\n"
        "    public function hitungBiayaAdmin(): float { return 1_000.0; }\n\n"
        "    public function bayar(): string {\n"
        "        $totalBayar = $this->totalTagihan + $this->hitungBiayaAdmin();\n"
        "        return \"💳 [E-WALLET GOPAY] Akun {$this->nomorHp} terdebet Rp \" . \n"
        "               number_format($totalBayar, 0, ',', '.') . \" [SUKSES]\";\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("6.3 Polymorphic Type Hinting & Koleksi Polimorfik")
    builder.add_paragraph(
        "Dengan Polymorphic Type Hinting, fungsi proses transaksi hanya perlu bergantung pada superclass SaluranPembayaran. Fungsi ini secara otomatis mematuhi Open/Closed Principle (terbuka untuk penambahan saluran baru, namun tertutup dari modifikasi kode lama)."
    )

    builder.add_code(
        "<?php\n"
        "// Fungsi Konsumen Polimorfik\n"
        "function prosesTransaksiKasir(SaluranPembayaran $saluran): void {\n"
        "    echo \"Menghubungkan ke payment gateway...\\n\";\n"
        "    $struk = $saluran->bayar(); // Dynamic Dispatch mengeksekusi implementasi child yang tepat\n"
        "    echo $struk . \"\\n\";\n"
        "    echo \"Biaya Admin: Rp \" . number_format($saluran->hitungBiayaAdmin(), 0, ',', '.') . \"\\n\";\n"
        "    echo \"--------------------------------------------------------\\n\";\n"
        "}\n\n"
        "// Array Polimorfik: Berisi ragam subclass berbeda dalam satu wadah seragam\n"
        "$antreanPembayaran = [\n"
        "    new TransferBank(500_000.0, \"Bank Syariah Indonesia (BSI)\", \"7123456789\"),\n"
        "    new QrisInstant(25_000.0, \"ID1020304050\"),\n"
        "    new EWalletGoPay(75_000.0, \"081269001122\"),\n"
        "    new QrisInstant(150_000.0, \"ID1020304050\")\n"
        "];\n\n"
        "foreach ($antreanPembayaran as $transaksi) {\n"
        "    prosesTransaksiKasir($transaksi);\n"
        "}\n"
    )

    builder.add_heading_2("6.4 Matriks Analisis: Polimorfisme Class vs Interface")
    table_poly = [
        ["Dimensi Analisis", "Polimorfisme Berbasis Class (extends)", "Polimorfisme Berbasis Interface (implements)"],
        ["Relasi Konseptual", "Hubungan keluarga taksonomi ketat (Is-A)", "Hubungan kontrak kemampuan perilaku (Can-Do)"],
        ["Pewarisan State", "Mewarisi properti dan method konkrit parent", "Murni kontrak antarmuka tanpa state bersama"],
        ["Fleksibilitas", "Terikat batas Single Inheritance (1 parent)", "Bebas diimplementasikan banyak class (Multiple)"],
        ["Standar Industri", "Hierarki entitas domain inti", "Desain Clean Architecture & Dependency Injection"]
    ]
    builder.add_table(table_poly[0], table_poly[1:])

    builder.add_heading_2("6.5 Type Narrowing Menggunakan Operator `instanceof`")
    builder.add_paragraph(
        "Operator instanceof digunakan ketika kode perlu memeriksa tipe objek spesifik sebelum menjalankan operasi yang hanya ada pada subclass tertentu. Namun, perancang sistem harus menghindari penggunaan instanceof yang berlebihan (Anti-pattern Type Checking), karena hal tersebut mengindikasikan bahwa logika seharusnya ditempatkan ke dalam method polimorfik milik subclass."
    )

    builder.add_code(
        "<?php\n"
        "class AuditorKeuangan\n"
        "{\n"
        "    public function audit(SaluranPembayaran $saluran): void {\n"
        "        echo \"Audit Transaksi: \" . $saluran->bayar() . \"\\n\";\n"
        "        if ($saluran instanceof QrisInstant) {\n"
        "            echo \"ℹ️ Validasi signature QRIS dengan Bank Indonesia.\\n\";\n"
        "        } elseif ($saluran instanceof TransferBank) {\n"
        "            echo \"ℹ️ Rekonsiliasi mutasi rekening koran perbankan.\\n\";\n"
        "        }\n"
        "    }\n"
        "}\n"
    )

    builder.add_tip(
        "Tips Polimorfisme: Desain Berbasis Kontrak untuk Kemudahan Pengujian",
        "Gunakan Polymorphic Type Hinting pada seluruh service layer aplikasi Anda. Dengan mengandalkan interface atau tipe abstrak, Anda dapat dengan mudah membuat Mock Object (objek tiruan) saat menjalankan Automated Unit Testing (PHPUnit) tanpa perlu menghubungkan sistem ke database atau payment gateway nyata."
    )

    builder.add_summary_and_questions([
        "Polimorfisme ('Satu Antarmuka, Banyak Wujud') memisahkan antarmuka umum dari implementasi spesifik.",
        "Dynamic Method Dispatch pada Zend Engine menyelesaikan pemanggilan method secara dinamis di runtime.",
        "Polymorphic Type Hinting memungkinkan penulisan kode modular yang mematuhi Open/Closed Principle.",
        "Polimorfisme berbasis Interface memberikan fleksibilitas tertinggi tanpa batasan Single Inheritance.",
        "Operator instanceof menjamin Type Safety ketika diperlukan inspeksi tipe objek spesifik."
    ], [
        "Jelaskan bagaimana polimorfisme memfasilitasi penambahan metode pembayaran Cryptocurrency baru tanpa memodifikasi kode kasir yang sudah ada!",
        "Apa perbedaan mendasar antara Polimorfisme Berbasis Pewarisan Class vs Polimorfisme Berbasis Interface?",
        "Rancanglah sebuah sistem notifikasi polimorfik (Email, SMS, WhatsApp) dengan antarmuka `kirim(string $tujuan, string $pesan)`!",
        "Mengapa terlalu banyak menggunakan percabangan `if ($obj instanceof X)` di dalam alur logika bisnis dianggap sebagai code smell?"
    ])

    # =========================================================================
    # BAB 7: ABSTRAKSI: INTERFACE, ABSTRACT CLASS, DAN BACKED ENUM
    # =========================================================================
    builder.add_bab_title(7, "Abstraksi: Abstract Class, Interface, dan Backed Enum")
    builder.add_learning_objectives("Sub-CPMK 3", [
        "Memahami filosofi fundamental Abstraction (Abstraksi) sebagai pilar ke-4 OOP, konsep Separation of Interface and Implementation, serta Design by Contract (DbC).",
        "Mendeklarasikan dan menerapkan Abstract Class serta Abstract Method menggunakan kata kunci abstract dan merancang pola Template Method Pattern.",
        "Merancang kontrak antarmuka murni menggunakan Interface, kata kunci implements, serta pewarisan antar-interface (Interface Inheritance).",
        "Mengimplementasikan Multiple Interfaces pada sebuah Class untuk mengatasi keterbatasan pewarisan tunggal.",
        "Membedakan secara tajam kapan harus menggunakan Abstract Class (IS-A) vs Interface (CAN-DO) dalam arsitektur perangkat lunak skala enterprise.",
        "Mengintegrasikan fitur modern Backed Enum (PHP 8.1+) dengan method fungsional dan pattern matching (match) untuk manajemen status yang type-safe."
    ])

    builder.add_heading_2("7.1 Filosofi dan Fondasi Teoretis Abstraksi")
    builder.add_paragraph(
        "Dalam rekayasa perangkat lunak, Abstraksi (Abstraction) adalah proses menyederhanakan kompleksitas sistem dengan hanya menampilkan karakteristik dan antarmuka penting kepada dunia luar, seraya menyembunyikan mekanisme teknis internal yang rumit."
    )
    builder.add_paragraph(
        "Sebagai analogi dunia nyata, ketika seseorang mengemudikan mobil, ia cukup berinteraksi dengan pedal gas, pedal rem, dan roda kemudi. Pengemudi tidak perlu mengetahui secara mikroskopis rasio kompresi bahan bakar di dalam ruang silinder atau perpindahan fluida transmisi hidrolik. Antarmuka pedal menyederhanakan kompleksitas mesin tersebut."
    )
    builder.add_paragraph(
        "Dalam bahasa PHP modern, pilar abstraksi diwujudkan melalui dua konstruksi utama: Abstract Class (kerangka dasar setengah jadi untuk hierarki keluarga erat IS-A) dan Interface (kontrak perilaku murni untuk kemampuan lintas modul CAN-DO)."
    )

    builder.add_heading_2("7.2 Abstract Class & Template Method Pattern")
    builder.add_paragraph(
        "Abstract Class adalah class induk yang tidak dapat diinstansiasi langsung menggunakan operator new. Class ini memuat gabungan antara Concrete Method (method yang sudah memiliki kode fungsional teruji) dan Abstract Method (method tanpa badan fungsi yang wajib disempurnakan oleh subclass turunan)."
    )
    builder.add_paragraph(
        "Pola desain Template Method Pattern mengunci alur kerja utama (master workflow) pada parent class menggunakan kata kunci final, sementara langkah-langkah detailnya diserahkan kepada subclass melalui abstract method:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Laporan;\n\n"
        "// Abstract Superclass\n"
        "abstract class TemplateLaporanAkademik\n"
        "{\n"
        "    public function __construct(\n"
        "        protected string $judulLaporan,\n"
        "        protected string $semester\n"
        "    ) {}\n\n"
        "    // 1. Concrete Method: Kop Surat Standar Universitas\n"
        "    public function cetakKopSurat(): void {\n"
        "        echo \"========================================================\\n\";\n"
        "        echo \"UNIVERSITAS UBUDIYAH INDONESIA\\n\";\n"
        "        echo \"FAKULTAS SAINS DAN TEKNOLOGI - PROGRAM STUDI INFORMATIKA\\n\";\n"
        "        echo \"Judul Laporan : {$this->judulLaporan}\\n\";\n"
        "        echo \"Semester      : {$this->semester}\\n\";\n"
        "        echo \"--------------------------------------------------------\\n\";\n"
        "    }\n\n"
        "    // 2. Abstract Methods: Wajib disediakan oleh subclass\n"
        "    abstract protected function ambilSumberData(): array;\n"
        "    abstract protected function susunBadanLaporan(array $data): string;\n"
        "    abstract public function exportFormat(): string;\n\n"
        "    // 3. Template Method (Final): Alur kerja utama terkunci aman\n"
        "    final public function generateDokumen(): void {\n"
        "        $this->cetakKopSurat();\n"
        "        $data = $this->ambilSumberData();\n"
        "        $konten = $this->susunBadanLaporan($data);\n"
        "        echo $konten . \"\\n\";\n"
        "        echo \"Format Dokumen: \" . $this->exportFormat() . \"\\n\";\n"
        "        echo \"========================================================\\n\";\n"
        "    }\n"
        "}\n\n"
        "// Subclass Konkrit: Rekapitulasi IPK Mahasiswa\n"
        "class LaporanIpMahasiswa extends TemplateLaporanAkademik\n"
        "{\n"
        "    protected function ambilSumberData(): array {\n"
        "        return [\n"
        "            ['nim' => '240101', 'nama' => 'Cut Meurah Intan', 'ipk' => 3.92],\n"
        "            ['nim' => '240102', 'nama' => 'Teuku Rayhan', 'ipk' => 3.85]\n"
        "        ];\n"
        "    }\n\n"
        "    protected function susunBadanLaporan(array $data): string {\n"
        "        $out = \"REKAPITULASI IPK MAHASISWA:\\n\";\n"
        "        foreach ($data as $mhs) {\n"
        "            $out .= sprintf(\"• [%s] %-20s : IPK %.2f\\n\", $mhs['nim'], $mhs['nama'], $mhs['ipk']);\n"
        "        }\n"
        "        return $out;\n"
        "    }\n\n"
        "    public function exportFormat(): string {\n"
        "        return \"Dokumen Portabel (PDF A4 Landscape)\";\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("7.3 Interface: Kontrak Perilaku Murni (*CAN-DO*)")
    builder.add_paragraph(
        "Interface adalah kontrak murni tanpa properti data dan tanpa implementasi method. Seluruh method yang dideklarasikan di dalam interface otomatis bersifat public abstract. Sebuah class dapat mengimplementasikan banyak interface sekaligus (Multiple Interface Implementation)."
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Kontrak;\n\n"
        "interface ExportablePdfInterface { public function renderPdf(): string; }\n"
        "interface KirimEmailInterface { public function kirimEmail(string $tujuan): bool; }\n"
        "interface AuditLoggableInterface { public function catatLog(string $pesan): void; }\n\n"
        "// Implementasi Multiple Interfaces\n"
        "class BerkasTranskripNilai implements ExportablePdfInterface, KirimEmailInterface, AuditLoggableInterface\n"
        "{\n"
        "    public function __construct(public readonly string $nim, public readonly string $nama) {}\n\n"
        "    public function renderPdf(): string {\n"
        "        return \"[PDF-BINARY] Transkrip Resmi {$this->nama} ({$this->nim}) siap diunduh.\";\n"
        "    }\n\n"
        "    public function kirimEmail(string $tujuan): bool {\n"
        "        echo \"📧 Mengirim transkrip ke <{$tujuan}>... Berhasil!\\n\";\n"
        "        $this->catatLog(\"Transkrip dikirim ke {$tujuan}\");\n"
        "        return true;\n"
        "    }\n\n"
        "    public function catatLog(string $pesan): void {\n"
        "        echo \"📝 [LOG] \" . date('Y-m-d H:i:s') . \" - {$pesan}\\n\";\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("7.4 Matriks Komparasi: Abstract Class vs Interface")
    table_abs = [
        ["Parameter Analisis", "Abstract Class", "Interface"],
        ["Kata Kunci", "abstract class + extends", "interface + implements"],
        ["Relasi Konseptual", "IS-A (Identitas kekeluargaan erat)", "CAN-DO (Kontrak kemampuan/perilaku)"],
        ["Pewarisan Ganda", "❌ Dilarang (Single Inheritance)", "✅ Diizinkan (Multiple Implementation)"],
        ["Properti / State", "✅ Boleh (public, protected, private)", "❌ Dilarang (Hanya konstanta const)"],
        ["Implementasi Method", "Campuran (Bisa konkret + abstract)", "Murni deklarasi tanpa kurung kurawal {}"],
        ["Constructor", "✅ Bisa memiliki __construct()", "❌ Tidak boleh memiliki constructor"]
    ]
    builder.add_table(table_abs[0], table_abs[1:])

    builder.add_heading_2("7.5 Backed Enum di PHP 8.1+ dan Pattern Matching")
    builder.add_paragraph(
        "PHP 8.1 memperkenalkan Backed Enum (tipe data enumerasi yang nilainya terikat pada string atau integer). Backed Enum sangat ideal dipadukan dengan Interface dan Abstract Class untuk mengelola status sistem secara type-safe serta mendukung ekspresi match:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Enums;\n\n"
        "interface LabelledEnumInterface {\n"
        "    public function getLabel(): string;\n"
        "    public function getBadgeColor(): string;\n"
        "}\n\n"
        "enum StatusKelulusan: string implements LabelledEnumInterface\n"
        "{\n"
        "    case LULUS_CUMLAUDE  = 'CUMLAUDE';\n"
        "    case LULUS_MEMUASKAN = 'MEMUASKAN';\n"
        "    case BERSYARAT       = 'BERSYARAT';\n"
        "    case MENGULANG       = 'MENGULANG';\n\n"
        "    public function getLabel(): string {\n"
        "        return match($this) {\n"
        "            self::LULUS_CUMLAUDE  => 'Lulus dengan Pujian (Cum Laude)',\n"
        "            self::LULUS_MEMUASKAN => 'Lulus Sangat Memuaskan',\n"
        "            self::BERSYARAT       => 'Lulus Bersyarat (Revisi Skripsi)',\n"
        "            self::MENGULANG       => 'Wajib Mengulang Sidang',\n"
        "        };\n"
        "    }\n\n"
        "    public function getBadgeColor(): string {\n"
        "        return match($this) {\n"
        "            self::LULUS_CUMLAUDE  => '#10B981',\n"
        "            self::LULUS_MEMUASKAN => '#3B82F6',\n"
        "            self::BERSYARAT       => '#F59E0B',\n"
        "            self::MENGULANG       => '#EF4444',\n"
        "        };\n"
        "    }\n"
        "}\n\n"
        "$status = StatusKelulusan::from('CUMLAUDE');\n"
        "echo \"Predikat : \" . $status->getLabel() . \"\\n\";\n"
        "echo \"Warna Tag: \" . $status->getBadgeColor() . \"\\n\";\n"
    )

    builder.add_tip(
        "Tips Abstraksi: Prinsip 'Program to an Interface, not an Implementation'",
        "Rancanglah lapisan bisnis aplikasi Anda dengan selalu bergantung pada Interface (abstraksi), bukan pada class konkrit. Dengan cara ini, Anda dapat mengganti pustaka pihak ketiga (misal: mengganti Mailgun dengan AWS SES, atau mengganti MySQL dengan PostgreSQL) tanpa perlu mengubah kode domain aplikasi sama sekali."
    )

    builder.add_summary_and_questions([
        "Abstraksi menyembunyikan detail teknis yang rumit dan menyajikan antarmuka esensial kepada pengguna.",
        "Abstract Class memadukan concrete methods dan abstract methods dalam pola Template Method Pattern.",
        "Interface adalah kontrak murni tanpa state data, memungkinkan Multiple Interfaces pada class.",
        "Pewarisan antar-interface (Interface Inheritance) dapat dilakukan menggunakan kata kunci extends.",
        "Backed Enum (PHP 8.1+) menghadirkan type safety untuk status data dengan dukungan method dan match expression."
    ], [
        "Jelaskan perbedaan mendasar antara relasi 'IS-A' pada Abstract Class dan 'CAN-DO' pada Interface!",
        "Bagaimana cara kerja Template Method Pattern dalam menjaga integritas master workflow sebuah laporan akademik?",
        "Rancanglah sebuah Backed Enum `StatusPembayaran` (PENDING, PAID, EXPIRED, FAILED) yang mengimplementasikan method `isFinal(): bool`!",
        "Mengapa sebuah Interface dilarang memiliki properti variabel dan constructor?"
    ])

    # =========================================================================
    # BAB 8: MANAJEMEN NAMESPACE, STANDAR PSR-4, DAN COMPOSER AUTOLOADING
    # =========================================================================
    builder.add_bab_title(8, "Manajemen Namespace, Standar PSR-4, dan Composer Autoloading")
    builder.add_learning_objectives("Sub-CPMK 4", [
        "Memahami urgensi dan filosofi Namespace untuk mencegah polusi ruang lingkup global (Global Namespace Pollution) dan tabrakan nama class (Name Collision).",
        "Menguasai sintaks deklarasi namespace, aturan resolusi nama simbol (Fully Qualified, Qualified, Unqualified), serta penggunaan instruksi use, aliasing (as), dan Group Use.",
        "Memahami evolusi mekanisme Autoloading di PHP dari era require_once, spl_autoload_register(), hingga standar industri modern PSR-4 (PHP Standards Recommendation).",
        "Mengonfigurasi berkas composer.json untuk mengelola pemetaan namespace aplikasi (autoload.psr-4) dan lingkungan pengujian (autoload-dev.psr-4).",
        "Menguasai teknik optimasi performa ClassLoader pada server produksi (composer dump-autoload -o dan --classmap-authoritative).",
        "Menyusun struktur direktori proyek PHP skala besar berstandar arsitektur industri (Clean / Layered Architecture)."
    ])

    builder.add_heading_2("8.1 Filosofi Namespace: Mengatasi Polusi Lingkup Global")
    builder.add_paragraph(
        "Pada masa awal perkembangan PHP (sebelum versi 5.3), seluruh class, fungsi, dan konstanta yang didefinisikan akan ditempatkan di dalam satu wadah global yang sama (Global Scope). Hal ini memicu masalah kritis ketika dua pustaka berbeda mendefinisikan nama class yang sama, menghasilkan Fatal Error: Cannot redeclare class."
    )
    builder.add_paragraph(
        "Untuk mencegah tabrakan nama tersebut, para pengembang zaman dahulu terpaksa memberi awalan nama class yang sangat panjang (misalnya Zend_Service_Amazon_Ec2_Instance_Configuration). Namespace menyediakan mekanisme pembungkus virtual hierarkis yang mengisolasi simbol-simbol kode ke dalam ruang nama masing-masing, mirip dengan struktur direktori di dalam sistem operasi komputer."
    )

    builder.add_heading_2("8.2 Anatomi Deklarasi dan Resolusi Simbol Namespace")
    builder.add_paragraph(
        "PHP mengenali tiga tingkatan cara penulisan nama class atau simbol:"
    )
    builder.add_bullet("1. Fully Qualified Name (FQN)", "Dimulai dengan garis miring terbalik (\\), menunjuk langsung dari akar global (contoh: \\App\\Domain\\Model\\Mahasiswa).")
    builder.add_bullet("2. Qualified Name", "Mengandung garis miring terbalik di tengah namun tidak diawali \\ (contoh: Model\\Mahasiswa).")
    builder.add_bullet("3. Unqualified Name", "Hanya menuliskan nama class secara langsung tanpa garis miring terbalik (contoh: Mahasiswa).")

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Controller;\n\n"
        "// 1. Mengimpor Class Tunggal dengan Alias untuk Mencegah Bentrok Nama\n"
        "use App\\Akademik\\Mahasiswa as MahasiswaAkademik;\n"
        "use App\\Penerimaan\\Mahasiswa as MahasiswaPendaftar;\n\n"
        "// 2. Mengimpor Fungsi dan Konstanta Spesifik\n"
        "use function App\\Helper\\formatMataUang;\n"
        "use const App\\Config\\MAX_SKS_SEMESTER;\n\n"
        "// 3. Group Use Statements (PHP 7+)\n"
        "use App\\Services\\{\n"
        "    RegistrasiKrsService,\n"
        "    ValidasiSyaratYudisiumService,\n"
        "    AuditLogService\n"
        "};\n\n"
        "// 4. Mengakses Global PHP Built-in Classes\n"
        "use DateTimeImmutable;\n"
        "use InvalidArgumentException;\n\n"
        "class PendaftaranController\n"
        "{\n"
        "    public function prosesPendaftaran(): void {\n"
        "        $mhsBaru = new MahasiswaPendaftar(\"REG-2025-001\", \"Cut Meurah\");\n"
        "        $mhsAktif = new MahasiswaAkademik(\"240101001\", \"Teuku Iskandar\");\n\n"
        "        $waktu = new DateTimeImmutable();\n"
        "        echo \"Pendaftaran diproses pada: \" . $waktu->format('Y-m-d H:i:s') . \"\\n\";\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("8.3 Standar Autoloading: Dari `require` ke Standar PSR-4")
    builder.add_paragraph(
        "Dahulu, setiap class harus disertakan secara manual menggunakan require_once, yang mengakibatkan fenomena Spaghetti Includes dan membebani performa I/O disk server. PHP kemudian memperkenalkan spl_autoload_register() yang memungkinkan pendaftaran fungsi penangkap (callback) otomatis ketika suatu class belum dimuat di memori."
    )
    builder.add_paragraph(
        "Konsorsium PHP-FIG (PHP Standards Recommendation) menetapkan standar PSR-4 yang mengatur pemetaan matematis antara Namespace terstruktur dengan direktori fisik file di server: Namespace Prefix App\\ dipetakan ke direktori src/, sehingga class App\\Domain\\Model\\Mahasiswa secara otomatis dimuat dari berkas fisik src/Domain/Model/Mahasiswa.php."
    )

    builder.add_heading_2("8.4 Konfigurasi Autoloading Menggunakan Composer")
    builder.add_paragraph(
        "Composer adalah Dependency & Package Manager resmi PHP yang mengelola autoloader PSR-4 secara otomatis melalui berkas konfigurasi composer.json:"
    )

    builder.add_code(
        "// composer.json\n"
        "{\n"
        "    \"name\": \"uui/sistem-informasi-akademik\",\n"
        "    \"description\": \"Sistem Informasi Akademik Berbasis OOP PHP 8+\",\n"
        "    \"type\": \"project\",\n"
        "    \"license\": \"MIT\",\n"
        "    \"autoload\": {\n"
        "        \"psr-4\": {\n"
        "            \"App\\\\\": \"src/\"\n"
        "        },\n"
        "        \"files\": [\n"
        "            \"src/Helper/GlobalFunctions.php\"\n"
        "        ]\n"
        "    },\n"
        "    \"autoload-dev\": {\n"
        "        \"psr-4\": {\n"
        "            \"Tests\\\\\": \"tests/\"\n"
        "        }\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_3("Perintah Kunci Manajemen Autoloader Composer:")
    builder.add_paragraph(
        "1. composer dump-autoload: Memperbarui peta berkas autoload saat ada penambahan file atau perubahan namespace.\n"
        "2. composer dump-autoload -o (--optimize): Menghasilkan tabel hash ClassMap Level 1 untuk performa tinggi di lingkungan staging/produksi.\n"
        "3. composer dump-autoload -a (--classmap-authoritative): Mengunci ClassMap secara absolut dan melarang pencarian sistem berkas dinamis saat runtime, menghasilkan kecepatan akses kelas maksimal."
    )

    builder.add_heading_2("8.5 Struktur Standar Proyek PHP Skala Enterprise (Clean Architecture)")
    builder.add_paragraph(
        "Berikut adalah pola struktur folder profesional yang memisahkan tanggung jawab domain, infrastruktur, dan service layer:"
    )

    builder.add_code(
        "sistem-akademik/\n"
        "├── composer.json\n"
        "├── vendor/                      # Direktori pustaka & autoloader\n"
        "│   └── autoload.php\n"
        "├── src/                         # Kode aplikasi (App\\)\n"
        "│   ├── Domain/\n"
        "│   │   ├── Model/               # App\\Domain\\Model\\Mahasiswa\n"
        "│   │   └── Repository/          # App\\Domain\\Repository\\MahasiswaRepoInterface\n"
        "│   ├── Infrastructure/\n"
        "│   │   └── Persistence/         # App\\Infrastructure\\Persistence\\JsonMahasiswaRepo\n"
        "│   ├── Service/                 # App\\Service\\RegistrasiKrsService\n"
        "│   └── Helper/                  # App\\Helper\\FormatRupiah\n"
        "├── tests/                       # Unit Testing (Tests\\)\n"
        "│   └── Domain/MahasiswaTest.php\n"
        "└── public/\n"
        "    └── index.php                # Single Entry Point Aplikasi\n"
    )

    builder.add_code(
        "<?php\n"
        "// File: public/index.php\n"
        "declare(strict_types=1);\n\n"
        "// Cukup muat 1 baris autoloader Composer untuk seluruh aplikasi:\n"
        "require_once __DIR__ . '/../vendor/autoload.php';\n\n"
        "use App\\Domain\\Model\\Mahasiswa;\n"
        "use App\\Service\\KrsEngineService;\n\n"
        "$mhs = new Mahasiswa(\"240101\", \"Cut Nyak Dhien\", 3.92);\n"
        "$engine = new KrsEngineService();\n"
        "$maxSks = $engine->hitungMaksimalSks($mhs);\n\n"
        "echo \"Mahasiswa {$mhs->nama} memenuhi syarat mengambil maksimal {$maxSks} SKS.\\n\";\n"
    )

    builder.add_tip(
        "Tips Composer: Optimasi Produksi dengan Flag --classmap-authoritative",
        "Saat melakukan deployment ke server produksi (Production CI/CD Pipeline), selalu jalankan 'composer install --no-dev --optimize-autoloader --classmap-authoritative'. Langkah ini menghilangkan overhead pemeriksaan filesystem stat() pada disk sehingga latency aplikasi PHP berkurang drastis."
    )

    builder.add_summary_and_questions([
        "Namespace mengisolasi class, fungsi, dan konstanta guna mencegah tabrakan nama simbol (Name Collision).",
        "Instruksi 'use' dan aliasing ('as') menyederhanakan penulisan Fully Qualified Name yang panjang.",
        "Standar PSR-4 menetapkan pemetaan matematis antara Namespace dan struktur folder fisik proyek.",
        "Composer mengotomatisasi autoloading class dan menyediakan fitur optimasi classmap untuk server produksi.",
        "Pemisahan folder terstruktur (Domain, Infrastructure, Service) mewujudkan Clean Architecture yang maintainable."
    ], [
        "Jelaskan bahaya Global Namespace Pollution pada proyek PHP skala besar dan bagaimana Namespace menyelesaikannya!",
        "Uraikan perbedaan antara Fully Qualified Name, Qualified Name, dan Unqualified Name dalam aturan resolusi simbol PHP!",
        "Bagaimana cara kerja perintah `composer dump-autoload --classmap-authoritative` dalam meningkatkan performa server produksi?",
        "Rancanglah konfigurasi `composer.json` untuk proyek yang memetakan namespace `Universitas\\Akademik\\` ke folder `src/` dan namespace `Tests\\` ke folder `tests/`!"
    ])

    # =========================================================================
    # BAB 9: PENANGANAN KESALAHAN (EXCEPTION HANDLING) & ROBUST ERROR FLOW
    # =========================================================================
    builder.add_bab_title(9, "Penanganan Kesalahan (Exception Handling) & Robust Error Flow")
    builder.add_learning_objectives("Sub-CPMK 4", [
        "Memahami filosofi Exception Handling dan prinsip Fail-Fast dibandingkan metode error status code konvensional.",
        "Menguasai pohon hierarki antarmuka Throwable, serta membedakan secara tajam antara Error, LogicException, dan RuntimeException.",
        "Menerapkan blok kontrol try-catch-finally, teknik Multi-Catch Exceptions, dan penangkapan anonim (Non-capturing Catches di PHP 8.0+).",
        "Menggunakan fitur throw as an Expression (PHP 8.0+) pada Null Coalescing dan ekspresi match.",
        "Merancang Custom Domain Exceptions yang kaya konteks (Rich Contextual Exceptions) serta menerapkan Exception Chaining / Wrapping untuk melestarikan root cause stack trace.",
        "Mengonfigurasi penanganan kesalahan terpusat (Global Exception Handler) berstandar keamanan enterprise."
    ])

    builder.add_heading_2("9.1 Filosofi Exception Handling & Hierarki `Throwable`")
    builder.add_paragraph(
        "Dalam rekayasa perangkat lunak modern, penanganan kesalahan dengan mengembalikan kode status (seperti return false atau return -1) terbukti berbahaya karena rawan diabaikan oleh kode pemanggil, memicu fenomena Silent Failure di mana data korup tetap diproses lebih jauh."
    )
    builder.add_paragraph(
        "Prinsip Fail-Fast menuntut program untuk segera menghentikan alur kerja normal dan melempar Exception terstruktur begitu terjadi anomali atau pelanggaran aturan bisnis. Sejak PHP 7+, seluruh galat disatukan di bawah payung antarmuka Throwable:"
    )
    builder.add_bullet("1. \\Error", "Mewakili kegagalan fatal pada mesin Zend Engine (misal TypeError, DivisionByZeroError, ValueError, UnhandledMatchError).")
    builder.add_bullet("2. \\LogicException", "Mewakili bug logika yang seharusnya dapat dicegah saat penulisan kode (misal InvalidArgumentException, DomainException, OutOfBoundsException).")
    builder.add_bullet("3. \\RuntimeException", "Mewakili kegagalan lingkungan runtime yang tidak dapat diprediksi saat kompilasi (misal PDOException, koneksi jaringan terputus, disk penuh).")

    builder.add_heading_2("9.2 Blok Kontrol `try-catch-finally` & Multi-Catch Syntax")
    builder.add_paragraph(
        "Blok try-catch menangkap eksepsi yang dilempar, di mana blok catch harus diurutkan dari tipe yang paling spesifik ke tipe yang paling umum. Fitur Multi-Catch memungkinkan beberapa tipe eksepsi ditangani dalam satu blok yang sama menggunakan simbol pipe (|):"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Keuangan;\n\n"
        "use InvalidArgumentException;\n"
        "use DomainException;\n"
        "use Exception;\n"
        "use Throwable;\n\n"
        "function kalkulasiBagiHasil(float $modal, float $rasio): float {\n"
        "    if ($modal <= 0) {\n"
        "        throw new InvalidArgumentException(\"Modal investasi harus lebih besar dari 0!\");\n"
        "    }\n"
        "    if ($rasio < 0.0 || $rasio > 1.0) {\n"
        "        throw new DomainException(\"Rasio bagi hasil harus berada pada rentang 0.0 s.d. 1.0!\");\n"
        "    }\n"
        "    return $modal * $rasio;\n"
        "}\n\n"
        "try {\n"
        "    echo \"Memulai kalkulasi investasi...\\n\";\n"
        "    $hasil = kalkulasiBagiHasil(100_000_000.0, 0.25);\n"
        "    echo \"Hasil Bagi Hasil: Rp \" . number_format($hasil, 0, ',', '.') . \"\\n\";\n"
        "} catch (InvalidArgumentException | DomainException $e) {\n"
        "    // Multi-Catch Syntax\n"
        "    echo \"⚠️ [VALIDASI GAGAL] \" . $e->getMessage() . \"\\n\";\n"
        "} catch (Exception $e) {\n"
        "    echo \"❌ [EXCEPTION UMUM] Terjadi kesalahan: \" . $e->getMessage() . \"\\n\";\n"
        "} catch (Throwable $e) {\n"
        "    echo \"🚨 [FATAL ENGINE ERROR] \" . $e->getMessage() . \"\\n\";\n"
        "} finally {\n"
        "    echo \"ℹ️ [CLEANUP] Pelepasan sumber daya / penutupan koneksi selesai.\\n\";\n"
        "}\n"
    )

    builder.add_heading_2("9.3 Fitur Modern PHP 8+: `throw` Expression & Non-Capturing Catches")
    builder.add_paragraph(
        "Sejak PHP 8.0, kata kunci throw bertindak sebagai ekspresi (expression), sehingga dapat dituliskan langsung di dalam operator Null Coalescing, Ternary, maupun Arrow Function:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "class ProfilPengguna\n"
        "{\n"
        "    public function __construct(\n"
        "        public readonly string $username,\n"
        "        public readonly string $email\n"
        "    ) {}\n\n"
        "    public static function dariArray(array $payload): self {\n"
        "        return new self(\n"
        "            // PHP 8.0: Throw Expression langsung pada operator ??\n"
        "            $payload['username'] ?? throw new \\InvalidArgumentException(\"Username wajib diisi!\"),\n"
        "            $payload['email'] ?? throw new \\InvalidArgumentException(\"Email wajib diisi!\")\n"
        "        );\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("9.4 Merancang Custom Domain Exceptions & Exception Chaining")
    builder.add_paragraph(
        "Custom Domain Exception menyimpan data kontekstual yang kaya (Rich Contextual Exception) agar mempermudah penanganan di lapisan presentasi. Exception Chaining / Wrapping membungkus error level rendah (misal query database gagal) ke dalam service exception tingkat tinggi dengan tetap melestarikan root cause asli melalui parameter $previous:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Exception;\n\n"
        "// Rich Contextual Custom Exception\n"
        "class SaldoTidakCukupException extends \\RuntimeException\n"
        "{\n"
        "    public function __construct(\n"
        "        public readonly string $nomorRekening,\n"
        "        public readonly float $saldoTersedia,\n"
        "        public readonly float $nominalDiminta,\n"
        "        int $code = 400,\n"
        "        ?\\Throwable $previous = null\n"
        "    ) {\n"
        "        $pesan = sprintf(\n"
        "            \"Penarikan gagal pada rekening [%s]. Saldo: Rp %s, Diminta: Rp %s.\",\n"
        "            $nomorRekening,\n"
        "            number_format($saldoTersedia, 0, ',', '.'),\n"
        "            number_format($nominalDiminta, 0, ',', '.')\n"
        "        );\n"
        "        parent::__construct($pesan, $code, $previous);\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("9.5 Studi Kasus Transaksi Perbankan Terpadu")
    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Domain;\n\n"
        "class RekeningBank\n"
        "{\n"
        "    public function __construct(\n"
        "        public readonly string $nomorRekening,\n"
        "        private float $saldo\n"
        "    ) {}\n\n"
        "    public function getSaldo(): float { return $this->saldo; }\n\n"
        "    public function tarik(float $nominal): void {\n"
        "        if ($nominal <= 0) {\n"
        "            throw new \\InvalidArgumentException(\"Nominal penarikan harus bernilai positif.\");\n"
        "        }\n"
        "        if ($nominal > $this->saldo) {\n"
        "            throw new \\App\\Exception\\SaldoTidakCukupException($this->nomorRekening, $this->saldo, $nominal);\n"
        "        }\n"
        "        $this->saldo -= $nominal;\n"
        "    }\n"
        "}\n\n"
        "$akun = new RekeningBank(\"UUI-ACC-1024\", 1_000_000.0);\n"
        "try {\n"
        "    $akun->tarik(1_500_000.0); // Memicu Exception\n"
        "} catch (\\App\\Exception\\SaldoTidakCukupException $e) {\n"
        "    echo \"❌ [TRANSAKSI DITOLAK] \" . $e->getMessage() . \"\\n\";\n"
        "    echo \"   Kekurangan Dana: Rp \" . number_format($e->nominalDiminta - $e->saldoTersedia, 0, ',', '.') . \"\\n\";\n"
        "}\n"
    )

    builder.add_tip(
        "Tips Error Handling: Jangan Pernah Menelan Eksepsi Kosong (Empty Catch)",
        "Hindari praktik buruk 'catch (Exception $e) {}' tanpa aksi apa pun. Menelan eksepsi secara diam-diam akan menyembunyikan letak bug sistem. Selalu catat eksepsi ke berkas log terpusat (PSR-3 Logger) atau lemparkan kembali jika tidak dapat ditangani pada lapisan tersebut."
    )

    builder.add_summary_and_questions([
        "Prinsip Fail-Fast menghentikan program seketika saat terjadi anomali untuk mencegah data korup.",
        "Throwable memayungi Error (masalah mesin) dan Exception (kondisi yang dapat ditangani aplikasi).",
        "Blok try-catch-finally mengelola alur eksepsi dan pelepasan sumber daya secara terjamin.",
        "PHP 8.0+ mendukung Throw as an Expression dan Non-capturing Catches yang meringkas kode.",
        "Exception Chaining ($previous) melestarikan riwayat root cause error tanpa membocorkan detail internal ke pengguna."
    ], [
        "Jelaskan perbedaan mendasar antara `LogicException` dan `RuntimeException` dalam arsitektur aplikasi!",
        "Bagaimana cara kerja fitur Multi-Catch Exception dalam menyederhanakan kode penanganan kesalahan?",
        "Rancanglah sebuah Rich Custom Exception `BatasKreditTerlampauiException` dengan properti `$limitKredit` dan `$pemakaianSaatIni`!",
        "Mengapa Exception Chaining (`$previous`) sangat krusial saat melakukan debugging sistem pada server produksi?"
    ])

    # =========================================================================
    # BAB 10: KOLEKSI OBJEK (OBJECT COLLECTIONS) DAN MANIPULASI ARRAY MODERN
    # =========================================================================
    builder.add_bab_title(10, "Koleksi Objek (Object Collections) dan Manipulasi Array Modern")
    builder.add_learning_objectives("Sub-CPMK 4", [
        "Memahami filosofi First-Class Collections (Object Calisthenics) dan batasan keamanan tipe (Type Safety) pada array bawaan PHP.",
        "Membangun Type-Safe Object Collection yang mengenkapsulasi array internal dan mencegah kontaminasi data heterogen.",
        "Mengintegrasikan antarmuka pustaka standar PHP (SPL): Countable, IteratorAggregate, ArrayAccess, dan JsonSerializable.",
        "Menguasai paradigma pemrograman fungsional pada manipulasi data menggunakan array_map, array_filter, array_reduce, dan Arrow Functions (fn() =>).",
        "Menerapkan fitur modern PHP 8+: Array Unpacking dengan String Keys (PHP 8.1+), array_is_list() (PHP 8.1+), serta Method Chaining (Fluent Interface).",
        "Merancang modul rekapitulasi data akademik skala besar berbasis Immutable Collection Transformation."
    ])

    builder.add_heading_2("10.1 Filosofi First-Class Collections & Masalah Type Safety")
    builder.add_paragraph(
        "Array bawaan PHP sangat fleksibel namun bersifat loosely-typed, di mana array dapat menampung campuran berbagai tipe data (integer, string, objek berbeda) dalam satu wadah tanpa adanya pengawasan kompilasi. Hal ini menimbulkan kerapuhan sistem ketika kode mengasumsikan seluruh elemen adalah instansi dari class tertentu."
    )
    builder.add_paragraph(
        "Pola desain First-Class Collection (dicetuskan oleh Jeff Bay dalam Object Calisthenics) mengajarkan bahwa setiap class yang memegang koleksi data tidak boleh memiliki variabel anggota lain selain koleksi itu sendiri. Pendekatan ini mengenkapsulasi array internal, menjamin Type Safety 100%, serta memusatkan logika analitik (seperti pencarian, penyaringan, dan agregasi) langsung pada objek koleksi."
    )

    builder.add_heading_2("10.2 Mengintegrasikan Antarmuka Standar PHP (SPL)")
    builder.add_paragraph(
        "Agar objek koleksi dapat diperlakukan senyaman array asli PHP (dihitung dengan count(), diulang dengan foreach, dan diserialisasi dengan json_encode()), implementasikan antarmuka SPL berikut:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Domain\\Model;\n\n"
        "use Countable;\n"
        "use IteratorAggregate;\n"
        "use JsonSerializable;\n"
        "use ArrayIterator;\n"
        "use Traversable;\n\n"
        "class Mahasiswa {\n"
        "    public function __construct(\n"
        "        public readonly string $nim,\n"
        "        public string $nama,\n"
        "        public float $ipk\n"
        "    ) {}\n"
        "}\n\n"
        "class MahasiswaCollection implements Countable, IteratorAggregate, JsonSerializable\n"
        "{\n"
        "    /** @var Mahasiswa[] */\n"
        "    private array $items = [];\n\n"
        "    public function __construct(Mahasiswa ...$mahasiswa) {\n"
        "        foreach ($mahasiswa as $mhs) {\n"
        "            $this->tambah($mhs);\n"
        "        }\n"
        "    }\n\n"
        "    public function tambah(Mahasiswa $mhs): self {\n"
        "        $this->items[$mhs->nim] = $mhs;\n"
        "        return $this;\n"
        "    }\n\n"
        "    public function count(): int {\n"
        "        return count($this->items);\n"
        "    }\n\n"
        "    public function getIterator(): Traversable {\n"
        "        return new ArrayIterator($this->items);\n"
        "    }\n\n"
        "    public function jsonSerialize(): array {\n"
        "        return array_values($this->items);\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("10.3 Manipulasi Fungsional: `map`, `filter`, `reduce`")
    builder.add_paragraph(
        "Pendekatan pemrograman fungsional memastikan data asli tidak mengalami mutasi liar (Side-Effect Free) dengan selalu mengembalikan instance MahasiswaCollection baru saat dilakukan transformasi data:"
    )

    builder.add_code(
        "<?php\n"
        "// A. FILTERING: Mengembalikan MahasiswaCollection baru (Cum Laude)\n"
        "public function filterCumLaude(): self {\n"
        "    $hasil = array_filter($this->items, fn(Mahasiswa $m) => $m->ipk >= 3.50);\n"
        "    $koleksiBaru = new self();\n"
        "    $koleksiBaru->items = $hasil;\n"
        "    return $koleksiBaru;\n"
        "}\n\n"
        "// B. REDUCING: Mengkalkulasi rata-rata IPK seluruh angkatan\n"
        "public function hitungRataRataIpk(): float {\n"
        "    if (empty($this->items)) return 0.0;\n"
        "    $total = array_reduce($this->items, fn(float $sum, Mahasiswa $m) => $sum + $m->ipk, 0.0);\n"
        "    return $total / count($this->items);\n"
        "}\n\n"
        "// C. SORTING: Mengurutkan IPK tertinggi ke terendah (Descending)\n"
        "public function urutkanBerdasarkanIpkTertinggi(): self {\n"
        "    $salinan = $this->items;\n"
        "    uasort($salinan, fn(Mahasiswa $a, Mahasiswa $b) => $b->ipk <=> $a->ipk);\n"
        "    $koleksiBaru = new self();\n"
        "    $koleksiBaru->items = $salinan;\n"
        "    return $koleksiBaru;\n"
        "}\n"
    )

    builder.add_heading_2("10.4 Studi Kasus Analisis Rekapitulasi Yudisium")
    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "$angkatan2024 = new MahasiswaCollection(\n"
        "    new Mahasiswa(\"240101\", \"Cut Meurah Intan\", 3.92),\n"
        "    new Mahasiswa(\"240102\", \"Teuku Rayhan\", 3.45),\n"
        "    new Mahasiswa(\"240103\", \"Siti Nurhaliza\", 3.88),\n"
        "    new Mahasiswa(\"240104\", \"Zulfa Safira\", 3.95)\n"
        ");\n\n"
        "echo \"Total Mahasiswa : \" . count($angkatan2024) . \" orang\\n\";\n"
        "echo sprintf(\"Rata-rata IPK   : %.2f\\n\", $angkatan2024->hitungRataRataIpk());\n\n"
        "// Fluent Chaining: Filter Cumlaude -> Urutkan Tertinggi\n"
        "$cumlaudeSorted = $angkatan2024->filterCumLaude()->urutkanBerdasarkanIpkTertinggi();\n"
        "foreach ($cumlaudeSorted as $mhs) {\n"
        "    echo sprintf(\"🏆 [%s] %-20s : IPK %.2f\\n\", $mhs->nim, $mhs->nama, $mhs->ipk);\n"
        "}\n"
    )

    builder.add_tip(
        "Tips Koleksi: Manfaatkan Immutable Methods pada Koleksi Objek",
        "Rancang method manipulasi koleksi (seperti filter(), sort(), slice()) agar selalu mengembalikan instance baru ('return new self(...)') alih-alih mengubah array internal secara in-place. Pola ini menjaga integritas data asli dari modifikasi yang tidak diinginkan di bagian modul lain."
    )

    builder.add_summary_and_questions([
        "First-Class Collections membungkus array di dalam class untuk menjamin type-safety 100%.",
        "Antarmuka SPL (Countable, IteratorAggregate, JsonSerializable) membuat koleksi objek dapat diproses layaknya array native.",
        "Manipulasi fungsional (map, filter, reduce) memproses data tanpa efek samping.",
        "PHP 8.1+ mendukung Array Unpacking dengan String Keys dan fungsi validasi array_is_list()."
    ], [
        "Jelaskan mengapa pendekatan First-Class Collection lebih unggul dibandingkan mengelola array mentah secara langsung!",
        "Bagaimana cara kerja antarmuka `IteratorAggregate` dalam memungkinkan sebuah objek diulang menggunakan perulangan `foreach`?",
        "Rancanglah method `ambilStokKritis(int $ambang)` pada `ProdukCollection` menggunakan fungsi `array_filter` dan Arrow Function!",
        "Apa fungsi dari fungsi bawaan `array_is_list()` yang diperkenalkan pada PHP 8.1?"
    ])

    # =========================================================================
    # BAB 11: MANAJEMEN BERKAS DAN ALIRAN DATA (FILE HANDLING & I/O STREAM)
    # =========================================================================
    builder.add_bab_title(11, "Manajemen Berkas dan Aliran Data (File Handling & I/O Stream)")
    builder.add_learning_objectives("Sub-CPMK 4", [
        "Memahami filosofi persistensi data berbasis berkas (File-based Persistence) serta perbedaan antara Stream-based I/O dan Buffered I/O.",
        "Menguasai teknik pencegahan Race Condition dan korupsi data menggunakan File Locking (flock) pada operasi konkuren.",
        "Menerapkan pustaka berorientasi objek modern SplFileObject dan SplFileInfo untuk membaca dan menulis berkas secara memory-efficient.",
        "Mengolah berkas berformat CSV dan JSON menggunakan standar keamanan modern (JSON_THROW_ON_ERROR, json_validate() di PHP 8.3+).",
        "Membangun pola desain File-Based Repository Pattern yang memisahkan logika persistensi berkas dari lapisan domain bisnis.",
        "Merancang modul Audit Logger persisten dengan rotasi berkas otomatis."
    ])

    builder.add_heading_2("11.1 Filosofi Persistensi Berkas & Pencegahan Race Condition")
    builder.add_paragraph(
        "Persistensi berkas merupakan fondasi paling mendasar dari penyimpanan data sebelum diperkenalkannya sistem basis data relasional. Pada aplikasi web yang melayani ratusan permintaan pengguna secara bersamaan (concurrent requests), operasi penulisan file rentan mengalami Race Condition yang menyebabkan isi berkas terpotong atau korup."
    )
    builder.add_paragraph(
        "PHP menyediakan mekanisme penguncian berkas menggunakan fungsi flock(): LOCK_SH (Shared Lock untuk operasi baca bersama), LOCK_EX (Exclusive Lock untuk operasi tulis eksklusif), dan LOCK_UN (Unlock untuk melepaskan kunci setelah operasi selesai)."
    )

    builder.add_heading_2("11.2 Pemrosesan Berkas Berorientasi Objek dengan `SplFileObject`")
    builder.add_paragraph(
        "Pustaka Standar PHP (SPL) menyediakan class SplFileObject yang membungkus fungsi I/O tradisional ke dalam antarmuka berorientasi objek yang hemat memori (Memory-Efficient Stream), sangat ideal untuk memproses berkas CSV besar baris demi baris:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Infrastructure\\IO;\n\n"
        "use SplFileObject;\n"
        "use RuntimeException;\n\n"
        "class CsvMahasiswaReader\n"
        "{\n"
        "    public function bacaData(string $pathFile): array {\n"
        "        if (!file_exists($pathFile)) {\n"
        "            throw new RuntimeException(\"Berkas CSV tidak ditemukan: {$pathFile}\");\n"
        "        }\n"
        "        $file = new SplFileObject($pathFile, 'r');\n"
        "        $file->setFlags(SplFileObject::READ_CSV | SplFileObject::SKIP_EMPTY | SplFileObject::DROP_NEW_LINE);\n\n"
        "        $daftarMahasiswa = [];\n"
        "        $isHeader = true;\n\n"
        "        foreach ($file as $baris) {\n"
        "            if ($isHeader) { $isHeader = false; continue; }\n"
        "            if (is_array($baris) && count($baris) >= 3) {\n"
        "                $daftarMahasiswa[] = [\n"
        "                    'nim'  => trim($baris[0]),\n"
        "                    'nama' => trim($baris[1]),\n"
        "                    'ipk'  => (float) trim($baris[2])\n"
        "                ];\n"
        "            }\n"
        "        }\n"
        "        return $daftarMahasiswa;\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("11.3 Manipulasi JSON Modern & Validasi Native di PHP 8.3+")
    builder.add_paragraph(
        "Pengolahan JSON modern diwajibkan menyertakan flag JSON_THROW_ON_ERROR untuk menangkap kesalahan parsing secara elegan. PHP 8.3 memperkenalkan fungsi native json_validate() untuk memverifikasi keabsahan sintaks JSON dengan performa tinggi tanpa alokasi memori berlebih:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "$dataAkademik = [\n"
        "    'fakultas' => 'Fakultas Sains dan Teknologi',\n"
        "    'prodi'    => 'Informatika',\n"
        "    'mahasiswa' => [\n"
        "        ['nim' => '240101', 'nama' => 'Cut Meurah Intan', 'ipk' => 3.92],\n"
        "        ['nim' => '240102', 'nama' => 'Teuku Rayhan', 'ipk' => 3.85]\n"
        "    ]\n"
        "];\n\n"
        "// 1. Serialization dengan Error Throwing & Pretty Print\n"
        "$jsonString = json_encode($dataAkademik, JSON_PRETTY_PRINT | JSON_UNESCAPED_UNICODE | JSON_THROW_ON_ERROR);\n\n"
        "// 2. Penulisan File Aman dengan Exclusive Lock (LOCK_EX)\n"
        "file_put_contents('data_akademik.json', $jsonString, LOCK_EX);\n\n"
        "// 3. PHP 8.3+: json_validate()\n"
        "if (function_exists('json_validate') && json_validate($jsonString)) {\n"
        "    echo \"Format JSON valid dan terverifikasi secara native.\\n\";\n"
        "}\n"
    )

    builder.add_heading_2("11.4 Pola Desain: File-Based Repository Pattern")
    builder.add_paragraph(
        "Pola Repository memisahkan lapisan proses bisnis dari detail media penyimpanan fisik. Berikut adalah implementasi JsonMahasiswaRepository lengkap dengan proteksi konkurensi:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Infrastructure\\Persistence;\n\n"
        "use App\\Domain\\Model\\Mahasiswa;\n\n"
        "class JsonMahasiswaRepository\n"
        "{\n"
        "    public function __construct(private string $filePath) {\n"
        "        if (!file_exists($this->filePath)) {\n"
        "            file_put_contents($this->filePath, json_encode([], JSON_PRETTY_PRINT), LOCK_EX);\n"
        "        }\n"
        "    }\n\n"
        "    public function ambilSemua(): array {\n"
        "        $json = file_get_contents($this->filePath);\n"
        "        $data = json_decode($json, true, 512, JSON_THROW_ON_ERROR);\n"
        "        $koleksi = [];\n"
        "        foreach ($data as $item) {\n"
        "            $koleksi[$item['nim']] = new Mahasiswa($item['nim'], $item['nama'], (float)$item['ipk']);\n"
        "        }\n"
        "        return $koleksi;\n"
        "    }\n\n"
        "    public function simpan(Mahasiswa $mhs): void {\n"
        "        $semua = $this->ambilSemua();\n"
        "        $semua[$mhs->nim] = $mhs;\n"
        "        $arrayData = array_map(fn(Mahasiswa $m) => [\n"
        "            'nim'  => $m->nim,\n"
        "            'nama' => $m->nama,\n"
        "            'ipk'  => $m->ipk\n"
        "        ], array_values($semua));\n\n"
        "        file_put_contents(\n"
        "            $this->filePath,\n"
        "            json_encode($arrayData, JSON_PRETTY_PRINT | JSON_THROW_ON_ERROR),\n"
        "            LOCK_EX // Menjamin keamanan penulisan konkuren\n"
        "        );\n"
        "    }\n"
        "}\n"
    )

    builder.add_tip(
        "Tips File I/O: Gunakan Stream I/O untuk Pemrosesan Berkas Skala Besar",
        "Untuk berkas berukuran besar (ratusan megabyte hingga gigabyte), hindari penggunaan file_get_contents() karena akan memuat seluruh isi berkas ke RAM sekaligus. Selalu gunakan SplFileObject atau fgetcsv() untuk memproses data secara bertahap (streaming)."
    )

    builder.add_summary_and_questions([
        "Persistensi berkas menyediakan media penyimpanan data mandiri tanpa ketergantungan DBMS luar.",
        "File Locking (flock) dengan flag LOCK_EX wajib diterapkan untuk mencegah Race Condition saat penulisan.",
        "SplFileObject menyediakan antarmuka OOP yang hemat memori untuk pemrosesan berkas teks dan CSV.",
        "Flag JSON_THROW_ON_ERROR dan fungsi json_validate() (PHP 8.3+) meningkatkan keandalan parsing JSON.",
        "Pola File-Based Repository memisahkan logika domain aplikasi dari sistem penyimpanan berkas."
    ], [
        "Jelaskan bahaya Race Condition pada penulisan berkas bersamaan dan bagaimana flag `LOCK_EX` mengatasinya!",
        "Apa keunggulan menggunakan `SplFileObject` dibandingkan kombinasi fungsi prosedural `fopen` dan `fgets`?",
        "Rancanglah class `DailyAuditLogger` yang mencatat log aktivitas dengan rotasi nama file harian (`audit_YYYY-MM-DD.log`)!",
        "Bagaimana peran fungsi `json_validate()` pada PHP 8.3+ dalam mengoptimalkan performa API?"
    ])

    # =========================================================================
    # BAB 12: PRINSIP DESAIN PERANGKAT LUNAK SOLID PADA PHP MODERN
    # =========================================================================
    builder.add_bab_title(12, "Prinsip Desain Perangkat Lunak SOLID pada PHP Modern")
    builder.add_learning_objectives("Sub-CPMK 5", [
        "Memahami latar belakang lahirnya Prinsip Desain SOLID (Robert C. Martin / Uncle Bob) dalam mengeliminasi 4 gejala pembusukan arsitektur (Rigidity, Fragility, Immobility, Viscosity).",
        "Menerapkan Single Responsibility Principle (SRP) untuk mencegah antipattern God Object dan memusatkan satu alasan perubahan per class.",
        "Menerapkan Open/Closed Principle (OCP) dengan pola Strategy Pattern sehingga sistem mudah diperluas tanpa mengubah kode yang sudah teruji.",
        "Menganalisis dan menegakkan Liskov Substitution Principle (LSP) termasuk aturan kovariansi (Covariance) dan kontravariansi (Contravariance) di PHP 8+.",
        "Menerapkan Interface Segregation Principle (ISP) dengan memecah Fat Interface menjadi Role-based Interfaces yang ramping.",
        "Menerapkan Dependency Inversion Principle (DIP) dan teknik Constructor Dependency Injection dengan fitur Constructor Property Promotion di PHP 8.0+."
    ])

    builder.add_heading_2("12.1 Fondasi Teoretis SOLID: Menghindari Pembusukan Perangkat Lunak")
    builder.add_paragraph(
        "Prinsip SOLID dicetuskan oleh Robert C. Martin (Uncle Bob) untuk memitigasi 4 gejala klasik pembusukan perangkat lunak (Software Rot): Rigidity (kekakuan sistem terhadap perubahan), Fragility (kerapuhan sistem di mana perubahan di satu modul merusak modul lain), Immobility (ketakmampuan kode untuk digunakan kembali akibat tight coupling), dan Viscosity (sulitnya mempertahankan desain arsitektur yang benar)."
    )

    builder.add_heading_2("12.2 Bedah Komprehensif 5 Prinsip SOLID di PHP 8+")
    builder.add_paragraph(
        "1. Single Responsibility Principle (SRP): Setiap class hanya boleh memiliki satu alasan untuk berubah (satu tanggung jawab bisnis). Menghindari antipattern God Object yang mencampuradukkan entitas, database, dan presentasi."
    )
    builder.add_paragraph(
        "2. Open/Closed Principle (OCP): Entitas perangkat lunak harus terbuka untuk perluasan (open for extension), namun tertutup untuk modifikasi (closed for modification). Menggunakan abstraksi interface (seperti Strategy Pattern) agar fitur baru dapat disisipkan tanpa mengedit kode lama."
    )
    builder.add_paragraph(
        "3. Liskov Substitution Principle (LSP): Objek turunan (subclass) harus dapat menggantikan objek induknya (superclass) tanpa merusak kebenaran program. Subclass dilarang memperketat prekondisi atau melempar eksepsi pada method yang sah di parent."
    )
    builder.add_paragraph(
        "4. Interface Segregation Principle (ISP): Klien tidak boleh dipaksa bergantung pada antarmuka yang tidak digunakannya. Pecah antarmuka gemuk (Fat Interface) menjadi kumpulan Role Interfaces yang ramping."
    )
    builder.add_paragraph(
        "5. Dependency Inversion Principle (DIP): Modul tingkat tinggi tidak boleh bergantung pada modul tingkat rendah; keduanya harus bergantung pada abstraksi interface."
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Ecommerce;\n\n"
        "// 1. OCP: Interface Strategi Diskon Terbuka\n"
        "interface DiskonStrategyInterface {\n"
        "    public function hitungDiskon(float $subtotal): float;\n"
        "}\n\n"
        "class DiskonMemberReguler implements DiskonStrategyInterface {\n"
        "    public function hitungDiskon(float $subtotal): float { return $subtotal * 0.05; }\n"
        "}\n\n"
        "class DiskonFlashSale implements DiskonStrategyInterface {\n"
        "    public function hitungDiskon(float $subtotal): float { return $subtotal * 0.20; }\n"
        "}\n\n"
        "// 2. DIP: Interface Payment Gateway Abstraksi\n"
        "interface PaymentProcessorInterface {\n"
        "    public function bayar(float $nominal): bool;\n"
        "}\n\n"
        "class QrisPaymentProcessor implements PaymentProcessorInterface {\n"
        "    public function bayar(float $nominal): bool {\n"
        "        echo \"📱 Pembayaran QRIS Rp \" . number_format($nominal, 0, ',', '.') . \" BERHASIL.\\n\";\n"
        "        return true;\n"
        "    }\n"
        "}\n\n"
        "// 3. ISP: Interface Notifikasi Khusus\n"
        "interface NotifierInterface {\n"
        "    public function kirimNotifikasi(string $tujuan, string $pesan): bool;\n"
        "}\n\n"
        "class WhatsAppNotifier implements NotifierInterface {\n"
        "    public function kirimNotifikasi(string $tujuan, string $pesan): bool {\n"
        "        echo \"📲 [WHATSAPP] Mengirim pesan ke {$tujuan}: {$pesan}\\n\";\n"
        "        return true;\n"
        "    }\n"
        "}\n\n"
        "// 4. SRP & DIP: Pipeline Checkout Tingkat Tinggi\n"
        "class TransaksiPipelineService\n"
        "{\n"
        "    // PHP 8.0: Constructor Property Promotion dengan Type-Hint Abstraksi\n"
        "    public function __construct(\n"
        "        private DiskonStrategyInterface $diskonStrategy,\n"
        "        private PaymentProcessorInterface $paymentProcessor,\n"
        "        private NotifierInterface $notifier\n"
        "    ) {}\n\n"
        "    public function selesaikanTransaksi(string $customer, float $subtotal, string $kontak): void {\n"
        "        $diskon = $this->diskonStrategy->hitungDiskon($subtotal);\n"
        "        $totalBayar = max(0.0, $subtotal - $diskon);\n\n"
        "        echo \"========================================================\\n\";\n"
        "        echo \"PROSES TRANSAKSI CHECKOUT E-COMMERCE (SOLID PIPELINE)\\n\";\n"
        "        echo \"Pelanggan   : {$customer}\\n\";\n"
        "        echo \"Subtotal    : Rp \" . number_format($subtotal, 0, ',', '.') . \"\\n\";\n"
        "        echo \"Diskon      : Rp \" . number_format($diskon, 0, ',', '.') . \"\\n\";\n"
        "        echo \"Total Bayar : Rp \" . number_format($totalBayar, 0, ',', '.') . \"\\n\";\n"
        "        echo \"--------------------------------------------------------\\n\";\n\n"
        "        $this->paymentProcessor->bayar($totalBayar);\n"
        "        $this->notifier->kirimNotifikasi($kontak, \"Pesanan Anda Rp \" . number_format($totalBayar) . \" lunas.\");\n"
        "        echo \"========================================================\\n\";\n"
        "    }\n"
        "}\n"
    )

    builder.add_tip(
        "Tips SOLID: Utamakan Constructor Dependency Injection",
        "Hindari pembuatan instansi langsung ('new ClassKonkret()') di dalam service layer Anda. Selalu minta objek dependensi melalui parameter constructor dengan type-hint antarmuka (interface). Langkah ini menjamin kepatuhan mutlak terhadap DIP dan membuat kode Anda 100% siap untuk automated unit testing dengan Mockery atau PHPUnit."
    )

    builder.add_summary_and_questions([
        "Prinsip SOLID mengeliminasi kerapuhan arsitektural (Software Rot) pada sistem perangkat lunak.",
        "SRP memusatkan satu tanggung jawab per class untuk mencegah God Objects.",
        "OCP memungkinkan perluasan fitur baru melalui Strategy Pattern tanpa modifikasi kode produksi.",
        "LSP menjamin bahwa seluruh subclass dapat disubstitusikan secara aman menggantikan superclass.",
        "ISP memecah interface gemuk menjadi role interfaces yang ramping dan tepat sasaran.",
        "DIP dan Constructor Injection membalikkan arah ketergantungan modul ke abstraksi interface."
    ], [
        "Jelaskan 4 gejala utama Software Rot (Rigidity, Fragility, Immobility, Viscosity) dan bagaimana prinsip SOLID menyelesaikannya!",
        "Bagaimana cara mengidentifikasi class yang melanggar Single Responsibility Principle (SRP)?",
        "Berikan contoh kasus nyata pelanggaran Liskov Substitution Principle (LSP) pada hierarki bentuk geometri atau entitas perbankan!",
        "Uraikan perbedaan konseptual antara Dependency Inversion Principle (prinsip), Dependency Injection (teknik), dan IoC Container (alat)!"
    ])

    # =========================================================================
    # BAB 13: ARSITEKTUR APLIKASI BERORIENTASI OBJEK (MODEL-SERVICE-REPOSITORY)
    # =========================================================================
    builder.add_bab_title(13, "Arsitektur Aplikasi Berorientasi Objek (Model-Service-Repository)")
    builder.add_learning_objectives("Sub-CPMK 6", [
        "Memahami filosofi Arsitektur Berlapis (Layered Architecture) dan Aturan Ketergantungan (Dependency Rule) dari Clean Architecture (Robert C. Martin, 2017).",
        "Membedakan peran dan tanggung jawab 4 lapisan konsentris: Entity, Use Case, Interface Adapter, dan Framework/Driver.",
        "Menguasai pola Model/Entity sebagai representasi data domain murni yang bebas dari dependensi infrastruktur.",
        "Menerapkan pola Repository sebagai abstraksi akses data yang memisahkan logika bisnis dari detail penyimpanan (Database Agnostic).",
        "Merancang Service Layer sebagai orkestrator aturan bisnis yang mengkoordinasikan Repository, Validator, dan Notifier.",
        "Membangun aplikasi CLI lengkap berpola Model-Service-Repository dengan persistensi JSON dan struktur folder terstandar PSR-4."
    ])

    builder.add_heading_2("13.1 Filosofi Clean Architecture & Arsitektur Berlapis")
    builder.add_paragraph(
        "Robert C. Martin dalam buku 'Clean Architecture: A Craftsman's Guide to Software Structure and Design' (2017) memperkenalkan arsitektur konsentris 4 lapisan dengan aturan ketergantungan (Dependency Rule) yang ketat: Dependensi kode sumber hanya boleh menunjuk ke arah dalam (dari lapisan luar ke lapisan dalam). Lapisan Entity dan Use Case tidak boleh mengetahui apa pun tentang lapisan Framework/Driver (database, web framework, UI)."
    )
    builder.add_paragraph(
        "Tanpa pemisahan lapisan yang tegas, sebuah aplikasi monolitik akan mengalami masalah serius: logika bisnis tertanam di dalam controller (tidak dapat diuji tanpa server web), SQL query tersebar di berbagai class (perubahan skema memaksa modifikasi puluhan berkas), dan dependensi framework merembes ke inti domain (migrasi framework memaksa penulisan ulang total)."
    )

    builder.add_heading_2("13.2 Lapisan 1: Entity / Domain Model")
    builder.add_paragraph(
        "Entity merepresentasikan objek bisnis inti yang bebas dari dependensi infrastruktur (tidak mengimpor database, file system, atau framework). Menggunakan Backed Enum PHP 8.1+ untuk mendefinisikan state yang type-safe:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Domain\\Model;\n\n"
        "enum StatusPinjam: string {\n"
        "    case Tersedia = 'tersedia';\n"
        "    case Dipinjam = 'dipinjam';\n"
        "    case Rusak    = 'rusak';\n"
        "}\n\n"
        "class Buku\n"
        "{\n"
        "    private StatusPinjam $status = StatusPinjam::Tersedia;\n\n"
        "    public function __construct(\n"
        "        public readonly string $isbn,\n"
        "        public readonly string $judul,\n"
        "        public readonly string $pengarang,\n"
        "        public readonly int $tahunTerbit\n"
        "    ) {}\n\n"
        "    public function getStatus(): StatusPinjam { return $this->status; }\n\n"
        "    public function pinjam(): void {\n"
        "        if ($this->status !== StatusPinjam::Tersedia) {\n"
        "            throw new \\App\\Domain\\Exception\\BukuTidakTersediaException(\n"
        "                $this->isbn, $this->status\n"
        "            );\n"
        "        }\n"
        "        $this->status = StatusPinjam::Dipinjam;\n"
        "    }\n\n"
        "    public function kembalikan(): void {\n"
        "        $this->status = StatusPinjam::Tersedia;\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("13.3 Repository Interface (Kontrak Abstraksi di Lapisan Domain)")
    builder.add_paragraph(
        "Repository Interface dideklarasikan di lapisan domain (lapisan dalam), bukan di lapisan infrastruktur. Hal ini memastikan Service Layer (Use Case) hanya mengenal kontrak abstraksi dan tidak bergantung pada implementasi konkret penyimpanan data:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Domain\\Repository;\n\n"
        "use App\\Domain\\Model\\Buku;\n\n"
        "interface BukuRepositoryInterface\n"
        "{\n"
        "    public function simpan(Buku $buku): void;\n"
        "    public function cariBerdasarkanIsbn(string $isbn): ?Buku;\n"
        "    /** @return Buku[] */\n"
        "    public function ambilSemua(): array;\n"
        "    public function hapus(string $isbn): bool;\n"
        "}\n"
    )

    builder.add_heading_2("13.4 Service Layer (Use Case / Business Orchestrator)")
    builder.add_paragraph(
        "Service Layer mengkoordinasikan Entity dan Repository tanpa mengetahui bagaimana data disimpan (JSON, database, atau API eksternal). Seluruh aturan bisnis dieksekusi pada lapisan ini:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Application\\Service;\n\n"
        "use App\\Domain\\Model\\Buku;\n"
        "use App\\Domain\\Repository\\BukuRepositoryInterface;\n\n"
        "class PerpustakaanService\n"
        "{\n"
        "    public function __construct(\n"
        "        private BukuRepositoryInterface $bukuRepository\n"
        "    ) {}\n\n"
        "    public function daftarkanBukuBaru(string $isbn, string $judul, string $pengarang, int $tahun): Buku {\n"
        "        if ($this->bukuRepository->cariBerdasarkanIsbn($isbn) !== null) {\n"
        "            throw new \\InvalidArgumentException(\"Buku [{$isbn}] sudah terdaftar!\");\n"
        "        }\n"
        "        $buku = new Buku($isbn, $judul, $pengarang, $tahun);\n"
        "        $this->bukuRepository->simpan($buku);\n"
        "        return $buku;\n"
        "    }\n\n"
        "    public function pinjamBuku(string $isbn): Buku {\n"
        "        $buku = $this->bukuRepository->cariBerdasarkanIsbn($isbn)\n"
        "            ?? throw new \\RuntimeException(\"Buku [{$isbn}] tidak ditemukan.\");\n"
        "        $buku->pinjam();\n"
        "        $this->bukuRepository->simpan($buku);\n"
        "        return $buku;\n"
        "    }\n\n"
        "    public function kembalikanBuku(string $isbn): Buku {\n"
        "        $buku = $this->bukuRepository->cariBerdasarkanIsbn($isbn)\n"
        "            ?? throw new \\RuntimeException(\"Buku [{$isbn}] tidak ditemukan.\");\n"
        "        $buku->kembalikan();\n"
        "        $this->bukuRepository->simpan($buku);\n"
        "        return $buku;\n"
        "    }\n"
        "}\n"
    )

    builder.add_heading_2("13.5 Implementasi Repository Konkret (Lapisan Infrastructure)")
    builder.add_paragraph(
        "Implementasi konkret berada di lapisan luar (Interface Adapter) dan bertanggung jawab terhadap detail penyimpanan. Jika suatu hari Anda mengganti penyimpanan dari file JSON ke database MySQL (PDO), Anda hanya perlu membuat class PdoBukuRepository implements BukuRepositoryInterface tanpa mengubah satu baris pun di folder Domain/ atau Application/:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Infrastructure\\Persistence;\n\n"
        "use App\\Domain\\Model\\Buku;\n"
        "use App\\Domain\\Model\\StatusPinjam;\n"
        "use App\\Domain\\Repository\\BukuRepositoryInterface;\n\n"
        "class JsonBukuRepository implements BukuRepositoryInterface\n"
        "{\n"
        "    public function __construct(private string $filePath) {\n"
        "        if (!file_exists($this->filePath)) {\n"
        "            file_put_contents($this->filePath, json_encode([], JSON_PRETTY_PRINT), LOCK_EX);\n"
        "        }\n"
        "    }\n\n"
        "    public function ambilSemua(): array {\n"
        "        $json = file_get_contents($this->filePath);\n"
        "        $data = json_decode($json, true, 512, JSON_THROW_ON_ERROR);\n"
        "        return array_map(function (array $item): Buku {\n"
        "            $buku = new Buku($item['isbn'], $item['judul'], $item['pengarang'], $item['tahun_terbit']);\n"
        "            if ($item['status'] === StatusPinjam::Dipinjam->value) { $buku->pinjam(); }\n"
        "            return $buku;\n"
        "        }, $data);\n"
        "    }\n\n"
        "    public function simpan(Buku $buku): void {\n"
        "        $semua = $this->ambilSemua();\n"
        "        $ditemukan = false;\n"
        "        foreach ($semua as $i => $existing) {\n"
        "            if ($existing->isbn === $buku->isbn) { $semua[$i] = $buku; $ditemukan = true; break; }\n"
        "        }\n"
        "        if (!$ditemukan) { $semua[] = $buku; }\n"
        "        $this->persistKeFile($semua);\n"
        "    }\n\n"
        "    public function cariBerdasarkanIsbn(string $isbn): ?Buku {\n"
        "        foreach ($this->ambilSemua() as $buku) {\n"
        "            if ($buku->isbn === $isbn) return $buku;\n"
        "        }\n"
        "        return null;\n"
        "    }\n\n"
        "    public function hapus(string $isbn): bool {\n"
        "        $semua = $this->ambilSemua();\n"
        "        $filtered = array_filter($semua, fn(Buku $b) => $b->isbn !== $isbn);\n"
        "        if (count($filtered) === count($semua)) return false;\n"
        "        $this->persistKeFile(array_values($filtered));\n"
        "        return true;\n"
        "    }\n\n"
        "    private function persistKeFile(array $koleksi): void {\n"
        "        $data = array_map(fn(Buku $b) => [\n"
        "            'isbn' => $b->isbn, 'judul' => $b->judul,\n"
        "            'pengarang' => $b->pengarang, 'tahun_terbit' => $b->tahunTerbit,\n"
        "            'status' => $b->getStatus()->value\n"
        "        ], $koleksi);\n"
        "        file_put_contents($this->filePath, json_encode($data, JSON_PRETTY_PRINT | JSON_THROW_ON_ERROR), LOCK_EX);\n"
        "    }\n"
        "}\n"
    )

    builder.add_tip(
        "Tips Arsitektur Bersih: Pisahkan Domain dari Infrastructure",
        "Selalu deklarasikan Repository Interface di folder Domain (lapisan dalam) dan implementasi konkretnya di folder Infrastructure (lapisan luar). Hal ini memastikan inti bisnis aplikasi Anda portabel, tidak bergantung pada framework atau teknologi penyimpanan tertentu, serta siap 100% untuk automated unit testing menggunakan mock repository."
    )

    builder.add_summary_and_questions([
        "Clean Architecture memisahkan aplikasi menjadi 4 lapisan konsentris dengan Dependency Rule mengarah ke dalam.",
        "Entity/Domain Model merepresentasikan aturan bisnis murni yang bebas dari dependensi infrastruktur.",
        "Repository Interface dideklarasikan di lapisan Domain sebagai kontrak abstraksi akses data.",
        "Service Layer mengkoordinasikan Entity dan Repository tanpa mengetahui detail penyimpanan.",
        "Implementasi konkret (JsonRepository, PdoRepository) berada di lapisan Infrastructure yang dapat diganti tanpa memodifikasi inti bisnis."
    ], [
        "Jelaskan Aturan Ketergantungan (Dependency Rule) dalam Clean Architecture dan mengapa arah dependensi harus ke dalam!",
        "Mengapa Repository Interface harus dideklarasikan di lapisan Domain dan bukan di lapisan Infrastructure?",
        "Rancang implementasi kedua `PdoBukuRepository` menggunakan database SQLite tanpa mengubah `PerpustakaanService`!",
        "Gambarkan diagram aliran ketergantungan (dependency flow) proyek perpustakaan ini dan pastikan tidak ada panah dari lapisan dalam ke lapisan luar!"
    ])

    # =========================================================================
    # BAB 14: STUDI KASUS MINI PROJECT: SISTEM POINT OF SALE (POS) TERPADU
    # =========================================================================
    builder.add_bab_title(14, "Studi Kasus Mini Project: Sistem Point of Sale (POS) Terpadu")
    builder.add_learning_objectives("Sub-CPMK 6", [
        "Mengintegrasikan seluruh 14 bab materi: 4 Pilar OOP, Backed Enum, First-Class Collections, Rich Exceptions, File I/O Stream dengan Locking, dan Prinsip SOLID.",
        "Menerapkan arsitektur Model-Service-Repository (Clean Architecture) dengan pemisahan lapisan yang tegas dan kepatuhan mutlak pada Dependency Rule.",
        "Mengembangkan modul transaksi kasir retail berbasis CLI interaktif dengan kalkulasi multi-diskon (Strategy Pattern) dan multi-metode pembayaran.",
        "Menghasilkan kode program yang memenuhi standar PSR-4 Autoloading, PSR-12 Coding Style, serta aman dari Race Condition pada persistensi data.",
        "Mempresentasikan proyek perangkat lunak terintegrasi dengan dokumentasi teknis dan diagram UML yang lengkap."
    ])

    builder.add_heading_2("14.1 Deskripsi Kasus & Kebutuhan Sistem POS Enterprise")
    builder.add_paragraph(
        "Sistem Point of Sale (POS) Kasir Retail Terpadu merupakan proyek capstone yang mensintesis seluruh kompetensi pemrograman berorientasi objek yang telah dipelajari. Sistem ini bertanggung jawab mengelola master produk dengan kategori bertipe Backed Enum, memproses keranjang belanja pelanggan menggunakan First-Class Collection, menerapkan skema diskon fleksibel berbasis Strategy Pattern, memvalidasi pembayaran tunai dan non-tunai, serta mendokumentasikan struk transaksi secara persisten ke dalam berkas JSON yang terproteksi File Locking (LOCK_EX)."
    )

    builder.add_heading_2("14.2 Desain Domain Model & First-Class Collection")
    builder.add_paragraph(
        "Lapisan domain dirancang menggunakan fitur-fitur modern PHP 8+: Backed Enum untuk kategori produk, readonly properties pada entitas, invariant validasi stok, dan First-Class Collection untuk keranjang belanja:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Domain\\Model;\n\n"
        "// 1. Backed Enum Kategori Produk\n"
        "enum KategoriProduk: string {\n"
        "    case Makanan    = 'Makanan';\n"
        "    case Minuman    = 'Minuman';\n"
        "    case Elektronik = 'Elektronik';\n"
        "    case Pendidikan = 'Pendidikan';\n"
        "}\n\n"
        "// 2. Entity ItemProduk dengan Enkapsulasi Invariant\n"
        "class ItemProduk {\n"
        "    public function __construct(\n"
        "        public readonly string $sku,\n"
        "        public readonly string $nama,\n"
        "        public readonly float $harga,\n"
        "        private int $stok,\n"
        "        public readonly KategoriProduk $kategori\n"
        "    ) {\n"
        "        if ($harga <= 0) throw new \\InvalidArgumentException(\"Harga harus positif!\");\n"
        "        if ($stok < 0) throw new \\InvalidArgumentException(\"Stok tidak boleh negatif!\");\n"
        "    }\n"
        "    public function getStok(): int { return $this->stok; }\n"
        "    public function kurangiStok(int $qty): void {\n"
        "        if ($qty > $this->stok) {\n"
        "            throw new \\App\\Domain\\Exception\\StokTidakCukupException($this->sku, $this->stok, $qty);\n"
        "        }\n"
        "        $this->stok -= $qty;\n"
        "    }\n"
        "}\n\n"
        "// 3. First-Class Collection Keranjang Belanja\n"
        "class KeranjangCollection implements \\Countable, \\IteratorAggregate {\n"
        "    /** @var array<string, ItemKeranjang> */\n"
        "    private array $items = [];\n"
        "    public function tambah(ItemProduk $p, int $qty): void {\n"
        "        $qtyTotal = ($this->items[$p->sku]->kuantitas ?? 0) + $qty;\n"
        "        $this->items[$p->sku] = new ItemKeranjang($p, $qtyTotal);\n"
        "    }\n"
        "    public function hitungTotalBruto(): float {\n"
        "        return array_reduce($this->items, fn($sum, $i) => $sum + $i->getSubtotal(), 0.0);\n"
        "    }\n"
        "    public function count(): int { return count($this->items); }\n"
        "    public function getIterator(): \\Traversable { return new \\ArrayIterator($this->items); }\n"
        "    public function isEmpty(): bool { return empty($this->items); }\n"
        "}\n"
    )

    builder.add_heading_2("14.3 Strategi Diskon & Pembayaran Polimorfik (SOLID OCP & DIP)")
    builder.add_paragraph(
        "Kalkulasi diskon dan proses pembayaran diabstraksikan melalui interface sehingga mendukung penambahan varian strategi baru tanpa modifikasi kode yang sudah ada:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Domain\\Diskon;\n\n"
        "interface DiskonStrategyInterface {\n"
        "    public function hitungDiskon(float $subtotal): float;\n"
        "    public function getNama(): string;\n"
        "}\n\n"
        "class DiskonPersentase implements DiskonStrategyInterface {\n"
        "    public function __construct(public readonly float $persen) {}\n"
        "    public function hitungDiskon(float $subtotal): float { return $subtotal * ($this->persen / 100); }\n"
        "    public function getNama(): string { return \"Diskon {$this->persen}%\"; }\n"
        "}\n"
    )

    builder.add_heading_2("14.4 Application Service & Persistensi File-Based")
    builder.add_paragraph(
        "KasirService mengorkestrasi alur transaksi, pemotongan stok otomatis, dan perekaman riwayat ke berkas JSON dengan proteksi LOCK_EX:"
    )

    builder.add_code(
        "<?php\n"
        "declare(strict_types=1);\n\n"
        "namespace App\\Application\\Service;\n\n"
        "use App\\Domain\\Model\\KeranjangCollection;\n"
        "use App\\Domain\\Diskon\\DiskonStrategyInterface;\n"
        "use App\\Domain\\Pembayaran\\PaymentProcessorInterface;\n"
        "use App\\Domain\\Repository\\ProdukRepositoryInterface;\n"
        "use App\\Domain\\Repository\\TransaksiRepositoryInterface;\n\n"
        "class KasirService {\n"
        "    public function __construct(\n"
        "        private ProdukRepositoryInterface $produkRepo,\n"
        "        private TransaksiRepositoryInterface $transaksiRepo\n"
        "    ) {}\n\n"
        "    public function checkout(\n"
        "        KeranjangCollection $keranjang,\n"
        "        DiskonStrategyInterface $diskon,\n"
        "        PaymentProcessorInterface $payment,\n"
        "        float $nominalBayar\n"
        "    ): array {\n"
        "        if ($keranjang->isEmpty()) throw new \\RuntimeException(\"Keranjang belanja kosong!\");\n"
        "        $subtotal = $keranjang->hitungTotalBruto();\n"
        "        $potongan = $diskon->hitungDiskon($subtotal);\n"
        "        $totalAkhir = max(0.0, $subtotal - $potongan);\n\n"
        "        $payment->prosesBayar($totalAkhir, $nominalBayar);\n\n"
        "        foreach ($keranjang as $item) {\n"
        "            $item->produk->kurangiStok($item->kuantitas);\n"
        "            $this->produkRepo->simpan($item->produk);\n"
        "        }\n\n"
        "        $dataTrx = [\n"
        "            'id' => 'TRX-' . date('YmdHis'),\n"
        "            'waktu' => date('Y-m-d H:i:s'),\n"
        "            'subtotal' => $subtotal,\n"
        "            'diskon' => $potongan,\n"
        "            'total' => $totalAkhir,\n"
        "            'metode' => $payment->getMetode()\n"
        "        ];\n"
        "        $this->transaksiRepo->simpan($dataTrx);\n"
        "        return $dataTrx;\n"
        "    }\n"
        "}\n"
    )

    builder.add_tip(
        "Tips Evaluasi Proyek Capstone: Pastikan Kualitas Kode Terstandar PSR-12",
        "Sebelum mengumpulkan proyek capstone, lakukan analisis statis menggunakan PHPStan (level 5+) dan periksa kepatuhan gaya penulisan kode menggunakan PHP_CodeSniffer (PSR-12). Pastikan seluruh unit test lolos 100% dan berkas database JSON terlindungi dengan penguncian LOCK_EX."
    )

    builder.add_summary_and_questions([
        "Proyek POS memadukan seluruh 14 bab materi: pilar OOP, collections, exceptions, I/O, dan SOLID.",
        "Model-Service-Repository menjamin kemandirian proses bisnis dari media penyimpanan.",
        "Strategy Pattern mempermudah ekspansi jenis diskon dan metode pembayaran baru.",
        "Penguncian berkas LOCK_EX wajib diterapkan untuk menjamin integritas data transaksi kasir."
    ], [
        "Jelaskan bagaimana prinsip Dependency Inversion (DIP) diterapkan pada `KasirService` dan keuntungan apa yang diperoleh saat pengujian otomatis!",
        "Bagaimana mekanisme penguncian berkas `LOCK_EX` mencegah kehilangan data (*Lost Update*) saat dua kasir melakukan checkout bersamaan?",
        "Rancanglah strategi diskon baru `DiskonBertingkat` (misal diskon 15% jika belanja di atas Rp 500.000) tanpa mengubah `KasirService`!",
        "Uraikan langkah-langkah implementasi `PdoTransaksiRepository` untuk menggantikan repositori JSON saat proyek beralih ke MySQL!"
    ])

    # Glosarium, Pustaka & Profil
    print("-> Menyusun Glosarium & Bagian Akhir...")
    builder.doc.add_page_break()
    p = builder.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GLOSARIUM ISTILAH OOP & PHP")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    glosarium_data = [
        ["Istilah", "Penjelasan / Definisi Teknis"],
        ["Abstract Class", "Class induk yang tidak dapat diinstansiasi langsung dan memuat definisi method abstrak."],
        ["Autoloading", "Mekanisme pemuatan berkas class PHP secara otomatis pada saat class tersebut dipanggil."],
        ["Backed Enum", "Tipe data enumerasi native PHP 8.1 yang nilainya terikat pada string atau integer skalar."],
        ["Class", "Cetak biru (*blueprint*) yang mendefinisikan atribut dan perilaku dari suatu tipe data objek."],
        ["Constructor Promotion", "Fitur PHP 8.0 yang menggabungkan deklarasi properti dan penugasan pada parameter constructor."],
        ["Dependency Injection", "Teknik memasukkan objek dependensi dari luar melalui constructor/setter."],
        ["Dynamic Dispatch", "Mekanisme runtime PHP dalam menentukan method mana yang dieksekusi berdasarkan tipe objek aktual."],
        ["Encapsulation", "Pilar pembungkusan data dan method serta pembatasan akses langsung dari luar class."],
        ["Inheritance", "Pewarisan atribut dan method dari superclass ke subclass menggunakan kata kunci `extends`."],
        ["Interface", "Kontrak murni tanpa implementasi yang mendikte method apa saja yang wajib disediakan oleh class."],
        ["Named Arguments", "Pemanggilan fungsi/method dengan menyebutkan nama parameternya secara eksplisit (PHP 8.0+)."],
        ["Polymorphism", "Prinsip satu antarmuka yang merespons dengan banyak implementasi perilaku yang berbeda."],
        ["PSR-4", "Standar rekomendasi PHP-FIG mengenai pemetaan otomatis antara namespace dan direktori berkas."],
        ["Readonly Property", "Properti yang nilainya hanya dapat diinisialisasi satu kali dan bersifat immutable (PHP 8.1+)."],
        ["SOLID", "Lima prinsip dasar rekayasa perangkat lunak berorientasi objek untuk arsitektur yang bersih."],
        ["Trait", "Mekanisme horizontal code reuse untuk berbagi kode lintas class di luar hierarki pewarisan tunggal."]
    ]
    builder.add_table(glosarium_data[0], glosarium_data[1:])

    builder.doc.add_page_break()
    p_pustaka = builder.doc.add_paragraph()
    p_pustaka.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_pustaka.add_run("DAFTAR PUSTAKA")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    pustaka_list = [
        "Bay, J. (2008). Object Calisthenics. In *The ThoughtWorks Anthology: Essays on Software Technology and Innovation* (pp. 83–98). Pragmatic Bookshelf.",
        "Beck, K. (2002). *Test-Driven Development: By Example*. Addison-Wesley Professional.",
        "Bloch, J. (2018). *Effective Java* (3rd ed.). Addison-Wesley Professional.",
        "Cardelli, L., & Wegner, P. (1985). On understanding types, data abstraction, and polymorphism. *ACM Computing Surveys (CSUR)*, 17(4), 471–523. https://doi.org/10.1145/6041.6042",
        "Evans, E. (2003). *Domain-Driven Design: Tackling Complexity in the Heart of Software*. Addison-Wesley Professional.",
        "Fowler, M. (2002). *Patterns of Enterprise Application Architecture*. Addison-Wesley Professional.",
        "Fowler, M. (2018). *Refactoring: Improving the Design of Existing Code* (2nd ed.). Addison-Wesley Professional.",
        "Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). *Design Patterns: Elements of Reusable Object-Oriented Software*. Addison-Wesley.",
        "Liskov, B. (1987). Data abstraction and hierarchy. *ACM SIGPLAN Notices*, 23(5), 17–34. https://doi.org/10.1145/62139.62141",
        "Lockhart, J. (2015). *Modern PHP: New Features and Good Practices*. O'Reilly Media.",
        "MacArthur, K. (2022). *PHP 8 Quick Scripting Reference* (3rd ed.). Apress. https://doi.org/10.1007/978-1-4842-8111-6",
        "Martin, R. C. (2002). *Agile Software Development, Principles, Patterns, and Practices*. Prentice Hall.",
        "Martin, R. C. (2008). *Clean Code: A Handbook of Agile Software Craftsmanship*. Prentice Hall.",
        "Martin, R. C. (2017). *Clean Architecture: A Craftsman's Guide to Software Structure and Design*. Prentice Hall.",
        "Meyer, B. (1988). *Object-Oriented Software Construction*. Prentice Hall.",
        "Parnas, D. L. (1972). On the criteria to be used in decomposing systems into modules. *Communications of the ACM*, 15(12), 1053–1058. https://doi.org/10.1145/361598.361623",
        "PHP Documentation Group. (2025). *PHP Manual: Classes and Objects (PHP 8.0 – 8.4)*. https://www.php.net/manual/en/language.oop5.php",
        "PHP Documentation Group. (2025). *PHP Manual: Predefined Exceptions and Standard PHP Library (SPL)*. https://www.php.net/manual/en/spl.exceptions.php",
        "PHP Foundation. (2024). *PHP 8.4 Release Announcement and Core Language Enhancements*. https://www.php.net/releases/8.4/en.php",
        "PHP-FIG. (2012). *PSR-3: Logger Interface Specification*. https://www.php-fig.org/psr/psr-3/",
        "PHP-FIG. (2014). *PSR-4: Autoloader Standard Specification*. https://www.php-fig.org/psr/psr-4/",
        "PHP-FIG. (2019). *PSR-12: Extended Coding Style Guide*. https://www.php-fig.org/psr/psr-12/",
        "Saradhna, M. (2021). *Mastering PHP 8: Build enterprise-ready applications using the new features, architecture, and design patterns*. Packt Publishing.",
        "Seemann, M., & Steven, M. (2019). *Dependency Injection Principles, Practices, and Patterns*. Manning Publications.",
        "Skvorc, B. (2017). *PHP Application Development with Composer, PHP-FIG, and More*. SitePoint.",
        "Zandstra, M. (2021). *PHP 8 Objects, Patterns, and Practice: Mastering OO Enhancements, Design Patterns, and Test-Driven Development* (6th ed.). Apress. https://doi.org/10.1007/978-1-4842-6791-2"
    ]
    for ref in pustaka_list:
        p = builder.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        # Render markdown *italic*
        parts = ref.split('*')
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            if i % 2 == 1:
                run.font.italic = True

    builder.doc.add_page_break()
    p_bio = builder.doc.add_paragraph()
    p_bio.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_bio.add_run("TENTANG PENULIS")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    bio_text = (
        "Mahendar Dwi Payana, S.ST., M.T. adalah dosen tetap pada Program Studi Informatika, Fakultas Sains dan Teknologi, "
        "Universitas Ubudiyah Indonesia (UUI), Banda Aceh. Beliau aktif mengampu mata kuliah dalam rumpun Rekayasa Perangkat Lunak, "
        "Pemrograman Berorientasi Objek, Pemrograman Web, Pemrograman Mobile, serta Algoritma dan Struktur Data.\n\n"
        "Penulis memiliki komitmen mendalam dalam pengembangan bahan ajar modern, kurikulum berbasis Outcome-Based Education (OBE), "
        "serta integrasi standar industri ke dalam perkuliahan akademik. Buku ajar ini merupakan wujud dedikasi penulis dalam "
        "membekali generasi mahasiswa dan calon rekayasawan perangkat lunak dengan fondasi keilmuan yang kokoh dan relevan dengan "
        "kemajuan teknologi dunia."
    )
    p_b = builder.doc.add_paragraph(bio_text)
    p_b.paragraph_format.first_line_indent = Cm(0.75)
    p_b.paragraph_format.line_spacing = 1.2

    # Save
    builder.save()
    shutil.copy2(output_path, latest_path)
    print(f"-> Salinan rilis terkini: {latest_path}")
    print("=" * 65)
    print(f" SELESAI: Dokumen '{filename}' berhasil digenerate!")
    print("=" * 65)

if __name__ == "__main__":
    args = parse_args()
    build_oop_php_book(version=args.version, output_dir=args.output_dir)
