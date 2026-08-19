#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Shared Academic Book Builder Core Engine for Universitas Ubudiyah Indonesia (UUI).
Author: Mahendar Dwi Payana, S.ST., M.T.

Features:
- Standard Academic Book A4 Page Setup (Left: 4cm for Binding, Right: 3cm, Top: 3cm, Bottom: 3cm)
- Professional Academic Metadata & KDT (Katalog Dalam Terbitan) Box
- High-Precision Table of Contents (Daftar Isi) Table Engine
- Prominent Green Highlight Callout Boxes (Tips Dosen / Petunjuk Praktik)
- Multi-Language Dark Monospaced Code Blocks (PHP, Java, Go, JavaScript, Python, etc.) strictly Left-Aligned
- Academic Table Styling with Deep Navy (#1E3A8A) Header & Zebra Striping
"""

import os
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets inner padding for a table cell in dxa."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """Sets specific borders on a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            val = edge_data.get('val', 'single')
            color = edge_data.get('color', 'auto')
            sz = edge_data.get('sz', '4')
            tag = f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            tcBorders.append(parse_xml(tag))
        else:
            tag = f'<w:{edge} {nsdecls("w")} w:val="none"/>'
            tcBorders.append(parse_xml(tag))
    tcPr.append(tcBorders)

def render_highlighted_code(paragraph, code_text, lang="php"):
    """Renders syntax-colored code runs inside a left-aligned monospaced paragraph."""
    COLOR_DEFAULT = RGBColor(0xF8, 0xFA, 0xFC)   # Light white
    COLOR_KEYWORD = RGBColor(0x38, 0xBD, 0xF8)   # Cyan / Sky blue (keywords)
    COLOR_TYPE    = RGBColor(0xFB, 0xBF, 0x24)   # Amber (types/class)
    COLOR_STRING  = RGBColor(0x4A, 0xDE, 0x80)   # Emerald green (strings)
    COLOR_COMMENT = RGBColor(0x94, 0xA3, 0xB8)   # Slate gray (comments)
    COLOR_VAR     = RGBColor(0xF4, 0x72, 0xB6)   # Pink/Magenta (variables)
    COLOR_OPEN    = RGBColor(0xFB, 0x71, 0x85)   # Rose (tags)

    keywords = {
        'class', 'interface', 'trait', 'enum', 'struct', 'package', 'import', 'extends', 'implements',
        'public', 'private', 'protected', 'readonly', 'final', 'abstract', 'static', 'const', 'var', 'let',
        'function', 'func', 'def', 'return', 'new', 'clone', 'use', 'namespace',
        'try', 'catch', 'finally', 'throw', 'throws', 'if', 'else', 'elseif', 'switch', 'case', 'match',
        'foreach', 'while', 'for', 'range', 'go', 'chan', 'select', 'defer',
        'declare', 'strict_types', 'as', 'instanceof', 'self', 'parent', 'this',
        'default', 'exit', 'die', 'async', 'await', 'type'
    }
    types = {
        'string', 'int', 'int32', 'int64', 'float', 'float32', 'float64', 'bool', 'boolean',
        'void', 'array', 'mixed', 'null', 'true', 'false', 'object', 'never', 'any', 'byte', 'rune', 'error'
    }

    lines = code_text.strip().split('\n')
    for line_idx, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith('//') or stripped.startswith('#') or stripped.startswith('*') or stripped.startswith('/**'):
            run = paragraph.add_run(line + ('\n' if line_idx < len(lines) - 1 else ''))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.0)
            run.font.italic = True
            run.font.color.rgb = COLOR_COMMENT
            continue

        pattern = r'("(?:\\.|[^"\\])*"|\'(?:\\.|[^\'\\])*\'|`[^`]*`|//.*$|\$[a-zA-Z_\x7f-\xff][a-zA-Z0-9_\x7f-\xff]*|[a-zA-Z_\x7f-\xff][a-zA-Z0-9_\x7f-\xff]*|[^\s\w\$]+|\s+)'
        tokens = re.split(pattern, line)
        for token in tokens:
            if not token:
                continue
            run = paragraph.add_run(token)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.0)
            
            if token.startswith('//') or token.startswith('#'):
                run.font.italic = True
                run.font.color.rgb = COLOR_COMMENT
            elif (token.startswith('"') and token.endswith('"')) or (token.startswith("'") and token.endswith("'")) or (token.startswith('`') and token.endswith('`')):
                run.font.color.rgb = COLOR_STRING
            elif token.startswith('$'):
                run.font.color.rgb = COLOR_VAR
            elif token in keywords:
                run.font.bold = True
                run.font.color.rgb = COLOR_KEYWORD
            elif token in types:
                run.font.bold = True
                run.font.color.rgb = COLOR_TYPE
            elif token in ('<?php', '?>'):
                run.font.bold = True
                run.font.color.rgb = COLOR_OPEN
            else:
                run.font.color.rgb = COLOR_DEFAULT
                
        if line_idx < len(lines) - 1:
            run_nl = paragraph.add_run('\n')
            run_nl.font.name = 'Consolas'
            run_nl.font.size = Pt(9.0)

def add_code_block(doc, code_text, lang="php"):
    """Adds a modern dark IDE-styled code block. STRICTLY LEFT ALIGNED."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Cm(14.0)
    set_cell_background(cell, "1E293B")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    left={'val': 'single', 'sz': '24', 'color': '38BDF8'},
                    top={'val': 'single', 'sz': '4', 'color': '0F172A'},
                    right={'val': 'single', 'sz': '4', 'color': '0F172A'},
                    bottom={'val': 'single', 'sz': '4', 'color': '0F172A'})
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # NEVER JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    
    render_highlighted_code(p, code_text, lang=lang)
    
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(4)

def add_callout_box(doc, title, text, box_type="tip"):
    """Adds a callout box. 'tip' and 'highlight' have prominent green highlight."""
    colors = {
        "tip": {"bg": "ECFDF5", "border": "059669", "icon": "🎯", "title_color": "047857", "text_color": "064E3B"},
        "highlight": {"bg": "F0FDF4", "border": "10B981", "icon": "💡", "title_color": "047857", "text_color": "065F46"},
        "info": {"bg": "EFF6FF", "border": "2563EB", "icon": "📘", "title_color": "1D4ED8", "text_color": "1E3A8A"},
        "warning": {"bg": "FFFBEB", "border": "F59E0B", "icon": "⚠️", "title_color": "B45309", "text_color": "78350F"},
        "caution": {"bg": "FEF2F2", "border": "EF4444", "icon": "🛑", "title_color": "B91C1C", "text_color": "7F1D1D"},
    }
    style = colors.get(box_type, colors["tip"])
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Cm(14.0)
    set_cell_background(cell, style["bg"])
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, 
                    left={'val': 'single', 'sz': '28', 'color': style["border"]},
                    top={'val': 'single', 'sz': '4', 'color': 'D1FAE5' if "F" in style["bg"] else 'E2E8F0'},
                    right={'val': 'single', 'sz': '4', 'color': 'D1FAE5' if "F" in style["bg"] else 'E2E8F0'},
                    bottom={'val': 'single', 'sz': '4', 'color': 'D1FAE5' if "F" in style["bg"] else 'E2E8F0'})
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(f"{style['icon']} {title}\n")
    r_title.bold = True
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor.from_string(style["title_color"])
    
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(10.0)
    r_text.font.italic = False
    r_text.font.color.rgb = RGBColor.from_string(style["text_color"])
    
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(4)

def add_academic_table(doc, headers, rows_data):
    """Creates a cleanly formatted academic table."""
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        set_cell_border(hdr_cells[i], 
                        bottom={'val': 'single', 'sz': '12', 'color': '0F172A'},
                        top={'val': 'single', 'sz': '4', 'color': '1E3A8A'},
                        left={'val': 'none'}, right={'val': 'none'})
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.0)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    for r_idx, row in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            set_cell_border(row_cells[c_idx], 
                            bottom={'val': 'single', 'sz': '4', 'color': 'E2E8F0'},
                            top={'val': 'none'}, left={'val': 'none'}, right={'val': 'none'})
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(4)

class AcademicBookBuilder:
    """Master Academic Textbook Builder Class."""
    def __init__(self, output_path, book_title, version="v1.0.0", course_code="IFR 214"):
        self.output_path = output_path
        self.book_title = book_title
        self.version = version
        self.course_code = course_code
        self.doc = Document()
        self._setup_page_layout()
        self._setup_styles()

    def _setup_page_layout(self):
        """Setup A4 page size and academic book margins (Left: 4cm, Right: 3cm, Top: 3cm, Bottom: 3cm)."""
        sections = self.doc.sections
        for section in sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(3.0)
            section.bottom_margin = Cm(3.0)
            section.left_margin = Cm(4.0)   # 4.0 cm for binding/jilid
            section.right_margin = Cm(3.0)
            
            # Header
            header = section.header
            p_head = header.paragraphs[0]
            p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_head.paragraph_format.space_after = Pt(0)
            r_head = p_head.add_run(f"{self.book_title} | {self.version}")
            r_head.font.name = "Calibri"
            r_head.font.size = Pt(8.5)
            r_head.font.italic = True
            r_head.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            
            # Footer
            footer = section.footer
            p_foot = footer.paragraphs[0]
            p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_foot.paragraph_format.space_after = Pt(0)
            r_foot = p_foot.add_run("Program Studi Informatika — Universitas Ubudiyah Indonesia")
            r_foot.font.name = "Calibri"
            r_foot.font.size = Pt(8.5)
            r_foot.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    def _setup_styles(self):
        """Setup default typography and paragraph formatting."""
        style_normal = self.doc.styles['Normal']
        style_normal.font.name = 'Calibri'
        style_normal.font.size = Pt(11.0)
        style_normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        style_normal.paragraph_format.line_spacing = 1.15
        style_normal.paragraph_format.space_after = Pt(4.0)
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_title_page(self, main_title, subtitle, author="Mahendar Dwi Payana, S.ST., M.T.", year="2025"):
        p_space = self.doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(30)
        
        p_badge = self.doc.add_paragraph()
        p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_badge = p_badge.add_run("BUKU AJAR PERKULIAHAN")
        r_badge.bold = True
        r_badge.font.name = 'Calibri'
        r_badge.font.size = Pt(12.0)
        r_badge.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(10)
        p_title.paragraph_format.space_after = Pt(8)
        r_title = p_title.add_run(main_title.upper())
        r_title.bold = True
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(22.0)
        r_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        p_sub = self.doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(40)
        r_sub = p_sub.add_run(subtitle)
        r_sub.font.name = 'Calibri'
        r_sub.font.size = Pt(12.0)
        r_sub.font.italic = True
        r_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        
        p_line = self.doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line.paragraph_format.space_after = Pt(40)
        r_line = p_line.add_run("____________________________________________________")
        r_line.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD)
        
        p_author_label = self.doc.add_paragraph()
        p_author_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_author_label = p_author_label.add_run("Penulis:")
        r_author_label.font.name = 'Calibri'
        r_author_label.font.size = Pt(11.0)
        r_author_label.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        
        p_author = self.doc.add_paragraph()
        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_author.paragraph_format.space_after = Pt(70)
        r_author = p_author.add_run(author)
        r_author.bold = True
        r_author.font.name = 'Calibri'
        r_author.font.size = Pt(14.0)
        r_author.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        p_inst = self.doc.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_inst.paragraph_format.space_after = Pt(2)
        r_inst = p_inst.add_run(f"PROGRAM STUDI INFORMATIKA\nFAKULTAS SAINS DAN TEKNOLOGI\nUNIVERSITAS UBUDIYAH INDONESIA\nBANDA ACEH\n{year}")
        r_inst.bold = True
        r_inst.font.name = 'Calibri'
        r_inst.font.size = Pt(11.0)
        r_inst.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
        self.doc.add_page_break()

    def add_copyright_page(self, meta_info, kdt_data):
        p_h = self.doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(10)
        p_h.paragraph_format.space_after = Pt(10)
        r = p_h.add_run("INFORMASI PENERBITAN & HAK CIPTA")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        # 2-Column Metadata Table
        tbl = self.doc.add_table(rows=len(meta_info), cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        for r_i, (k, v) in enumerate(meta_info):
            c0 = tbl.cell(r_i, 0)
            c1 = tbl.cell(r_i, 1)
            c0.width = Cm(4.2)
            c1.width = Cm(9.8)
            set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
            set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
            set_cell_border(c0, bottom={'val': 'single', 'sz': '4', 'color': 'E2E8F0'}, top={'val': 'none'}, left={'val': 'none'}, right={'val': 'none'})
            set_cell_border(c1, bottom={'val': 'single', 'sz': '4', 'color': 'E2E8F0'}, top={'val': 'none'}, left={'val': 'none'}, right={'val': 'none'})
            
            p0 = c0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r0 = p0.add_run(k)
            r0.font.name = 'Calibri'
            r0.font.size = Pt(9.5)
            r0.bold = True
            r0.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            
            p1 = c1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r1 = p1.add_run(v)
            r1.font.name = 'Calibri'
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        p_sp = self.doc.add_paragraph()
        p_sp.paragraph_format.space_before = Pt(8)
        p_sp.paragraph_format.space_after = Pt(4)

        # KDT Box
        kdt_table = self.doc.add_table(rows=1, cols=1)
        kdt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        kdt_table.autofit = False
        kdt_cell = kdt_table.cell(0, 0)
        kdt_cell.width = Cm(14.0)
        set_cell_background(kdt_cell, "F8FAFC")
        set_cell_margins(kdt_cell, top=140, bottom=140, left=180, right=180)
        set_cell_border(kdt_cell, 
                        top={'val': 'single', 'sz': '12', 'color': '64748B'},
                        bottom={'val': 'single', 'sz': '12', 'color': '64748B'},
                        left={'val': 'single', 'sz': '12', 'color': '64748B'},
                        right={'val': 'single', 'sz': '12', 'color': '64748B'})
        
        p_kdt = kdt_cell.paragraphs[0]
        p_kdt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_kdt.paragraph_format.line_spacing = 1.15
        
        r_kdt_title = p_kdt.add_run("Katalog Dalam Terbitan (KDT)\n\n")
        r_kdt_title.bold = True
        r_kdt_title.font.name = 'Calibri'
        r_kdt_title.font.size = Pt(10.0)
        r_kdt_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        r_kdt_body = p_kdt.add_run(kdt_data)
        r_kdt_body.font.name = 'Consolas'
        r_kdt_body.font.size = Pt(8.5)
        r_kdt_body.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        p_sp2 = self.doc.add_paragraph()
        p_sp2.paragraph_format.space_before = Pt(8)
        p_sp2.paragraph_format.space_after = Pt(2)

        # UU Hak Cipta Box
        uu_table = self.doc.add_table(rows=1, cols=1)
        uu_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        uu_table.autofit = False
        uu_cell = uu_table.cell(0, 0)
        uu_cell.width = Cm(14.0)
        set_cell_background(uu_cell, "FEF2F2")
        set_cell_margins(uu_cell, top=100, bottom=100, left=150, right=150)
        set_cell_border(uu_cell, left={'val': 'single', 'sz': '24', 'color': 'EF4444'},
                        top={'val': 'single', 'sz': '4', 'color': 'FECACA'},
                        right={'val': 'single', 'sz': '4', 'color': 'FECACA'},
                        bottom={'val': 'single', 'sz': '4', 'color': 'FECACA'})
        
        p_uu = uu_cell.paragraphs[0]
        p_uu.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_uu_t = p_uu.add_run("PERINGATAN HAK CIPTA & SANKSI PELANGGARAN:\n")
        r_uu_t.bold = True
        r_uu_t.font.name = 'Calibri'
        r_uu_t.font.size = Pt(8.5)
        r_uu_t.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
        
        uu_desc = (
            "Hak Cipta © 2025 pada Penulis. Hak Penerbitan pada UUI Press. "
            "Dilarang memperbanyak atau memindahkan sebagian atau seluruh isi buku ini dalam bentuk apa pun, "
            "baik secara elektronik maupun mekanik, termasuk memfotokopi, merekam, atau dengan sistem penyimpanan "
            "lainnya tanpa izin tertulis dari Penulis dan Penerbit.\n"
            "Sanksi Pelanggaran Pasal 113 UU No. 28 Tahun 2014: Setiap orang yang dengan tanpa hak melakukan "
            "pelanggaran hak ekonomi dapat dipidana dengan pidana penjara paling lama 10 (sepuluh) tahun "
            "dan/atau denda paling banyak Rp 4.000.000.000,00 (empat miliar rupiah)."
        )
        r_uu_d = p_uu.add_run(uu_desc)
        r_uu_d.font.name = 'Calibri'
        r_uu_d.font.size = Pt(8.0)
        r_uu_d.font.color.rgb = RGBColor(0x7F, 0x1D, 0x1D)
        
        self.doc.add_page_break()

    def add_preface(self, paragraphs, author="Mahendar Dwi Payana, S.ST., M.T.", location="Banda Aceh, 2025"):
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(10)
        p_title.paragraph_format.space_after = Pt(16)
        r = p_title.add_run("KATA PENGANTAR")
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        for para in paragraphs:
            p = self.doc.add_paragraph(para)
            p.paragraph_format.first_line_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(8)
            
        p_sign = self.doc.add_paragraph()
        p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sign.paragraph_format.space_before = Pt(20)
        p_sign.add_run(f"{location}\n\n\n\n").font.italic = True
        r_penulis = p_sign.add_run(author)
        r_penulis.bold = True
        
        self.doc.add_page_break()

    def add_table_of_contents(self, toc_entries):
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(10)
        p_title.paragraph_format.space_after = Pt(14)
        r = p_title.add_run("DAFTAR ISI")
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        toc_table = self.doc.add_table(rows=len(toc_entries), cols=2)
        toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        toc_table.autofit = False
        
        for idx, (entry_type, title, page) in enumerate(toc_entries):
            cell_title = toc_table.cell(idx, 0)
            cell_page = toc_table.cell(idx, 1)
            cell_title.width = Cm(12.5)
            cell_page.width = Cm(1.5)
            
            set_cell_border(cell_title, top={'val': 'none'}, bottom={'val': 'none'}, left={'val': 'none'}, right={'val': 'none'})
            set_cell_border(cell_page, top={'val': 'none'}, bottom={'val': 'none'}, left={'val': 'none'}, right={'val': 'none'})
            
            p_t = cell_title.paragraphs[0]
            p_p = cell_page.paragraphs[0]
            p_t.paragraph_format.space_before = Pt(0)
            p_t.paragraph_format.space_after = Pt(1)
            p_p.paragraph_format.space_before = Pt(0)
            p_p.paragraph_format.space_after = Pt(1)
            p_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            if entry_type == 'SECTION':
                set_cell_background(cell_title, "EFF6FF")
                set_cell_background(cell_page, "EFF6FF")
                set_cell_margins(cell_title, top=80, bottom=80, left=100, right=60)
                set_cell_margins(cell_page, top=80, bottom=80, left=60, right=100)
                
                r_sec = p_t.add_run(f"■ {title}")
                r_sec.bold = True
                r_sec.font.name = 'Calibri'
                r_sec.font.size = Pt(10.0)
                r_sec.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            elif entry_type == 'BAB':
                set_cell_margins(cell_title, top=40, bottom=20, left=60, right=40)
                set_cell_margins(cell_page, top=40, bottom=20, left=40, right=60)
                
                r_t = p_t.add_run(title)
                r_t.bold = True
                r_t.font.name = 'Calibri'
                r_t.font.size = Pt(9.5)
                r_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                
                r_p = p_p.add_run(page)
                r_p.bold = True
                r_p.font.name = 'Calibri'
                r_p.font.size = Pt(9.5)
                r_p.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            elif entry_type == 'SUB':
                set_cell_margins(cell_title, top=10, bottom=10, left=160, right=40)
                set_cell_margins(cell_page, top=10, bottom=10, left=40, right=60)
                
                r_t = p_t.add_run("   " + title)
                r_t.font.name = 'Calibri'
                r_t.font.size = Pt(8.5)
                r_t.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                
                r_p = p_p.add_run(page)
                r_p.font.name = 'Calibri'
                r_p.font.size = Pt(8.5)
                r_p.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            else: # PRE or POST
                set_cell_margins(cell_title, top=20, bottom=20, left=60, right=40)
                set_cell_margins(cell_page, top=20, bottom=20, left=40, right=60)
                
                r_t = p_t.add_run(title)
                r_t.bold = (entry_type == 'POST')
                r_t.font.name = 'Calibri'
                r_t.font.size = Pt(9.0)
                r_t.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                
                r_p = p_p.add_run(page)
                r_p.bold = (entry_type == 'POST')
                r_p.font.name = 'Calibri'
                r_p.font.size = Pt(9.0)
                r_p.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
        self.doc.add_page_break()

    def add_bab_title(self, bab_num, title):
        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(2)
        r_bab = p.add_run(f"BAB {bab_num}\n")
        r_bab.bold = True
        r_bab.font.size = Pt(14)
        r_bab.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        
        r_title = p.add_run(title.upper())
        r_title.bold = True
        r_title.font.size = Pt(16)
        r_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        p_div = self.doc.add_paragraph()
        p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_div.paragraph_format.space_after = Pt(14)
        r_div = p_div.add_run("═" * 45)
        r_div.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD)

    def add_heading_2(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12.5)
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    def add_heading_3(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11.0)
        r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    def add_paragraph(self, text, indent=False):
        p = self.doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11.0)
        r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        return p

    def add_bullet(self, bold_prefix, text):
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r_pre = p.add_run(bold_prefix + ": ")
            r_pre.bold = True
            r_pre.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        r_text = p.add_run(text)
        r_text.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    def add_code(self, code_str, lang="php"):
        add_code_block(self.doc, code_str, lang=lang)

    def add_callout(self, title, text, box_type="tip"):
        add_callout_box(self.doc, title, text, box_type)

    def add_tip(self, title, text):
        add_callout_box(self.doc, title, text, box_type="tip")

    def add_table(self, headers, rows):
        add_academic_table(self.doc, headers, rows)

    def add_learning_objectives(self, sub_cpmk, objectives):
        self.add_callout(
            f"Tujuan & Capaian Pembelajaran ({sub_cpmk})",
            "Setelah mempelajari dan menyelesaikan seluruh materi pada bab ini, mahasiswa diharapkan memiliki kompetensi untuk:\n" +
            "\n".join([f"• {obj}" for obj in objectives]),
            box_type="tip"
        )

    def add_summary_and_questions(self, summary_points, questions):
        self.add_heading_2("Rangkuman Materi")
        for pt in summary_points:
            self.add_bullet("", pt)
            
        self.add_heading_2("Soal Latihan & Penugasan Mandiri")
        for i, q in enumerate(questions, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            r_num = p.add_run(f"{i}. ")
            r_num.bold = True
            r_num.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            p.add_run(q)

    def save(self):
        os.makedirs(os.path.dirname(os.path.abspath(self.output_path)), exist_ok=True)
        self.doc.save(self.output_path)
        print(f"Buku berhasil digenerate dan disimpan ke: {self.output_path}")

print("AcademicBookBuilder core module loaded.")
